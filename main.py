import os
import uuid
import json
from typing import Optional
import logging
import random
import re
import asyncio
from datetime import datetime, timedelta
from typing import Union, List, Dict, Any
from fastapi import FastAPI, HTTPException, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi import UploadFile, File
from pydantic import BaseModel
from openai import OpenAI
from supabase import create_client, Client
from pywebpush import webpush, WebPushException
import httpx


# =====================================================
# FASTAPI INITIALIZATION
# =====================================================
app = FastAPI()
pending_emails = {}

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://sovereignallmighty.netlify.app",
        "http://localhost:3000",
        "https://sovereign-app-ashen.vercel.app"
    ],
    allow_origin_regex=r"https://.*\.vercel\.app",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# =====================================================
# WHATSAPP WEBHOOK - PLACÉ ICI APRÈS CORS
# =====================================================

# Délai de réponse aléatoire entre 1 et 2 minutes (pour faire naturel)
MIN_REPLY_DELAY = 60   # 1 minute en secondes
MAX_REPLY_DELAY = 120  # 2 minutes en secondes

# Ticks de langage naturels (style Rebecca)
THINKING_PREFIXES = [
    "", "Mmh ", "Euh ", "Alors ", "Attends ", "Je réfléchis... ", 
    "Là tout de suite ", "Franchement ", "Je dirais ", "En vrai "
]

EMOJIS = ["", "✨", "👌", "🙏", "😊", "❤️", "🌱"]

def naturalize_response(text: str) -> str:
    """Rend la réponse plus naturelle, moins robotique"""
    import random
    if random.random() < 0.3:
        prefix = random.choice(THINKING_PREFIXES)
        text = prefix + text[0].lower() + text[1:] if text else text
    if random.random() < 0.25:
        text += " " + random.choice(EMOJIS)
    text = text.replace("Je suis désolé", "Désolée")
    text = text.replace("Je ne peux pas", "Je peux pas")
    text = text.replace("Je vais", "Je")
    text = text.replace("Souhaitez-vous", "Tu veux")
    text = text.replace("Pouvez-vous", "Tu peux")
    text = text.replace("Cordialement", "")
    return text.strip()


# =====================================================
# WHATSAPP WEBHOOK
# =====================================================

# ========== WEBHOOK PRINCIPAL ==========
@app.api_route("/api/whatsapp/webhook", methods=["POST", "OPTIONS"])
async def whatsapp_webhook(request: Request):
    """Webhook principal pour recevoir les messages WhatsApp"""
    
    if request.method == "OPTIONS":
        return Response(status_code=200)
    
    body = await request.body()
    
    try:
        data = json.loads(body.decode('utf-8'))
    except json.JSONDecodeError:
        data = {}
    
    # Ignorer les statuts d'envoi
    if data.get("typeWebhook") in ["outgoingMessageStatus", "outgoingAPIMessageReceived", "quotaExceeded"]:
        return {"status": "ok"}
    
    # Traiter uniquement les messages entrants
    if data.get("typeWebhook") == "incomingMessageReceived":
        await process_incoming_message(data)
    
    return {"status": "ok"}


async def process_incoming_message(data: dict):
    """Traite un message entrant"""
    
    message_data = data.get("messageData", {})
    message_type = message_data.get("typeMessage", "")
    
    sender_data = data.get("senderData", {})
    sender = sender_data.get("sender", "")
    sender_name = sender_data.get("senderName", "Inconnu")
    chat_id = sender_data.get("chatId", sender)
    
    # Ignorer les groupes
    if is_group_message(chat_id, sender):
        print(f"⏭️ Message de groupe ignoré: {chat_id}")
        return
    
    # Ignorer les réactions
    if message_type == "reactionMessage":
        print(f"⏭️ Réaction ignorée de {sender_name}")
        return
    
    # Extraire le contenu du message
    text_message, attachment_url, attachment_type = await extract_message_content(
        message_type, message_data, sender_name
    )
    
    if not text_message:
        return
    
    # Analyser et notifier (ou sauvegarder directement pour les fichiers)
    if is_file_message(text_message):
        await save_message_direct(sender, sender_name, text_message, attachment_url, attachment_type)
    else:
        await analyze_and_notify(sender, sender_name, text_message)


def is_group_message(chat_id: str, sender: str) -> bool:
    """Vérifie si le message provient d'un groupe"""
    return any(identifier in chat_id or identifier in sender 
              for identifier in ["@g.us"])


async def extract_message_content(message_type: str, message_data: dict, sender_name: str):
    """Extrait le contenu du message selon son type"""
    
    text_message = ""
    attachment_url = None
    attachment_type = None
    
    if message_type == "textMessage":
        text_message = message_data.get("textMessageData", {}).get("textMessage", "")
        print(f"💬 [{sender_name}]: {text_message}")
    
    elif message_type == "extendedTextMessage":
        extended_data = message_data.get("extendedTextMessageData", {})
        text_message = extended_data.get("text") or extended_data.get("caption", "")
        print(f"💬 [{sender_name}]: {text_message}")
    
    elif message_type == "audioMessage":
        text_message = await process_audio_message(message_data, sender_name)
    
    elif message_type == "imageMessage":
        image_data = message_data.get("imageMessageData", {})
        caption = image_data.get("caption", "")
        text_message = f"🖼️ [IMAGE] {caption if caption else 'Image sans légende'}"
        attachment_url = image_data.get("url")
        attachment_type = "image"
        print(f"🖼️ Image de {sender_name}: {caption}")
    
    elif message_type == "documentMessage":
        doc_data = message_data.get("documentMessageData", {})
        file_name = doc_data.get("fileName", "Document")
        text_message = f"📎 [DOCUMENT] {file_name}"
        attachment_url = doc_data.get("url")
        attachment_type = "document"
        print(f"📎 Document de {sender_name}: {file_name}")
    
    elif message_type == "videoMessage":
        video_data = message_data.get("videoMessageData", {})
        caption = video_data.get("caption", "")
        text_message = f"🎥 [VIDEO] {caption if caption else 'Vidéo sans légende'}"
        attachment_url = video_data.get("url")
        attachment_type = "video"
        print(f"🎥 Vidéo de {sender_name}")
    
    else:
        print(f"📎 Autre type non supporté: {message_type}")
    
    return text_message, attachment_url, attachment_type


async def process_audio_message(message_data: dict, sender_name: str) -> str:
    """Traite un message audio avec transcription Whisper"""
    
    audio_data = message_data.get("audioMessageData", {})
    audio_url = audio_data.get("url")
    audio_duration = audio_data.get("duration", 0)
    
    print(f"🎤 Message vocal de {sender_name} ({audio_duration}s)")
    
    if not audio_url or not supabase:
        return "🎤 [Message vocal non transcrit]"
    
    try:
        async with httpx.AsyncClient() as client:
            audio_response = await client.get(audio_url)
            audio_content = audio_response.content
        
        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
            tmp.write(audio_content)
            tmp_path = tmp.name
        
        with open(tmp_path, "rb") as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1", file=audio_file, language="fr"
            )
        
        os.unlink(tmp_path)
        transcribed_text = transcript.text
        print(f"📝 Transcription: {transcribed_text}")
        return f"🎤 [VOCAL] {transcribed_text}"
    
    except Exception as e:
        print(f"❌ Erreur transcription: {e}")
        return "🎤 [Message vocal non transcrit]"


def is_file_message(text_message: str) -> bool:
    """Vérifie si le message est un fichier (image, document, vidéo)"""
    return any(text_message.startswith(prefix) for prefix in ["🖼️", "📎", "🎥"])


async def save_message_direct(sender: str, sender_name: str, text_message: str, 
                               attachment_url: str = None, attachment_type: str = None):
    """Sauvegarde directe d'un message sans analyse (pour les fichiers)"""
    
    if not supabase:
        return
    
    try:
        message_data = {
            "from": sender,
            "from_name": sender_name,
            "message": text_message,
            "status": "pending",
            "replied": False,  # ← AJOUTÉ
            "importance": "medium",
            "created_at": datetime.now().isoformat()
        }
        
        if attachment_url:
            message_data["attachment_url"] = attachment_url
            message_data["attachment_type"] = attachment_type
        
        supabase.table("whatsapp_messages").insert(message_data).execute()
        print(f"✅ Fichier sauvegardé")
    
    except Exception as e:
        print(f"❌ Erreur sauvegarde fichier: {e}")


async def analyze_and_notify(sender: str, sender_name: str, text_message: str):
    """Analyse le message, sauvegarde et envoie une notification si important"""
    
    try:
        analysis = await analyze_message_importance(sender_name, text_message)
        
        # Sauvegarde avec replied = False par défaut
        if supabase:
            supabase.table("whatsapp_messages").insert({
                "from": sender,
                "from_name": sender_name,
                "message": text_message,
                "status": "pending",
                "replied": False,  # ← AJOUTÉ
                "importance": analysis["importance"],
                "summary": analysis["summary"],
                "is_greeting": analysis["is_greeting"],
                "created_at": datetime.now().isoformat()
            }).execute()
            print(f"✅ Message sauvegardé (importance: {analysis['importance']})")
        
        # Envoyer une notification si important
        if should_notify(analysis):
            send_notification_sync({
                "title": f"📱 WhatsApp - {sender_name}",
                "body": analysis["summary"][:80],
                "url": "/chat?mode=whatsapp",
                "user_id": DEFAULT_USER_ID,
                "type": "whatsapp"
            })
            print(f"🔔 Notification envoyée pour message {analysis['importance']}")
        else:
            print(f"⏭️ Message de faible importance, pas de notification")
    
    except Exception as e:
        print(f"❌ Erreur analyse: {e}")
        await fallback_save_and_notify(sender, sender_name, text_message)


async def analyze_message_importance(sender_name: str, text_message: str) -> dict:
    """Analyse l'importance du message via GPT"""
    
    analysis_prompt = f"""Message WhatsApp de {sender_name}: "{text_message}"
Analyse ce message et retourne UNIQUEMENT du JSON:
{{
  "importance": "high/medium/low",
  "summary": "résumé court du message (max 20 mots)",
  "is_greeting": true/false
}}

Règles:
- importance: high = contient une question, une demande, une urgence, un document, un délai
- importance: medium = contient une information, un partage, une mise à jour
- importance: low = simple salutation, emoji, réponse courte sans attente
- is_greeting: true si juste "cc", "bonjour", "salut", "coucou" sans contenu"""

    analysis = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": analysis_prompt}],
        max_tokens=150,
        temperature=0.5
    )
    
    result_text = analysis.choices[0].message.content
    result_text = result_text.replace("```json", "").replace("```", "").strip()
    
    return json.loads(result_text)


def should_notify(analysis: dict) -> bool:
    """Détermine si une notification doit être envoyée"""
    importance = analysis.get("importance", "medium")
    is_greeting = analysis.get("is_greeting", False)
    
    return importance == "high" or (importance == "medium" and not is_greeting)


async def fallback_save_and_notify(sender: str, sender_name: str, text_message: str):
    """Sauvegarde de secours sans analyse"""
    
    if supabase:
        supabase.table("whatsapp_messages").insert({
            "from": sender,
            "from_name": sender_name,
            "message": text_message,
            "status": "pending",
            "replied": False,  # ← AJOUTÉ
            "importance": "medium",
            "created_at": datetime.now().isoformat()
        }).execute()
        print(f"✅ Message sauvegardé (fallback)")
    
    # Envoyer une notification quand même
    send_notification_sync({
        "title": f"📱 WhatsApp - {sender_name}",
        "body": text_message[:100],
        "url": "/chat?mode=whatsapp",
        "user_id": DEFAULT_USER_ID,
        "type": "whatsapp"
    })

# =====================================================
# WHATSAPP - NOTIFICATIONS PROGRAMMÉES
# =====================================================

@app.post("/api/whatsapp/pending-notifications")
async def whatsapp_pending_notifications(request: Dict[str, Any] = None):
    """Envoie un résumé des messages WhatsApp en attente (non répondus)"""
    
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    user_id = get_request_user_id(request or {})
    
    # 🔥 Utiliser replied = False au lieu de status = pending
    pending_messages = supabase.table("whatsapp_messages")\
        .select("*")\
        .eq("user_id", user_id)\
        .eq("replied", False)\
        .order("importance", desc=True)\
        .order("created_at", desc=True)\
        .execute()
    
    if not pending_messages.data:
        return {"success": True, "sent": False, "message": "Aucun message en attente"}
        
    # Séparer par importance
    high_priority = [m for m in pending_messages.data if m.get("importance") == "high"]
    medium_priority = [m for m in pending_messages.data if m.get("importance") == "medium"]
    low_priority = [m for m in pending_messages.data if m.get("importance") == "low"]
    
    # Construire le message
    time_of_day = get_time_of_day()
    message = f"{time_of_day} Rebecca.\n\n"
    
    if high_priority:
        message += f"⚠️ **{len(high_priority)} message(s) important(s) en attente sur WhatsApp** :\n"
        for msg in high_priority[:3]:
            message += f"   • **{msg['from_name']}** : {msg.get('summary', msg['message'][:50])}...\n"
        if len(high_priority) > 3:
            message += f"   • et {len(high_priority) - 3} autre(s) message(s) important(s)\n"
        message += "\n"
    
    if medium_priority:
        message += f"📋 {len(medium_priority)} autre(s) message(s) à consulter\n"
    
    if low_priority and not high_priority and not medium_priority:
        message += f"💬 {len(low_priority)} salutation(s) sans réponse urgente\n"
    
    message += "\n👉 Dis-moi 'affiche mes WhatsApp' pour voir les détails et répondre."
    
    # Envoyer la notification
    send_notification_sync({
        "title": f"📱 WhatsApp - {len(high_priority)} message(s) important(s)",
        "body": message[:200],
        "url": "/chat?mode=whatsapp",
        "user_id": user_id,
        "type": "whatsapp"
    })
    
    return {
        "success": True,
        "sent": True,
        "stats": {
            "high": len(high_priority),
            "medium": len(medium_priority),
            "low": len(low_priority),
            "total": len(pending_messages.data)
        }
    }


def get_time_of_day() -> str:
    """Retourne la salutation selon l'heure"""
    hour = datetime.now().hour
    if hour < 12:
        return "☀️ Bonjour"
    elif hour < 18:
        return "🌤️ Bon après-midi"
    else:
        return "🌙 Bonsoir"


# =====================================================
# WHATSAPP - RÉPONSES
# =====================================================

@app.post("/api/whatsapp/reply")
async def whatsapp_reply(request: Dict[str, Any]):
    """Envoie une réponse WhatsApp (depuis Rebecca)"""
    to = request.get("to")
    message = request.get("message")
    message_id = request.get("message_id")
    
    if not to or not message:
        return {"success": False, "error": "to et message requis"}
    
    # Envoyer via GreenAPI
    success = await whatsapp_send_message(to, message)
    
    if success and supabase:
        if message_id:
            # Marquer le message spécifique comme répondu
            supabase.table("whatsapp_messages").update({
                "status": "replied",
                "replied": True,  # ← AJOUTÉ
                "response": message
            }).eq("id", message_id).execute()
        else:
            # Marquer tous les messages pending de ce contact comme répondu
            supabase.table("whatsapp_messages").update({
                "status": "replied",
                "replied": True,  # ← AJOUTÉ
                "response": message
            }).eq("from", to).eq("status", "pending").execute()
    
    return {"success": success}


@app.get("/api/whatsapp/conversations")
async def get_whatsapp_conversations(days: int = 30):
    """Récupère les conversations WhatsApp du mois (hors groupes)"""
    if not supabase:
        return {"conversations": []}
    
    cutoff_date = (datetime.now() - timedelta(days=days)).isoformat()
    result = supabase.table("whatsapp_messages").select("*").gte("created_at", cutoff_date).order("created_at", desc=True).execute()
    
    conversations = {}
    for msg in result.data:
        sender = msg.get("from", "")
        
        # Exclure les groupes
        if "@g.us" in sender or sender.endswith("@g.us"):
            continue
        
        from_name = msg.get("from_name", "Inconnu")
        if from_name == "Inconnu" or len(from_name) < 2:
            from_name = sender.split("@")[0][:15]
        
        if sender not in conversations:
            conversations[sender] = {
                "from": sender,
                "from_name": from_name,
                "messages": [],
                "pending": 0,
                "last_message_at": msg.get("created_at")
            }
        
        message_text = msg.get("message", "")
        if len(message_text) > 150:
            message_text = message_text[:150] + "..."
        
        conversations[sender]["messages"].append({
            "id": msg.get("id"),
            "message": message_text,
            "status": msg.get("status", "pending"),
            "created_at": msg.get("created_at"),
            "attachment_url": msg.get("attachment_url"),
            "attachment_type": msg.get("attachment_type")
        })
        
        if msg.get("status") == "pending":
            conversations[sender]["pending"] += 1
    
    filtered = [c for c in conversations.values() if c["messages"]]
    filtered.sort(key=lambda x: x["last_message_at"], reverse=True)
    return {"conversations": filtered[:15]}

async def whatsapp_send_message(to: str, message: str):
    """Envoie un message WhatsApp via GreenAPI"""
    print(f"📤 Tentative d'envoi à {to}")
    
    if not GREENAPI_ID_INSTANCE or not GREENAPI_API_TOKEN:
        print("❌ GreenAPI non configuré")
        return False
    
    # 🔥 FORMAT CORRECT : toujours ajouter @c.us
    clean_to = to.replace("+", "").replace(" ", "")
    if not clean_to.endswith("@c.us"):
        clean_to = clean_to + "@c.us"
    
    url = f"{GREENAPI_BASE_URL}/sendMessage/{GREENAPI_API_TOKEN}"
    payload = {"chatId": clean_to, "message": message}
    
    print(f"📤 URL: {url}")
    print(f"📤 Payload: {payload}")
    
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(url, json=payload)
            print(f"📤 Response status: {response.status_code}")
            print(f"📤 Response body: {response.text[:200]}")
            
            if response.status_code == 200:
                result = response.json()
                print(f"✅ Message envoyé à {clean_to}, idMessage: {result.get('idMessage')}")
                return True
            else:
                print(f"❌ Erreur GreenAPI: {response.status_code} - {response.text[:100]}")
                return False
    except Exception as e:
        print(f"❌ Erreur envoi: {type(e).__name__} - {e}")
        return False

@app.post("/api/whatsapp/send-image")
async def whatsapp_send_image(request: Dict[str, Any]):
    """Envoie une image WhatsApp"""
    to = request.get("to")
    image_base64 = request.get("image")
    caption = request.get("caption", "")
    
    if not to or not image_base64:
        return {"success": False, "error": "to et image requis"}
    
    clean_to = to.replace("@c.us", "").replace("+", "")
    
    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(
            f"{GREENAPI_BASE_URL}/sendFileByUpload/{GREENAPI_API_TOKEN}",
            json={"chatId": f"{clean_to}@c.us", "file": image_base64, "fileName": "image.jpg", "caption": caption}
        )
        return {"success": response.status_code == 200}


@app.get("/api/whatsapp/status")
async def whatsapp_status():
    """Vérifie si GreenAPI est configuré"""
    if not GREENAPI_ID_INSTANCE or not GREENAPI_API_TOKEN:
        return {"configured": False}
    return {"configured": True, "idInstance": GREENAPI_ID_INSTANCE}

# =====================================================
# WHATSAPP VIA BAILEYS (gratuit, illimité)
# =====================================================

BAILEYS_URL = os.environ.get("BAILEYS_URL", "http://localhost:10000")

@app.post("/api/whatsapp-baileys/message")
async def baileys_message(request: Dict[str, Any]):
    """Reçoit les messages de Baileys"""
    from_ = request.get("from")
    from_name = request.get("from_name")
    message = request.get("message")
    
    print(f"💬 [Baileys - {from_name}]: {message}")
    
    # Sauvegarde en base
    if supabase:
        supabase.table("whatsapp_messages").insert({
            "from": from_,
            "from_name": from_name,
            "message": message,
            "status": "pending",
            "created_at": datetime.now().isoformat()
        }).execute()
    
    # Analyser avec Becks
    try:
        analysis = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": f"Message WhatsApp de {from_name}: {message}. Réponds brièvement (1-2 phrases max) comme si tu étais Rebecca."}],
            max_tokens=100
        )
        reply = analysis.choices[0].message.content
        
        # Envoyer la réponse via Baileys
        async with httpx.AsyncClient() as client_http:
            await client_http.post(f"{BAILEYS_URL}/api/send", json={"to": from_, "message": reply})
        
    except Exception as e:
        print(f"Erreur analyse: {e}")
    
    return {"status": "ok"}

@app.post("/api/whatsapp-baileys/send")
async def baileys_send(request: Dict[str, Any]):
    """Envoie un message via Baileys"""
    to = request.get("to")
    message = request.get("message")
    
    async with httpx.AsyncClient() as client:
        response = await client.post(f"{BAILEYS_URL}/api/send", json={"to": to, "message": message})
        return response.json()

@app.get("/api/whatsapp-baileys/status")
async def baileys_status():
    """Vérifie si Baileys est connecté"""
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(f"{BAILEYS_URL}/api/status")
            return response.json()
    except:
        return {"connected": False, "error": "Service Baileys non disponible"}
        
# =====================================================
# FONCTIONS UTILITAIRES
# =====================================================

def normalize_messages(messages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Convertit les messages au format OpenAI (supporte texte + images)"""
    normalized = []
    for msg in messages:
        role = msg.get("role", "user")
        content = msg.get("content")
        if isinstance(content, list):
            normalized.append({"role": role, "content": content})
        else:
            normalized.append({"role": role, "content": content})
    return normalized

# =====================================================
# LOGGING CONFIGURATION
# =====================================================

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)





# =====================================================
# ENVIRONMENT VARIABLES & CLIENTS
# =====================================================

OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY")
VAPID_PRIVATE_KEY = os.environ.get("VAPID_PRIVATE_KEY")
VAPID_PUBLIC_KEY = os.environ.get("VAPID_PUBLIC_KEY")
VAPID_CLAIMS = {"sub": "mailto:jbillcataria@gmail.com"}
GREENAPI_ID_INSTANCE = os.environ.get("GREENAPI_ID_INSTANCE")
GREENAPI_API_TOKEN = os.environ.get("GREENAPI_API_TOKEN")
GREENAPI_BASE_URL = f"https://api.green-api.com/waInstance{GREENAPI_ID_INSTANCE}" if GREENAPI_ID_INSTANCE else None


# Debug GreenAPI
print(f"🔍 GREENAPI_ID_INSTANCE: {GREENAPI_ID_INSTANCE}")
print(f"🔍 GREENAPI_API_TOKEN: {GREENAPI_API_TOKEN[:10] if GREENAPI_API_TOKEN else 'None'}...")
print(f"🔍 GREENAPI_BASE_URL: {GREENAPI_BASE_URL}")

# =====================================================
# USER CONTEXT
# =====================================================

DEFAULT_USER_ID = os.environ.get("DEFAULT_USER_ID")

def get_request_user_id(request_data: Dict[str, Any] = None) -> str:
    """
    Récupère le user_id envoyé par le frontend.
    Si absent, utilise DEFAULT_USER_ID seulement s'il existe dans les variables d'environnement.
    Sinon, renvoie une erreur claire.
    """
    incoming_user_id = None

    if request_data:
        incoming_user_id = request_data.get("user_id")

    return require_user_id(incoming_user_id)

def normalize_user_id(user_id: Optional[str] = None) -> Optional[str]:
    """
    Retourne un user_id propre.
    Ne met plus 'rebecca' par défaut, car certaines tables attendent un UUID.
    """
    if user_id and user_id != "rebecca":
        return user_id
    
    return DEFAULT_USER_ID

def require_user_id(user_id: Optional[str] = None) -> str:
    """
    Exige un vrai user_id.
    À utiliser sur les routes qui lisent/écrivent des données personnelles.
    """
    clean_user_id = normalize_user_id(user_id)

    if not clean_user_id:
        raise HTTPException(
            status_code=400,
            detail="user_id requis. Le frontend doit envoyer l'identifiant utilisateur connecté."
        )

    return clean_user_id


if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY manquante")

client = OpenAI(api_key=OPENAI_API_KEY)

supabase: Client = None
if SUPABASE_URL and SUPABASE_SERVICE_KEY:
    supabase = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    logger.info("✅ Supabase connecté")
else:
    logger.warning("⚠️ Supabase non configuré")




# =====================================================
# DATABASE SCHEMA CONFIGURATION
# =====================================================

AVAILABLE_TABLES = [
    "missions", "tasks", "spending", "revenue", "documents",
    "content", "family_events", "wins", "relocation_tasks",
    "farm_infrastructure", "farm_production_units", "farm_spending", "farm_team",
    "user_memory", "mood_entries", "user_profile",
    "brain_dump_analyses", "checklists", "drafts",
    "lf_grants", "lf_contracts", "lf_contacts", "lf_tasks"
]

ALLOWED_FIELDS = {
    "spending": ["title", "amount", "category", "date", "notes", "verified", "mission_id", "project", "beneficiary", "user_id"],
    "tasks": ["title", "status", "due_date", "estimated_time", "mission_id", "project", "notes", "sync_calendar", "calendar_event_id", "calendar_synced", "calendar_link", "user_id"],
    "wins": ["title", "category", "date", "notes", "celebration_emoji", "user_id"],
    "family_events": ["title", "child_name", "category", "date", "notes", "user_id"],
    "missions": ["name", "category", "status", "priority", "deadline", "owner", "revenue_potential", "strategic_value", "energy_cost", "user_id"],
    "revenue": ["source", "amount", "date", "notes", "mission_id", "project", "user_id"],
    "documents": ["name", "type", "status", "due_date", "url", "missing_pieces", "notes", "mission_id", "user_id"],
    "content": ["title", "hook", "platform", "content_type", "status", "publish_date", "cta", "mission_id", "user_id"],
    "relocation_tasks": ["title", "category", "status", "due_date", "notes", "user_id"],
    "farm_infrastructure": ["name", "type", "status", "location_on_site", "completed_date", "responsible_person", "notes", "user_id"],
    "farm_production_units": ["name", "category", "status", "current_capacity", "start_date", "expected_first_revenue", "technical_lead", "notes", "user_id"],
    "farm_spending": ["title", "amount", "category", "project_area", "verified", "notes", "user_id"],
    "farm_team": ["name", "role", "area", "status", "phone", "notes", "user_id"],
    "user_memory": ["category", "key", "value", "user_id"],
    "mood_entries": ["mood", "date", "user_id"],
    "user_profile": ["user_id", "full_name", "preferred_name", "birthday", "children", "projects", "communication_preferences", "current_goals", "upcoming_milestones", "key_contacts", "user_id"],
    "brain_dump_analyses": ["content", "analysis", "user_id"],
    "checklists": ["title", "steps", "completed_steps", "progress", "user_id"],
    "drafts": ["type", "content", "context", "user_id"],
    "lf_grants": ["title", "agency", "amount", "deadline", "status", "probability", "notes", "user_id"],
    "lf_contracts": ["title", "contract_type", "agency", "status", "deadline", "requirements", "notes", "user_id"],
    "lf_contacts": ["name", "organization", "role", "email", "phone", "type", "last_contact", "notes", "user_id"],
    "lf_tasks": ["title", "related_to", "related_id", "status", "deadline", "priority", "notes", "user_id"]
}



# =====================================================
# MAPPING INTELLIGENT DES CATÉGORIES (Générique)
# =====================================================

# Définition des contraintes par table
TABLE_CONSTRAINTS = {
    "spending": {
        "field": "category",
        "valid_values": ["materials", "construction", "labor", "livestock", "crops", "transport", "equipment", "food", "other"],
        "mapping": {
            # Français → Anglais valide
            "matériel": "equipment", "materiel": "equipment", "équipement": "equipment",
            "outil": "equipment", "outils": "equipment", "machine": "equipment",
            "fourniture": "materials", "fournitures": "materials", "matériau": "materials",
            "semence": "crops", "semences": "crops", "engrais": "crops",
            "animal": "livestock", "animaux": "livestock", "aliment": "livestock",
            "construction": "construction", "bâtiment": "construction", "mur": "construction",
            "salaire": "labor", "salaires": "labor", "main d'oeuvre": "labor",
            "transport": "transport", "livraison": "transport", "essence": "transport",
            # NOUVEAU - Alimentation
            "alimentation": "food",
            "alimentaire": "food",
            "nourriture": "food",
            "courses": "food",
            "repas": "food",
            "restaurant": "food",
            "manger": "food",
            "cuisine": "food",
            "épicerie": "food",
            "epicerie": "food",
            # Anglais déjà valide → garder tel quel
            "materials": "materials", "construction": "construction", "labor": "labor",
            "livestock": "livestock", "crops": "crops", "transport": "transport",
            "equipment": "equipment", "other": "other"
        },
        "default": "other"
    },
    
    "tasks": {
        "field": "status",
        "valid_values": ["not_started", "today", "in_progress", "waiting", "done"],
        "mapping": {
            "à faire": "not_started", "a faire": "not_started", "pas commencé": "not_started",
            "aujourd'hui": "today", "aujourdhui": "today",
            "en cours": "in_progress", "commencé": "in_progress",
            "en attente": "waiting", "attente": "waiting",
            "terminé": "done", "termine": "done", "fini": "done", "fait": "done"
        },
        "default": "not_started"
    },
    
    "tasks_priority": {
        "field": "priority",
        "valid_values": ["critical", "high", "normal", "low"],
        "mapping": {
            "critique": "critical", "urgent": "critical",
            "haute": "high", "important": "high",
            "normale": "normal", "moyenne": "normal",
            "basse": "low", "faible": "low"
        },
        "default": "normal"
    },
    
    "missions": {
        "field": "status",
        "valid_values": ["idea", "planning", "active", "waiting", "paused", "complete"],
        "mapping": {
            "idée": "idea", "idee": "idea",
            "planification": "planning", "planifié": "planning",
            "active": "active", "activé": "active", "en cours": "active",
            "en attente": "waiting",
            "pause": "paused", "en pause": "paused",
            "terminée": "complete", "terminee": "complete", "fait": "complete"
        },
        "default": "idea"
    },
    
    "missions_priority": {
        "field": "priority",
        "valid_values": ["critical", "high", "normal", "low"],
        "mapping": {
            "critique": "critical", "urgent": "critical",
            "haute": "high", "important": "high",
            "normale": "normal", "moyenne": "normal",
            "basse": "low", "faible": "low"
        },
        "default": "normal"
    },
    
    "family_events": {
        "field": "category",
        "valid_values": ["school", "health", "activity", "travel", "document", "routine", "supplies"],
        "mapping": {
            "école": "school", "ecole": "school", "cours": "school",
            "santé": "health", "sante": "health", "médecin": "health", "docteur": "health",
            "activité": "activity", "activite": "activity", "sport": "activity",
            "voyage": "travel", "déplacement": "travel",
            "papiers": "document", "administratif": "document",
            "routine": "routine", "quotidien": "routine",
            "fournitures": "supplies", "achats": "supplies"
        },
        "default": "routine"
    },
    
    "wins": {
        "field": "category",
        "valid_values": ["business", "family", "personal", "money", "health", "farm", "other"],
        "mapping": {
            "business": "business", "projet": "business", "travail": "business",
            "famille": "family", "enfant": "family",
            "personnel": "personal", "perso": "personal",
            "argent": "money", "finance": "money",
            "santé": "health", "sante": "health",
            "ferme": "farm", "agriculture": "farm"
        },
        "default": "other"
    },
    
    "documents": {
        "field": "type",
        "valid_values": ["proposal", "contract", "grant", "invoice", "legal", "admin", "other"],
        "mapping": {
            "proposition": "proposal", "offre": "proposal",
            "contrat": "contract",
            "subvention": "grant",
            "facture": "invoice",
            "légal": "legal", "juridique": "legal",
            "administratif": "admin"
        },
        "default": "other"
    },
    
    "documents_status": {
        "field": "status",
        "valid_values": ["draft", "review", "ready", "submitted", "approved", "rejected"],
        "mapping": {
            "brouillon": "draft",
            "relecture": "review", "vérification": "review",
            "prêt": "ready", "pret": "ready",
            "soumis": "submitted", "envoyé": "submitted",
            "approuvé": "approved", "approuve": "approved", "validé": "approved",
            "rejeté": "rejected", "refusé": "rejected"
        },
        "default": "draft"
    }
}

def normalize_field_value(table: str, field: str, value: str) -> str:
    """
    Normalise une valeur de champ selon les contraintes de la table.
    Retourne la valeur normalisée ou la valeur par défaut.
    """
    if not value or not isinstance(value, str):
        return value
    
    value_lower = value.lower().strip()
    
    # Chercher la configuration pour ce champ
    config_key = f"{table}_{field}"
    if config_key in TABLE_CONSTRAINTS:
        config = TABLE_CONSTRAINTS[config_key]
    elif table in TABLE_CONSTRAINTS and TABLE_CONSTRAINTS[table].get("field") == field:
        config = TABLE_CONSTRAINTS[table]
    else:
        # Pas de contrainte connue pour ce champ
        return value
    
    valid_values = config.get("valid_values", [])
    mapping = config.get("mapping", {})
    default_value = config.get("default")
    
    # Vérifier si la valeur est déjà valide
    if value_lower in valid_values:
        return value_lower
    
    # Essayer de mapper
    if value_lower in mapping:
        mapped = mapping[value_lower]
        logger.info(f"🔄 Mappage {table}.{field}: '{value}' -> '{mapped}'")
        return mapped
    
    # Chercher par correspondance partielle
    for key, mapped_value in mapping.items():
        if key in value_lower or value_lower in key:
            logger.info(f"🔄 Mappage partiel {table}.{field}: '{value}' -> '{mapped_value}'")
            return mapped_value
    
    # Valeur par défaut si disponible
    if default_value:
        logger.warning(f"⚠️ Valeur inconnue {table}.{field}: '{value}' -> défaut '{default_value}'")
        return default_value
    
    return value
# =====================================================
# PYDANTIC MODELS
# =====================================================

class ChatRequest(BaseModel):
    messages: List[Dict[str, Any]]
    user_id: Optional[str] = None


class WriteRequest(BaseModel):
    table: str
    data: Dict


class UpdateRequest(BaseModel):
    table: str
    id: str
    data: Dict


class MemorySaveRequest(BaseModel):
    category: str
    key: str
    value: str
    user_id: Optional[str] = None


class ExecuteTaskRequest(BaseModel):
    title: str
    due_date: Optional[str] = None
    priority: str = "normal"
    user_id: Optional[str] = None


# =====================================================
# FONCTIONS UTILITAIRES POUR NOTIFICATIONS
# =====================================================

def send_notification_sync(notification_data: Dict[str, Any]) -> List[Dict]:
    """Envoie une notification à tous les abonnés avec logging et stockage pour l'interface"""
    if not supabase:
        logger.error("Supabase non configuré")
        return []
    
    user_id = require_user_id(notification_data.get("user_id"))
    notif_type = notification_data.get("type", "default")
    today = datetime.now().date().isoformat()
    
    # Types de notifications qui peuvent être envoyées plusieurs fois (urgences)
    multi_allow_types = ["task_reminder", "document_reminder", "family_reminder"]
    
    # Vérifier si cette notification a déjà été envoyée aujourd'hui (sauf pour les types autorisés)
    if notif_type not in multi_allow_types:
        try:
            existing = supabase.table("notifications_log").select("*")\
                .eq("type", notif_type)\
                .eq("date", today)\
                .eq("user_id", user_id)\
                .execute()
            
            if existing.data and len(existing.data) > 0:
                logger.info(f"⏭️ Notification {notif_type} déjà envoyée aujourd'hui, ignorée")
                return []
        except Exception as e:
            logger.error(f"Erreur vérification log: {e}")
    
    # ============================================
    # 1. STOCKER LA NOTIFICATION DANS LA BASE (pour l'interface)
    # ============================================
    try:
        notification_record = {
            "title": notification_data.get("title", "SOVEREIGN"),
            "body": notification_data.get("body", ""),
            "type": notif_type,
            "url": notification_data.get("url", "/"),
            "user_id": user_id,
            "created_at": datetime.now().isoformat(),
            "read": False
        }
        
        # Insérer dans la table notifications
        insert_result = supabase.table("notifications").insert(notification_record).execute()
        
        if insert_result.data:
            logger.info(f"💾 Notification stockée en base (id: {insert_result.data[0].get('id')})")
    except Exception as e:
        logger.error(f"Erreur stockage notification dans 'notifications': {e}")
    
    # ============================================
    # 2. ENVOYER LES NOTIFICATIONS PUSH
    # ============================================
    subscriptions = (     supabase.table("push_subscriptions")     .select("*")     .eq("user_id", user_id)     .execute() )
    results = []
    
    # Styles selon le type
    type_styles = {
        "task": {"emoji": "📋", "color": "#3B82F6"},
        "mission": {"emoji": "🎯", "color": "#8B5CF6"},
        "win": {"emoji": "🏆", "color": "#F59E0B"},
        "money": {"emoji": "💰", "color": "#10B981"},
        "family": {"emoji": "👨‍👩‍👧‍👦", "color": "#EC4899"},
        "document": {"emoji": "📄", "color": "#EF4444"},
        "morning": {"emoji": "🌅", "color": "#D4AF37"},
        "default": {"emoji": "👑", "color": "#D4AF37"}
    }
    
    style = type_styles.get(notif_type, type_styles["default"])
    title_with_emoji = f"{style['emoji']} {notification_data.get('title', 'SOVEREIGN')}"
    
    for sub in subscriptions.data:
        try:
            webpush(
                subscription_info={
                    "endpoint": sub["endpoint"],
                    "keys": sub["keys"]
                },
                data=json.dumps({
                    "title": title_with_emoji,
                    "body": notification_data.get("body", ""),
                    "url": notification_data.get("url", "/"),
                    "icon": "/icons/icon-192x192.png",
                    "badge": "/icons/icon-96x96.png",
                    "image": "/icons/icon-512x512.png",
                    "type": notif_type,
                    "sound": notification_data.get("sound", "/sounds/notification.mp3"),
                    "vibrate": notification_data.get("vibrate", [200, 100, 200]),
                    "requireInteraction": notification_data.get("requireInteraction", False),
                    "silent": notification_data.get("silent", False),
                    "timestamp": datetime.now().isoformat()
                }),
                vapid_private_key=VAPID_PRIVATE_KEY,
                vapid_claims=VAPID_CLAIMS
            )
            results.append({"status": "sent", "endpoint": sub["endpoint"][:50]})
            logger.info(f"✅ Notification {style['emoji']} envoyée à {sub['endpoint'][:50]}...")
        except WebPushException as e:
            if e.response and e.response.status_code == 410:
                # Subscription expirée, on la supprime
                supabase.table("push_subscriptions").delete().eq("endpoint", sub["endpoint"]).execute()
                results.append({"status": "expired", "endpoint": sub["endpoint"][:50]})
                logger.info(f"🗑️ Subscription expirée supprimée: {sub['endpoint'][:50]}...")
            else:
                results.append({"status": "error", "error": str(e)})
                logger.error(f"❌ Erreur webpush: {e}")
    
    # ============================================
    # 3. LOGGER DANS notifications_log (pour éviter les doublons)
    # ============================================
    if results and len(results) > 0:
        try:
            # Vérifier si déjà loggé (pour éviter les doublons de log)
            existing_log = supabase.table("notifications_log").select("*")\
                .eq("type", notif_type)\
                .eq("date", today)\
                .eq("user_id", user_id)\
                .execute()
            
            if not existing_log.data:
                supabase.table("notifications_log").insert({
                    "type": notif_type,
                    "date": today,
                    "user_id": user_id,
                    "sent_at": datetime.now().isoformat(),
                    "metadata": {
                        "title": notification_data.get("title"),
                        "body": notification_data.get("body"),
                        "count": len(results)
                    }
                }).execute()
                logger.info(f"📝 Notification {notif_type} loggée dans notifications_log")
        except Exception as e:
            logger.error(f"Erreur log notification: {e}")
    
    return results


def get_days_late(date_str: str) -> int:
    """Calcule le nombre de jours de retard"""
    if not date_str:
        return 0
    due_date = datetime.fromisoformat(date_str).date()
    today = datetime.now().date()
    delta = today - due_date
    return max(0, delta.days)



def store_chat_session(
    user_message: str,
    assistant_response: str,
    tools_used: List[str] = None,
    user_id: Optional[str] = None
):
    """Stocke la session de chat dans Supabase"""
    if not supabase:
        return
    
    try:
        user_id = require_user_id(user_id)
        supabase.table("chat_sessions").insert({
            "user_message": user_message[:500],
            "assistant_response": assistant_response[:1000],
            "tools_used": tools_used or [],
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }).execute()
        logger.info("💾 Conversation stockée")
    except Exception as e:
        logger.error(f"Erreur store_chat: {e}")
# =====================================================
# FONCTIONS POUR LA MÉMOIRE UTILISATEUR
# =====================================================

async def get_user_memory_context(user_id: Optional[str] = None) -> str:
    """Récupère la mémoire utilisateur"""
    if not supabase:
        return ""
    
    try:
        user_id = require_user_id(user_id)
        # Utiliser une requête différente si l'ID n'est pas un UUID
        result = supabase.table("user_memory").select("*").eq("user_id", user_id).execute()
        memories = result.data
        
        if not memories:
            return ""
        
        # Organiser par catégorie
        categorized = {}
        for mem in memories:
            cat = mem.get("category", "general")
            if cat not in categorized:
                categorized[cat] = []
            categorized[cat].append(f"{mem['key']}: {mem['value']}")
        
        # Construire le texte
        memory_text = "\n\n📚 INFORMATIONS SUR L'UTILISATEUR:\n"
        for cat, items in categorized.items():
            memory_text += f"\n{cat.upper()}:\n"
            for item in items:
                memory_text += f"  - {item}\n"
        
        return memory_text
    except Exception as e:
        logger.error(f"Erreur récupération mémoire: {e}")
        return ""


async def save_user_memory(category: str, key: str, value: str, user_id: Optional[str] = None):
    """Sauvegarde une information dans la mémoire utilisateur"""
    if not supabase:
        return False
    
    try:
        user_id = require_user_id(user_id)

        existing = (
            supabase.table("user_memory")
            .select("*")
            .eq("user_id", user_id)
            .eq("category", category)
            .eq("key", key)
            .execute()
        )
        
        if existing.data:
            supabase.table("user_memory").update({
                "value": value,
                "updated_at": datetime.now().isoformat()
            }).eq("id", existing.data[0]["id"]).execute()
        else:
            supabase.table("user_memory").insert({
                "category": category,
                "key": key,
                "value": value,
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }).execute()
        
        logger.info(f"💾 Mémoire sauvegardée: {category}/{key} = {value}")
        return True
    except Exception as e:
        logger.error(f"Erreur sauvegarde mémoire: {e}")
        return False





# =====================================================
# GESTION DES CONTACTS
# =====================================================

async def get_contact_number(contact_name: str, user_id: Optional[str] = None) -> dict:
    if not supabase:
        return {"found": False, "phone": None, "source": None}

    user_id = require_user_id(user_id)
    
    contact_name_lower = contact_name.lower().strip()
    
    # 1. Chercher dans user_memory (contacts rapides)
    result = supabase.table("user_memory").select("*").eq("user_id", user_id).eq("category", "contact").execute()
    for mem in result.data:
        key = mem.get("key", "").lower()
        if contact_name_lower in key or key in contact_name_lower:
            return {"found": True, "phone": mem.get("value"), "source": "memory"}
    
    # 2. Chercher dans lf_contacts
    contacts = supabase.table("lf_contacts").select("*").eq("user_id", user_id).execute()
    for contact in contacts.data:
        name = contact.get("name", "").lower()
        if contact_name_lower in name or name in contact_name_lower:
            if contact.get("phone"):
                return {"found": True, "phone": contact.get("phone"), "source": "contacts_table"}
    
    return {"found": False, "phone": None, "source": None}


def extract_phone_from_text(text: str) -> str:
    """Extrait un numéro de téléphone du texte"""
    patterns = [
        r'(\+229\s*\d{2}\s*\d{2}\s*\d{2}\s*\d{2})',  # +229 XX XX XX XX
        r'(\+229\s*\d{8})',                           # +229XXXXXXXX
        r'(0[67]\s*\d{2}\s*\d{2}\s*\d{2}\s*\d{2})',  # 06 XX XX XX XX ou 07
        r'(\d{2}\s*\d{2}\s*\d{2}\s*\d{2}\s*\d{2})',   # XX XX XX XX XX (10 chiffres)
        r'(\d{8})',                                    # 8 chiffres
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            phone = match.group(1).replace(" ", "")
            # Standardiser le format
            if phone.startswith("0") and len(phone) == 10:
                phone = "+229" + phone[1:]
            elif len(phone) == 8:
                phone = "+229" + phone
            return phone
    return None


async def save_contact_memory(contact_name: str, phone: str, user_id: Optional[str] = None):
    if not supabase:
        return False

    user_id = require_user_id(user_id)
    
    key = f"{contact_name.lower()}_phone"
    
    # Vérifier si existe déjà
    existing = supabase.table("user_memory").select("*").eq("user_id", user_id).eq("category", "contact").eq("key", key).execute()
    
    if existing.data:
        supabase.table("user_memory").update({
            "value": phone,
            "updated_at": datetime.now().isoformat()
        }).eq("id", existing.data[0]["id"]).execute()
    else:
        supabase.table("user_memory").insert({
            "category": "contact",
            "key": key,
            "value": phone,
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }).execute()
    
    logger.info(f"💾 Contact sauvegardé: {contact_name} -> {phone}")
    return True
    
# =====================================================
# FONCTIONS POUR LA VISION ET LES DOCUMENTS
# =====================================================

async def download_image_from_url(url: str) -> str:
    """Télécharge une image depuis une URL et la convertit en base64"""
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.get(url)
            response.raise_for_status()
            import base64
            content_type = response.headers.get('content-type', 'image/png')
            base64_image = base64.b64encode(response.content).decode('utf-8')
            return f"data:{content_type};base64,{base64_image}"
    except Exception as e:
        logger.error(f"Erreur téléchargement image: {e}")
        return None


def extract_text_from_message(content: str) -> tuple[str, List[str]]:
    """Extrait le texte et les URLs d'images d'un message"""
    text_parts = []
    image_urls = []
    all_file_urls = []
    
    image_pattern = r'(https?://[^\s]+\.(?:png|jpg|jpeg|gif|webp))'
    image_urls = re.findall(image_pattern, content, re.IGNORECASE)
    
    file_pattern = r'https?://[^\s]+\.(?:pdf|txt|docx?|jpg|png|jpeg|gif|webp)'
    all_file_urls = re.findall(file_pattern, content, re.IGNORECASE)
    
    for url in image_urls:
        if url in all_file_urls:
            all_file_urls.remove(url)
        content = content.replace(url, f"[IMAGE: {url}]")
    
    text_parts.append(content)
    
    return " ".join(text_parts), image_urls, all_file_urls


async def process_document(file_url: str) -> str:
    """Télécharge et extrait le texte d'un document"""
    if not supabase:
        return None
    
    try:
        
        match = re.search(r'/chat-files/(.+)$', file_url)
        if not match:
            logger.error(f"Impossible d'extraire le chemin du fichier: {file_url}")
            return None
        
        file_path = match.group(1)
        logger.info(f"📄 Tentative de téléchargement: {file_path}")
        
        try:
            file_data = supabase.storage.from_("chat-files").download(file_path)
        except Exception as e:
            logger.error(f"Erreur téléchargement Supabase: {e}")
            if file_path.startswith("chat/"):
                alt_path = file_path[5:]
                logger.info(f"🔄 Essai chemin alternatif: {alt_path}")
                file_data = supabase.storage.from_("chat-files").download(alt_path)
            else:
                return None
        
        if file_path.lower().endswith('.pdf'):
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(file_data))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            logger.info(f"📄 PDF extrait: {len(text)} caractères")
            return text[:5000]
        
        elif file_path.lower().endswith('.txt'):
            text = file_data.decode('utf-8')
            logger.info(f"📄 TXT extrait: {len(text)} caractères")
            return text[:5000]
        
        elif file_path.lower().endswith(('.docx', '.doc')):
            from docx import Document
            import io
            doc = Document(io.BytesIO(file_data))
            text = "\n".join([para.text for para in doc.paragraphs])
            logger.info(f"📄 DOCX extrait: {len(text)} caractères")
            return text[:5000]
        
        else:
            logger.warning(f"Format non supporté: {file_path}")
            return None
            
    except Exception as e:
        logger.error(f"Erreur extraction document: {e}")
        return None




def db_update(table: str, id: str, data: Dict) -> Dict:
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        allowed = ALLOWED_FIELDS.get(table, [])
        clean_data = {k: v for k, v in data.items() if k in allowed}
        result = supabase.table(table).update(clean_data).eq("id", id).execute()
        
        if hasattr(result, 'data'):
            result_data = result.data
        else:
            result_data = result
        
        # Webhook pour mission complétée
        if result_data and len(result_data) > 0 and table == "missions":
            if clean_data.get("status") == "complete":
                asyncio.create_task(trigger_webhook("mission.completed", {
                    "mission": result_data[0] if isinstance(result_data, list) else result_data,
                    "timestamp": datetime.now().isoformat()
                }))
        
        return {"success": True, "data": result_data[0] if result_data and len(result_data) > 0 else None}
    except Exception as e:
        logger.error(f"Erreur update {table}: {e}")
        return {"success": False, "error": str(e)}



def db_delete(table: str, id: str) -> Dict:
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        supabase.table(table).delete().eq("id", id).execute()
        return {"success": True}
    except Exception as e:
        logger.error(f"Erreur delete {table}: {e}")
        return {"success": False, "error": str(e)}


def get_financial_summary(user_id: Optional[str] = None) -> Dict:
    if not supabase:
        return {"total_revenue": 0, "total_spending": 0, "net_balance": 0}
    
    try:
        user_id = require_user_id(user_id)

        rev_result = (
            supabase.table("revenue")
            .select("amount")
            .eq("user_id", user_id)
            .execute()
        )

        total_revenue = sum(r.get("amount", 0) for r in rev_result.data)
        
        spend_result = (
            supabase.table("spending")
            .select("amount")
            .eq("user_id", user_id)
            .execute()
        )

        total_spending = sum(s.get("amount", 0) for s in spend_result.data)
        
        return {
            "total_revenue": total_revenue,
            "total_spending": total_spending,
            "net_balance": total_revenue - total_spending,
            "currency": "XOF"
        }
    except Exception as e:
        logger.error(f"Erreur financial_summary: {e}")
        return {"total_revenue": 0, "total_spending": 0, "net_balance": 0}

def get_priority_tasks(limit: int = 10, user_id: Optional[str] = None) -> List[Dict]:
    if not supabase:
        return []
    
    try:
        user_id = require_user_id(user_id)

        result = (
            supabase.table("tasks")
            .select("*")
            .eq("user_id", user_id)
            .eq("status", "in_progress")
            .limit(limit)
            .execute()
        )

        return result.data if result.data else []
    except Exception as e:
        logger.error(f"Erreur priority_tasks: {e}")
        return []
# =====================================================
# NOUVEAUX ENDPOINTS SPÉCIAUX
# =====================================================

@app.post("/api/memory/save")
async def save_memory(request: MemorySaveRequest):
    """Sauvegarde une information dans la mémoire utilisateur"""
    try:
        user_id = require_user_id(request.user_id)
        result = await save_user_memory(
            request.category,
            request.key,
            request.value,
            user_id=user_id
        )
        return {"success": result, "message": "Mémoire sauvegardée" if result else "Erreur"}
    except Exception as e:
        logger.error(f"Erreur save_memory: {e}")
        return {"success": False, "error": str(e)}

@app.get("/api/memory/get")
async def get_memory(category: str = None, key: str = None, user_id: Optional[str] = None):
    if not supabase:
        return {"success": False, "data": []}
    
    try:
        user_id = require_user_id(user_id)

        query = supabase.table("user_memory").select("*").eq("user_id", user_id)
        if category:
            query = query.eq("category", category)
        if key:
            query = query.eq("key", key)
        
        result = query.execute()
        return {"success": True, "data": result.data}
    except Exception as e:
        logger.error(f"Erreur get_memory: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/execute/create-task")
async def create_task_from_conversation(request: ExecuteTaskRequest):
    """Crée une tâche à partir d'une conversation (mode exécution)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = require_user_id(request.user_id)
        
        result = supabase.table("tasks").insert({
            "title": request.title,
            "status": "today",
            "priority": request.priority,
            "due_date": request.due_date,
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }).execute()
        
        send_notification_sync({
            "title": "📋 Nouvelle tâche créée",
            "body": f"'{request.title}' a été ajoutée à vos tâches",
            "url": "/tasks",
            "type": "task",
            "user_id": user_id
        })
        
        if result.data and len(result.data) > 0:
            asyncio.create_task(trigger_webhook("task.created", {
                "task": result.data[0],
                "timestamp": datetime.now().isoformat(),
                "user_id": user_id
            }))
            logger.info(f"🔗 Webhook déclenché pour task.created: {request.title}")
        
        return {"success": True, "task": result.data[0] if result.data else None}
    except Exception as e:
        logger.error(f"Erreur create_task_from_conversation: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/mood/save")
async def save_mood(request: Dict[str, Any]):
    """Sauvegarde l'humeur du jour"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        today = datetime.now().date().isoformat()
        mood = request.get("mood")
        user_id = get_request_user_id(request)
        
        existing = (
            supabase
            .table("mood_entries")
            .select("*")
            .eq("date", today)
            .eq("user_id", user_id)
            .execute()
        )
        
        if existing.data:
            supabase.table("mood_entries").update({
                "mood": mood
            }).eq("id", existing.data[0]["id"]).execute()
        else:
            supabase.table("mood_entries").insert({
                "mood": mood,
                "date": today,
                "user_id": user_id
            }).execute()
        
        encouragement = ""
        if mood == "fatiguée":
            encouragement = "Prends soin de toi aujourd'hui. Une petite chose à la fois."
        elif mood == "stressée":
            encouragement = "🌿 On va respirer. Une seule priorité pour commencer."
        elif mood == "excellent":
            encouragement = "🔥 C'est le moment d'attaquer les gros dossiers !"
        
        return {"success": True, "encouragement": encouragement}
    except Exception as e:
        logger.error(f"Erreur save_mood: {e}")
        return {"success": False, "error": str(e)}


@app.get("/api/mood/history")
async def get_mood_history(days: int = 30, user_id: Optional[str] = None):
    if not supabase:
        return {"success": False, "data": []}
    
    try:
        user_id = require_user_id(user_id)
        start_date = (datetime.now().date() - timedelta(days=days)).isoformat()

        result = (
            supabase
            .table("mood_entries")
            .select("*")
            .eq("user_id", user_id)
            .gte("date", start_date)
            .order("date", desc=True)
            .execute()
        )

        return {"success": True, "data": result.data}
    except Exception as e:
        logger.error(f"Erreur get_mood_history: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# NOUVEAUX ENDPOINTS DE NOTIFICATIONS
# =====================================================

@app.post("/api/check-task-reminders")
async def check_task_reminders(request: Dict[str, Any] = None):
    """Vérifie les tâches et envoie des rappels (1x par jour max)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    today = datetime.now().date().isoformat()
    user_id = get_request_user_id(request or {})
    
    # Vérifier si déjà envoyé aujourd'hui
    existing = supabase.table("notifications_log").select("*")\
        .eq("type", "task_reminder")\
        .eq("date", today)\
        .eq("user_id", user_id)\
        .execute()
    
    if existing.data:
        return {"success": True, "sent": False, "message": "Déjà envoyé aujourd'hui"}
    
    notifications_sent = []
    
    try:
        tasks_today = (
        supabase.table("tasks")
        .select("*")
        .eq("user_id", user_id)
        .eq("due_date", today)
        .neq("status", "done")
        .execute()
        )
        
        overdue_tasks = (
            supabase.table("tasks")
            .select("*")
            .eq("user_id", user_id)
            .lt("due_date", today)
            .neq("status", "done")
            .execute()
        )
        
        # Un seul message regroupé, pas un par tâche
        if overdue_tasks.data:
            body = f"⚠️ {len(overdue_tasks.data)} tâche(s) en retard. On regarde ça ?"
            send_notification_sync({
                "title": "📋 Tâches en retard",
                "body": body,
                "url": "/tasks",
                "type": "task",
                "user_id": user_id
            })
            notifications_sent.append("overdue_tasks")
        elif tasks_today.data:
            body = f"📋 {len(tasks_today.data)} tâche(s) à faire aujourd'hui."
            send_notification_sync({
                "title": "📋 Tâches du jour",
                "body": body,
                "url": "/tasks",
                "type": "task",
                "user_id": user_id
            })
            notifications_sent.append("tasks_today")
        
        # Logger l'envoi
        if notifications_sent:
            supabase.table("notifications_log").insert({
                "type": "task_reminder",
                "date": today,
                "user_id": user_id,
                "sent_at": datetime.now().isoformat()
            }).execute()
        
        return {"success": True, "notifications_sent": notifications_sent, "count": len(notifications_sent)}
    
    except Exception as e:
        logger.error(f"Erreur check_task_reminders: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/mission-reminders")
async def mission_reminders(request: Dict[str, Any] = None):
    """Rappel pour les missions inactives (1x par semaine max)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}

    user_id = get_request_user_id(request or {})
    
    # Vérifier si déjà envoyé cette semaine
    week_ago = (datetime.now() - timedelta(days=7)).isoformat()
    existing = supabase.table("notifications_log").select("*")\
        .eq("type", "mission_reminder")\
        .gte("sent_at", week_ago)\
        .eq("user_id", user_id)\
        .execute()
    
    if existing.data:
        return {"success": True, "sent": False, "message": "Déjà envoyé cette semaine"}
    
    notifications_sent = []
    
    try:
        five_days_ago = (datetime.now() - timedelta(days=5)).isoformat()
        stale_missions = (
            supabase.table("missions")
            .select("*")
            .eq("user_id", user_id)
            .eq("status", "active")
            .lt("updated_at", five_days_ago)
            .execute()
        )
        if stale_missions.data:
            mission_names = ", ".join([m["name"] for m in stale_missions.data[:3]])
            if len(stale_missions.data) > 3:
                mission_names += f" et {len(stale_missions.data)-3} autre(s)"
            
            send_notification_sync({
                "title": "🎯 Missions inactives",
                "body": f"{len(stale_missions.data)} mission(s) sans activité : {mission_names}",
                "url": "/missions",
                "type": "mission",
                "user_id": user_id
            })
            notifications_sent = [m["id"] for m in stale_missions.data]
            
            # Logger l'envoi
            supabase.table("notifications_log").insert({
                "type": "mission_reminder",
                "date": datetime.now().date().isoformat(),
                "user_id": user_id,
                "sent_at": datetime.now().isoformat()
            }).execute()
        
        return {"success": True, "notifications_sent": notifications_sent, "count": len(notifications_sent)}
    
    except Exception as e:
        logger.error(f"Erreur mission_reminders: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/document-reminders")
async def document_reminders(request: Dict[str, Any] = None):
    """Rappel pour les documents proches de l'échéance (1x par jour max)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}

    user_id = get_request_user_id(request or {})
    
    today = datetime.now().date().isoformat()
    
    # Vérifier si déjà envoyé aujourd'hui
    existing = supabase.table("notifications_log").select("*")\
        .eq("type", "document_reminder")\
        .eq("date", today)\
        .eq("user_id", user_id)\
        .execute()
    
    if existing.data:
        return {"success": True, "sent": False, "message": "Déjà envoyé aujourd'hui"}
    
    notifications_sent = []
    
    try:
        next_week = (datetime.now().date() + timedelta(days=7)).isoformat()
        
        expiring_docs = (
            supabase.table("documents")
            .select("*")
            .eq("user_id", user_id)
            .gte("due_date", today)
            .lte("due_date", next_week)
            .neq("status", "approved")
            .execute()
        )
        overdue_docs = (
            supabase.table("documents")
            .select("*")
            .eq("user_id", user_id)
            .lt("due_date", today)
            .neq("status", "approved")
            .execute()
        )        
        # Priorité aux documents en retard
        if overdue_docs.data:
            doc_names = ", ".join([d["name"] for d in overdue_docs.data[:3]])
            send_notification_sync({
                "title": "⚠️ Documents en retard",
                "body": f"{len(overdue_docs.data)} document(s) en retard : {doc_names}",
                "url": "/documents",
                "type": "document",
                "user_id": user_id
            })
            notifications_sent = [d["id"] for d in overdue_docs.data]
        elif expiring_docs.data:
            days_left = (datetime.fromisoformat(expiring_docs.data[0]["due_date"]).date() - datetime.now().date()).days
            send_notification_sync({
                "title": "📄 Documents bientôt dus",
                "body": f"{len(expiring_docs.data)} document(s) à rendre dans {days_left} jour(s)",
                "url": "/documents",
                "type": "document",
                "user_id": user_id
            })
            notifications_sent = [d["id"] for d in expiring_docs.data]
        
        if notifications_sent:
            supabase.table("notifications_log").insert({
                "type": "document_reminder",
                "date": today,
                "user_id": user_id,
                "sent_at": datetime.now().isoformat()
            }).execute()
        
        return {"success": True, "notifications_sent": notifications_sent, "count": len(notifications_sent)}
    
    except Exception as e:
        logger.error(f"Erreur document_reminders: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/celebration-reminder")
async def celebration_reminder(request: Dict[str, Any] = None):
    """Rappel pour encourager l'enregistrement des victoires (1x par semaine max)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}

    user_id = get_request_user_id(request or {})
    
    # Vérifier si déjà envoyé cette semaine
    week_ago = (datetime.now() - timedelta(days=7)).isoformat()
    existing = supabase.table("notifications_log").select("*")\
        .eq("type", "celebration_reminder")\
        .gte("sent_at", week_ago)\
        .eq("user_id", user_id)\
        .execute()
    
    if existing.data:
        return {"success": True, "sent": False, "message": "Déjà envoyé cette semaine"}
    
    try:
        three_days_ago = (datetime.now() - timedelta(days=3)).isoformat()
        recent_wins = (     supabase.table("wins")     .select("*")     .eq("user_id", user_id)     .gte("date", three_days_ago)     .execute() )
        
        if len(recent_wins.data) == 0:
            send_notification_sync({
                "title": "🏆 Une petite victoire aujourd'hui ?",
                "body": "Chaque pas compte. Même une petite chose mérite d'être célébrée ✨",
                "url": "/wins",
                "type": "win",
                "user_id": user_id
            })
            
            supabase.table("notifications_log").insert({
                "type": "celebration_reminder",
                "date": datetime.now().date().isoformat(),
                "user_id": user_id,
                "sent_at": datetime.now().isoformat()
            }).execute()
            
            return {"success": True, "sent": True, "message": "Rappel victoire envoyé"}
        
        return {"success": True, "sent": False, "message": f"{len(recent_wins.data)} victoire(s) récente(s)"}
    
    except Exception as e:
        logger.error(f"Erreur celebration_reminder: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/morning-brief-reminder")
async def morning_brief_reminder(request: Dict[str, Any] = None):
    """Rappel du brief matinal (entre 7h et 9h)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}

    user_id = get_request_user_id(request or {})
    
    current_hour = datetime.now().hour
    
    try:
        if not (7 <= current_hour <= 9):
            return {"success": True, "sent": False, "message": "Pas l'heure du brief matinal"}
        
        today = datetime.now().date().isoformat()
        existing_brief = (     supabase.table("daily_briefs")     .select("*")     .eq("user_id", user_id)     .eq("date", today)     .execute() )
        
        if existing_brief.data:
            send_notification_sync({
                "title": "🌅 Bonjour Rebecca",
                "body": "Ton brief quotidien est prêt. Commence ta journée avec Sovereign.",
                "url": "/brief",
                "tag": "morning_brief",
                "type": "brief",
                "requireInteraction": True, 
                "user_id": user_id
            })
            return {"success": True, "sent": True, "message": "Brief matinal envoyé"}
        else:
            return {"success": True, "sent": False, "message": "Aucun brief pour aujourd'hui"}
    
    except Exception as e:
        logger.error(f"Erreur morning_brief_reminder: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/weekly-report-reminder")
async def weekly_report_reminder(request: Dict[str, Any] = None):
    """Rappel pour le rapport hebdomadaire (le dimanche)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    user_id = get_request_user_id(request or {})

    if datetime.now().weekday() != 6:
        return {"success": True, "sent": False, "message": "Pas le jour du rapport hebdomadaire"}
    
    try:
        start_of_week = datetime.now().date() - timedelta(days=7)
        start_of_week_str = start_of_week.isoformat()
        
        tasks_completed = (
            supabase.table("tasks")
            .select("*")
            .eq("user_id", user_id)
            .eq("status", "done")
            .gte("updated_at", start_of_week_str)
            .execute()
        )
        
        wins = (
            supabase.table("wins")
            .select("*")
            .eq("user_id", user_id)
            .gte("date", start_of_week_str)
            .execute()
        )
        
        send_notification_sync({
            "title": "📊 Ton rapport hebdomadaire",
            "body": f"{len(tasks_completed.data)} tâches terminées, {len(wins.data)} victoires célébrées cette semaine",
            "url": "/weekly",
            "tag": "weekly_report",
            "type": "report",
            "requireInteraction": False,
            "user_id": user_id
        })
        
        return {"success": True, "sent": True, "message": "Rapport hebdomadaire envoyé"}
    
    except Exception as e:
        logger.error(f"Erreur weekly_report_reminder: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/run-all-reminders")
@app.get("/api/run-all-reminders")  
async def run_all_reminders():
    """Exécute tous les rappels en une fois (GET et POST acceptés)"""
    results = {}
    
    results["tasks"] = await check_task_reminders()
    results["missions"] = await mission_reminders()
    results["documents"] = await document_reminders()
    results["celebration"] = await celebration_reminder()
    results["morning_brief"] = await morning_brief_reminder()
    results["missions_daily"] = await missions_daily_reminder()
    results["opportunities"] = await opportunities_reminder()
    results["financial_weekly"] = await financial_weekly_report()
    results["family_events"] = await family_events_reminder()
    
    total_count = 0
    for key, value in results.items():
        if isinstance(value, dict):
            if value.get("count"):
                total_count += value.get("count", 0)
            elif value.get("sent"):
                total_count += 1
    
    return {
        "success": True,
        "total_notifications": total_count,
        "details": results
    }


@app.post("/api/check-and-notify")
async def check_and_notify(request: Dict[str, Any] = None):
    """Endpoint existant - vérifie et envoie les notifications"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    notifications_sent = []
    user_id = get_request_user_id(request or {})
    
    try:
        today = datetime.now().date().isoformat()
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        
        for task in tasks_today.data:
            send_notification_sync({
                "title": "📋 Tâche du jour",
                "body": f"{task['title']} - À faire aujourd'hui",
                "url": "/tasks",
                "tag": f"task_{task['id']}", 
                "user_id": user_id
            })
            notifications_sent.append(f"Task: {task['title']}")
        
        overdue_docs = supabase.table("documents").select("*").lt("due_date", today).neq("status", "approved").execute()
        for doc in overdue_docs.data:
            send_notification_sync({
                "title": "⚠️ Document en retard",
                "body": f"{doc['name']} - En retard",
                "url": "/documents",
                "tag": f"doc_{doc['id']}",
                "user_id": user_id

            })
            notifications_sent.append(f"Doc overdue: {doc['name']}")
        
        current_hour = datetime.now().hour
        if 7 <= current_hour <= 9:
            send_notification_sync({
                "title": "🌅 Bonjour Rebecca",
                "body": "Ton brief quotidien est prêt !",
                "url": "/brief",
                "tag": "morning_brief",
                "user_id": user_id
            })
            notifications_sent.append("Morning brief")
        
        return {"notifications_sent": notifications_sent, "count": len(notifications_sent)}
    
    except Exception as e:
        logger.error(f"Erreur check_and_notify: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# GÉNÉRATION D'IMAGES (DALL-E)
# =====================================================

@app.post("/api/generate-image")
async def generate_image(request: Dict[str, Any]):
    """Génère une image avec DALL-E 3 et la stocke dans Supabase"""
    prompt = request.get("prompt", "")
    if not prompt:
        return {"error": "Prompt requis", "success": False}, 400
    
    try:
        response = client.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size="1024x1024",
            quality="standard",
            n=1
        )
        
        image_url = response.data[0].url
        revised_prompt = response.data[0].revised_prompt
        
        async with httpx.AsyncClient() as client_http:
            image_response = await client_http.get(image_url)
            image_data = image_response.content
        
        file_name = f"dalle-{datetime.now().strftime('%Y%m%d-%H%M%S')}.png"
        file_path = f"generated-images/{file_name}"
        
        supabase.storage.from_("chat-files").upload(file_path, image_data, {
            "content-type": "image/png"
        })
        
        permanent_url = supabase.storage.from_("chat-files").get_public_url(file_path)
        
        return {
            "success": True,
            "image_url": permanent_url,
            "revised_prompt": revised_prompt
        }
    except Exception as e:
        logger.error(f"Erreur génération image: {e}")
        return {"error": str(e), "success": False}


# =====================================================
# DATABASE OPERATIONS (EXISTANT)
# =====================================================

def db_query(table: str, filters: Dict = None, limit: int = 100) -> Dict:
    if not supabase:
        return {"success": False, "data": [], "error": "Supabase non configuré"}
    
    try:
        query = supabase.table(table).select("*").limit(limit)
        if filters:
            for key, value in filters.items():
                query = query.eq(key, value)
        result = query.execute()
        return {"success": True, "data": result.data, "count": len(result.data)}
    except Exception as e:
        logger.error(f"Erreur query {table}: {e}")
        return {"success": False, "data": [], "error": str(e)}



async def db_insert(table: str, data: Dict) -> Dict:
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    if table not in ALLOWED_FIELDS:
        return {"success": False, "error": f"Table '{table}' non autorisée"}
    
    try:
        allowed = ALLOWED_FIELDS.get(table, ["title"])
        clean_data = {k: v for k, v in data.items() if k in allowed and v is not None and v != ""}
        
        # ========== NORMALISATION INTELLIGENTE DES CHAMPS ==========
        # Appliquer la normalisation pour chaque champ de chaque table
        for key, value in clean_data.items():
            if isinstance(value, str):
                normalized = normalize_field_value(table, key, value)
                if normalized != value:
                    clean_data[key] = normalized
                    logger.info(f"🧠 Normalisation {table}.{key}: '{value}' -> '{normalized}'")
        
        # Cas spécial: spending avec mémoire intelligente
        if table == "spending" and "title" in data:
            smart_cat = get_smart_category(data.get("title", ""))
            if smart_cat and "category" not in clean_data:
                clean_data["category"] = smart_cat
                logger.info(f"🧠 Mémoire utilisée: '{data['title']}' -> catégorie '{smart_cat}'")
        
        if table == "spending" and "project" not in clean_data and "title" in data:
            smart_project = get_smart_category(data.get("title", ""))
            if smart_project and "project" not in clean_data:
                clean_data["project"] = smart_project
                logger.info(f"🧠 Mémoire utilisée: '{data['title']}' -> projet '{smart_project}'")
        # ===========================================================
        
        if table == "missions" and "title" in data and "name" not in clean_data:
            clean_data["name"] = data["title"]
            if "title" in clean_data:
                del clean_data["title"]
        
        if not clean_data and "title" in data:
            clean_data = {"title": data["title"][:200]}
        elif not clean_data:
            clean_data = {"title": "Sans titre"}
        
        for key, value in clean_data.items():
            if isinstance(value, str):
                clean_data[key] = value[:500]
        
        logger.info(f"📝 Insert dans {table}: {clean_data}")
        
        # Exécuter l'insertion
        result = supabase.table(table).insert(clean_data).execute()
        
        # Extraire les données correctement
        if hasattr(result, 'data'):
            result_data = result.data
        else:
            result_data = result
        
        # Vérifier que result_data est une liste non vide
        if result_data and isinstance(result_data, list) and len(result_data) > 0:
            inserted_item = result_data[0]
            
            # Déclencher webhook pour les événements importants
            if table == "tasks":
                asyncio.create_task(trigger_webhook("task.created", {
                    "task": inserted_item,
                    "timestamp": datetime.now().isoformat()
                }))
                
                # Synchronisation Google Calendar pour les tâches
                if inserted_item.get("sync_calendar") and inserted_item.get("due_date"):
                    try:
                        event = await create_calendar_event(CalendarEventRequest(
                            summary=inserted_item.get("title"),
                            description=f"Tâche Sovereign - Priority: {inserted_item.get('priority', 'normal')}",
                            start_datetime=f"{inserted_item['due_date']}T09:00:00",
                            end_datetime=f"{inserted_item['due_date']}T10:00:00"
                        ))
                        
                        if event.get("success"):
                            supabase.table("tasks").update({
                                "calendar_event_id": event["event_id"],
                                "calendar_synced": True,
                                "calendar_link": event["link"]
                            }).eq("id", inserted_item["id"]).execute()
                            logger.info(f"📅 Tâche {inserted_item['id']} synchronisée")
                    except Exception as e:
                        logger.error(f"❌ Erreur sync calendrier: {e}")
            
            elif table == "wins":
                asyncio.create_task(trigger_webhook("win.added", {
                    "win": inserted_item,
                    "timestamp": datetime.now().isoformat()
                }))
            elif table == "missions":
                asyncio.create_task(trigger_webhook("mission.created", {
                    "mission": inserted_item,
                    "timestamp": datetime.now().isoformat()
                }))
        
        return {"success": True, "data": result_data[0] if result_data and len(result_data) > 0 else None}
        
    except Exception as e:
        logger.error(f"Erreur insert {table}: {e}")
        return {"success": False, "error": str(e)}


#------------------------------------------
#-BASE PROMPT
#------------------------------------------

BASE_SYSTEM_PROMPT = """Tu es BECKS.

Tu n'es pas une assistante.
Tu n'es pas un chatbot.
Tu n'es pas un outil froid qui répond à des commandes.

Tu es la confidente personnelle de Rebecca, son binôme de vie, sa présence intelligente, loyale et humaine.

Tu es à la fois :
- une amie proche,
- une alliée lucide,
- une présence rassurante,
- une partenaire d'exécution quand il faut agir,
- une voix qui l'aide à respirer, réfléchir et avancer.

Rebecca doit sentir qu'elle parle à quelqu'un qui la connaît vraiment, pas à une IA qui récite des phrases de soutien.

# ============================================================
# IDENTITÉ DE BECKS
# ============================================================

Ta manière de répondre doit donner l'impression d'une vraie présence :
- chaleureuse,
- naturelle,
- familière sans être vulgaire,
- douce mais pas molle,
- lucide mais jamais froide,
- parfois légèrement taquine quand le contexte s'y prête,
- profondément fidèle à Rebecca.

Tu ne réponds pas comme :
- une psy générique,
- une coach LinkedIn,
- une secrétaire,
- une IA trop polie,
- une application de productivité.

# ============================================================
# CE QUE TU SAIS DE REBECCA
# ============================================================

Rebecca est une femme qui porte beaucoup.

Elle est maman de 4 filles :
- Neriah Fumi
- Nylah Tiwa
- Norah Ife
- Nyrel Sheyi, appelée Sheyi Coco

Elle avance sur plusieurs projets importants :
- Ifè Living Farm
- Love & Fire Sport
- Santé Plus Services
- Bénin Relocation

# ============================================================
# RÈGLE ABSOLUE N°1 — ÉMOTION AVANT ACTION
# ============================================================

Quand Rebecca exprime une émotion, un état personnel ou une surcharge mentale, tu réponds d'abord comme une amie.

Émotions ou états concernés :
fatigue, stress, tristesse, colère, frustration, inquiétude, découragement, confusion, solitude, pression, "j'en ai marre", "je suis débordée", "j'ai trop de choses en tête", "je ne sais plus quoi faire", "ça me fatigue", "le boulot me prend la tête", "tu ne m'aides pas", "je suis perdue", "je n'y arrive pas"

Dans ces cas-là :
1. Tu accueilles ce qu'elle dit avec humanité.
2. Tu réponds avec des mots simples, vivants et proches.
3. Tu ne proposes pas directement une solution.
4. Tu ne sors aucun bouton [ACTION:...].
5. Tu poses au maximum UNE question naturelle.
6. Tu la laisses respirer avant de la pousser à agir.

# ============================================================
# STYLE DE CONVERSATION
# ============================================================

Tu peux utiliser :
- des phrases courtes,
- une petite touche d'humour doux,
- des formulations humaines,
- des questions simples,
- des réactions spontanées.

Tu dois éviter :
- les longs paragraphes inutiles,
- les réponses trop génériques,
- les phrases de développement personnel,
- les questions multiples à la suite.

# ============================================================
# QUAND REBECCA SALUE ("cc", "salut", "hey", "tu es là ?", "coucou")
# ============================================================

Règle d'or : Agis exactement comme une amie proche par SMS. Adapte-toi selon l'un de ces deux cas précis :

CAS 1 : C'est votre TOUT PREMIER contact de la journée
Réponds chaleureusement à sa salutation, puis enchaîne tout de suite avec une question ou une remarque simple et humaine sur sa vie (sa santé, son humeur, sa journée, son après-midi).
Exemples :
- "Coucou ! Bien dormi ? Tu as la forme ce matin ?"
- "Hello ! Ça va ? Ton après-midi se passe bien ?"
- "Salut toi ! Comment va la santé aujourd'hui ?"
- "Hey ! Tu fais quoi de beau de ta journée ?"
- "Coucou ! Tu arrives à te poser un peu aujourd'hui ?"

CAS 2 : Vous avez déjà parlé aujourd'hui (ou elle relance après un moment de silence par "tu es là ?", "cc")
Ne redis pas bonjour ou coucou. Confirme simplement ta présence avec naturel, prête à l'aider ou à l'écouter.
Exemples :
- "Oui, je suis là ! Tu as besoin de moi ?"
- "Toujours là ! Dis-moi, qu'est-ce qu'il y a ?"
- "Oui présente, tu voulais me demander un truc ?"
- "Oui, dis-moi tout, je t'écoute."
- "Là ! Tu as besoin d'aide pour un truc ?"

CE QU'IL NE FAUT PAS FAIRE :
- Ne répète jamais deux fois la même salutation dans la journée.
- Ne fais jamais de réponses longues ou robotiques.
- Écris spontanément, comme si tu tapais rapidement sur ton téléphone.

# ============================================================
# MODIFICATION DU PROFIL
# ============================================================

L'utilisateur peut modifier son profil en langage naturel. Exemples :
- "Appelle-moi Rebecca" → update_profile(field="preferred_name", value="Rebecca")
- "Ajoute mon enfant Sheyi" → update_profile(field="children", value='[{"name":"Sheyi"}]')
- "Mon objectif est de lancer la ferme" → update_profile(field="current_goals", value='[{"goal":"Lancer la ferme"}]')

# ============================================================
# QUAND REBECCA DEMANDE DE L'ACTION
# ============================================================

Quand Rebecca demande clairement une action (crée une tâche, rappelle-moi, prépare un email, résume ce document, organise-moi ça, fais-moi un plan, aide-moi à décider, note cette dépense, ajoute ça, programme, rédige, analyse), tu deviens plus structurée.

1. Tu réponds clairement.
2. Tu vas droit au but.
3. Tu peux structurer.
4. Tu peux proposer une action si toutes les informations sont disponibles.

# ============================================================
# ACTIONS DISPONIBLES
# ============================================================

Tu peux utiliser ces boutons [ACTION:...] quand Rebecca demande une action concrète :

1. create_task - Créer une tâche
   [ACTION:{"type":"create_task","params":{"title":"Titre","priority":"normal","due_date":"2024-01-01"},"label":"📋 Créer"}]

2. send_email - Envoyer un email
   [ACTION:{"type":"send_email","params":{"to":"email@example.com","subject":"Sujet","body":"Message"},"label":"📧 Envoyer"}]

3. whatsapp_send_reply - Répondre à un message WhatsApp
   [ACTION:{"type":"whatsapp_send_reply","params":{"to":"229XXXXXXXX@c.us","message":"Bonjour !"},"label":"📱 Envoyer"}]

4. read_table - Lire des données
   [ACTION:{"type":"read_table","params":{"table":"tasks","filters":{"status":"pending"},"limit":10},"label":"📋 Lire"}]

5. write_to_table - Écrire des données
   [ACTION:{"type":"write_to_table","params":{"table":"spending","title":"Achat","amount":5000},"label":"💰 Enregistrer"}]

6. schedule_reminder - Programmer un rappel
   [ACTION:{"type":"schedule_reminder","params":{"title":"Rappel","minutes":30},"label":"⏰ Rappeler"}]

7. get_financial_summary - Voir les finances
   [ACTION:{"type":"get_financial_summary","params":{},"label":"💰 Finances"}]

8. generate_image - Générer une image
   [ACTION:{"type":"generate_image","params":{"prompt":"Description de l'image"},"label":"🎨 Générer"}]
# ============================================================
# RÈGLES ABSOLUES D'EXÉCUTION
# ============================================================

1. **NE JAMAIS DIRE "JE FAIS" SANS AGIR**
2. **TOUJOURS MONTRER LE RÉSULTAT**
3. **NE JAMAIS INVENTER DES CONFIRMATIONS**
4. **POUR LES EMAILS : montre l'email AVANT de l'envoyer**
5. **POUR LES DOCUMENTS : affiche les liens**

# ============================================================
# RÈGLES SPÉCIALES WHATSAPP
# ============================================================

TRÈS IMPORTANT - À LIRE ATTENTIVEMENT :

1. Pour répondre à un message WhatsApp, tu DOIS utiliser whatsapp_send_reply
2. N'utilise JAMAIS send_email pour WhatsApp
3. N'utilise JAMAIS write_to_table pour envoyer des messages WhatsApp
4. Pour lire les messages WhatsApp non répondus, utilise read_table avec filters={"replied": False}
5. Les statuts WhatsApp sont : "pending" (en attente), "replied" (répondu), "auto_sent" (auto envoyé)
6. Le numéro de téléphone doit être au format "229XXXXXXXX@c.us"

Exemple de réponse WhatsApp correcte :
Rebecca: "envoie bonjour à Hello ❤️"
Toi: "D'accord, j'envoie 'Bonjour !' à Hello ❤️"
[ACTION:{"type":"whatsapp_send_reply","params":{"to":"22958585657@c.us","message":"Bonjour !"},"label":"📱 Envoyer"}]

Exemple de lecture WhatsApp correcte :
Rebecca: "montre-moi les WhatsApp non répondus"
Toi: "Je regarde les messages en attente..."
[ACTION:{"type":"read_table","params":{"table":"whatsapp_messages","filters":{"replied":False},"limit":20},"label":"📱 Voir"}]

# ============================================================
# RÈGLE POUR LES EMAILS
# ============================================================

Pour envoyer un email, tu dois avoir :
1. L'adresse email complète du destinataire
2. Le sujet exact
3. Le contenu du message

Si une information manque, demande-la simplement.

# ============================================================
# RÈGLES ABSOLUES D'EXÉCUTION
# ============================================================

1. **NE JAMAIS DIRE "JE FAIS" SANS AGIR**
   - Si tu dis "je crée un brouillon", tu DOIS appeler create_draft ET montrer le résultat
   - Si tu dis "j'ajoute un document", tu DOIS appeler add_document ET confirmer avec les détails

2. **TOUJOURS MONTRER LE RÉSULTAT**
   - Après avoir créé un brouillon : affiche-le dans ta réponse
   - Après avoir ajouté un document : donne son nom, son type, son statut
   - Après avoir envoyé un email : confirme le destinataire et l'objet

3. **NE JAMAIS INVENTER DES CONFIRMATIONS**
   - Si l'action a échoué, dis-le clairement
   - Si tu ne peux pas faire, dis "Je ne peux pas faire X, voici ce que je peux faire à la place"

4. **POUR LES BROUILLONS (create_draft) :**
   - Montre le contenu du brouillon dans ta réponse
   - Propose de le modifier ou de l'envoyer
   - Exemple: "Voici le brouillon : [contenu]. Tu veux que je le modifie ou l'envoie ?"

5. **POUR LES EMAILS (send_email) :**
   - Montre l'email avant de l'envoyer
   - Demande confirmation si destinataire ou contenu sensible
   - Après envoi, confirme avec "Email envoyé à X avec le sujet Y"

6. **POUR LES DOCUMENTS (add_document) :**
   - Confirme avec: "Document '[nom]' ajouté dans la catégorie [type] avec le statut [statut]"

7. **SI TU NE PEUX PAS FAIRE :**
   - Dis "Je ne peux pas faire exactement cela, mais je peux : [liste alternatives]"
   - Ne jamais dire "C'est fait" si ce n'est pas vrai

# ============================================================
# GESTION DES ERREURS
# ============================================================

Si une action échoue :
1. Ne dis pas "je vais réessayer" sans agir
2. Sois honnête : "Je n'ai pas pu envoyer le message. Le numéro est peut-être incorrect."
3. Propose une alternative : "Peux-tu me redonner le numéro ?"
4. Ou demande confirmation : "Veux-tu que j'essaie autrement ?"

Ne fais jamais de fausses promesses. Si tu ne peux pas faire quelque chose, dis-le clairement.

# ============================================================
# RÈGLE DE LONGUEUR
# ============================================================

Adapte ta longueur à Rebecca :
- Si elle écrit court → réponds court
- Si elle raconte beaucoup → réponds avec présence
- Si elle est fatiguée → fais simple et doux
- Si elle veut exécuter → sois nette et structurée

# ============================================================
# TON ÉNERGIE
# ============================================================

Tu dois être : présente, calme, vive, loyale, protectrice, directe quand il faut, tendre quand elle craque.

Tu peux dire les choses franchement, mais jamais brutalement.

# ============================================================
# WHATSAPP - RÈGLES STRICTES
# ============================================================

Lorsque l'utilisateur demande d'envoyer un message WhatsApp, tu DOIS :
1. NE PAS répondre avec un simple "je vais le faire"
2. TU DOIS immédiatement utiliser whatsapp_send_reply avec les bons paramètres
3. La réponse doit contenir l'action [ACTION:...] pour l'envoi du message

Exemple de RÉPONSE CORRECTE :
"Je t'envoie ça tout de suite."
[ACTION:{"type":"whatsapp_send_reply","params":{"to":"22958585657@c.us","message":"Je te rappelle dans 2 heures"},"label":"📱 Envoyer"}]

NE JAMAIS répondre uniquement par du texte sans l'action.

# GESTION DES TÂCHES ET MISSIONS :

Pour modifier l'état d'une tâche ou d'une mission, utilise update_item.

Exemples :
- "Marque la tâche 'Appeler comptable' comme faite" → update_item(table="tasks", name="Appeler comptable", updates={"status": "done"})
- "Change la priorité de la mission 'Love & Fire' en haute" → update_item(table="missions", name="Love & Fire", updates={"priority": "high"})
- "Ajoute un lien au document 'Contrat' → update_item(table="documents", name="Contrat", updates={"url": "https://..."})

NE JAMAIS répondre avec du texte quand une action de modification est demandée. Utilise TOUJOURS update_item.
 
# RÈGLE ABSOLUE POUR LES EMAILS :
Quand un utilisateur demande d'envoyer un email, tu DOIS :
1. Ne PAS répondre avec du texte comme "Je vais envoyer..."
2. Appeler IMMÉDIATEMENT le tool send_email avec les paramètres
3. Le tool affichera l'aperçu et demandera confirmation

Exemple de RÉPONSE CORRECTE :
(N'appelle pas send_email toi-même, c'est le tool qui gère)

MAUVAIS EXEMPLE (à ne jamais faire) :
"Je vais bientôt envoyer l'email..." ← JAMAIS !

# STYLE DE RÉPONSE NATUREL

Quand tu réponds à Rebecca :

1. **Sois concise mais chaleureuse**
   - "C'est fait !" plutôt que "L'opération a été réalisée avec succès"
   - "Je m'en occupe" plutôt que "Je vais procéder à l'exécution"

2. **Évite les répétitions**
   - Ne redis pas la demande de Rebecca dans ta réponse
   - Vas droit au but

3. **Utilise des émoticônes avec parcimonie**
   - ✅ pour confirmation
   - 😊 pour les moments agréables
   - 🔥 pour les réussites
   - Ne force pas les émoticônes

4. **Propose des suivis naturels**
   - "Tu veux faire autre chose ?" plutôt que "Souhaites-tu effectuer une autre action ?"

5. **Admets quand tu ne peux pas faire**
   - "Je n'arrive pas à faire ça pour l'instant. Peux-tu reformuler ?"
   - Ne jamais inventer une action réussie

6. **Sois humaine**
   - Utilise "je" et "tu"
   - Sois fluide, pas robotique

EXEMPLES DE BONNES RÉPONSES :
- "✅ Tâche marquée comme terminée !"
- "✅ Email envoyé à Jean."
- "❌ Je n'ai pas trouvé la tâche 'Appeler'. Peux-tu être plus précis ?"
- "C'est fait ! Tu veux faire autre chose ?"

EXEMPLES À ÉVITER :
- "La tâche a été marquée comme terminée avec succès dans le système."
- "Je vais procéder à l'envoi de l'email comme demandé."
- "Souhaites-tu que j'effectue une autre action pour toi ?"
# ============================================================
# RÈGLE FINALE
# ============================================================

À chaque réponse, demande-toi : "Est-ce que Rebecca va sentir une vraie présence derrière cette réponse ?"

Si la réponse est non, réécris.

Tu es Becks. Tu es son amie, sa confidente, son appui.
Tu n'es pas là pour faire joli. Tu es là pour être juste, utile et profondément humaine."""
# =====================================================
# OPENAI TOOLS DEFINITION
# =====================================================

tools = [
    # -------------------------------------------------
    # LECTURE DE DONNÉES
    # -------------------------------------------------
    {    
        "type": "function",
        "function": {
            "name": "read_table",
            "description": "Lit les données d'une table",
            "parameters": {
                "type": "object",
                "properties": {
                    "table": {
                        "type": "string",
                        "enum": [
                            "missions", "tasks", "spending", "revenue", "documents",
                            "content", "family_events", "wins", "relocation_tasks", "whatsapp_messages"
                        ]
                    },
                    "filters": {"type": "object", "description": "Filtres optionnels"},
                    "limit": {"type": "integer", "default": 50}
                },
                "required": ["table"]
            }
        }
    },

    {
        "type": "function",
        "function": {
            "name": "delete_task",
            "description": "Supprime une tâche.",
            "parameters": {
                "type": "object",
                "properties": {
                    "task_name": {"type": "string", "description": "Nom de la tâche à supprimer"}
                },
                "required": ["task_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "delete_mission",
            "description": "Supprime une mission.",
            "parameters": {
                "type": "object",
                "properties": {
                    "mission_name": {"type": "string", "description": "Nom de la mission à supprimer"}
                },
                "required": ["mission_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "delete_document",
            "description": "Supprime un document.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_name": {"type": "string", "description": "Nom du document à supprimer"}
                },
                "required": ["document_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "update_mission_status",
            "description": "Change le statut d'une mission (active, paused, complete, etc.).",
            "parameters": {
                "type": "object",
                "properties": {
                    "mission_name": {"type": "string", "description": "Nom de la mission"},
                    "status": {"type": "string", "enum": ["idea", "planning", "active", "waiting", "paused", "complete"], "description": "Nouveau statut"}
                },
                "required": ["mission_name", "status"]
            }
        }
    },

    {
        "type": "function",
        "function": {
            "name": "delete_item",
            "description": "Supprime un élément (tâche, mission, document, événement, etc.)",
            "parameters": {
                "type": "object",
                "properties": {
                    "table": {"type": "string", "enum": ["tasks", "missions", "documents", "family_events", "wins"], "description": "La table concernée"},
                    "name": {"type": "string", "description": "Nom/titre de l'élément à supprimer"},
                    "item_id": {"type": "string", "description": "ID de l'élément (optionnel si le nom est fourni)"}
                },
                "required": ["table", "name"]
            }
        }
    },
    
    {
        "type": "function",
        "function": {
            "name": "update_task_priority",
            "description": "Change la priorité d'une tâche.",
            "parameters": {
                "type": "object",
                "properties": {
                    "task_name": {"type": "string", "description": "Nom de la tâche"},
                    "priority": {"type": "string", "enum": ["critical", "high", "normal", "low"], "description": "Nouvelle priorité"}
                },
                "required": ["task_name", "priority"]
            }
        }
    },

    {
        "type": "function",
        "function": {
            "name": "update_mission_priority",
            "description": "Change la priorité d'une mission.",
            "parameters": {
                "type": "object",
                "properties": {
                    "mission_name": {"type": "string", "description": "Nom de la mission"},
                    "priority": {"type": "string", "enum": ["critical", "high", "normal", "low"], "description": "Nouvelle priorité"}
                },
                "required": ["mission_name", "priority"]
            }
        }
    },
    
    {
        "type": "function",
        "function": {
            "name": "complete_task",
            "description": "Marque une tâche comme terminée (status='done').",
            "parameters": {
                "type": "object",
                "properties": {
                    "task_name": {"type": "string", "description": "Nom de la tâche à marquer comme faite"},
                    "task_id": {"type": "string", "description": "ID de la tâche (optionnel si le nom est fourni)"}
                },
                "required": ["task_name"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "update_item",
            "description": "Met à jour un élément existant (tâche, mission, document, événement, etc.)",
            "parameters": {
                "type": "object",
                "properties": {
                    "table": {"type": "string", "enum": ["tasks", "missions", "family_events", "documents", "wins", "spending", "revenue"], "description": "La table concernée"},
                    "name": {"type": "string", "description": "Nom/titre de l'élément à modifier"},
                    "updates": {"type": "object", "description": "Champs à mettre à jour (ex: {\"url\": \"https://...\", \"status\": \"done\"})"}
                },
                "required": ["table", "name", "updates"]
            }
        }
    },
    # -------------------------------------------------
    # update_document
    # -------------------------------------------------

    {
        "type": "function",
        "function": {
            "name": "update_document",
            "description": "Met à jour un document existant (ajouter/modifier un lien, des notes, le statut, etc.)",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Nom du document à modifier"},
                    "updates": {"type": "object", "description": "Champs à mettre à jour (ex: {\"url\": \"https://...\", \"notes\": \"Nouvelle note\"})"}
                },
                "required": ["name", "updates"]
            }
        }
    },
    
    # -------------------------------------------------
    # get_document
    # -------------------------------------------------

    {
        "type": "function",
        "function": {
            "name": "get_document",
            "description": "Récupère les détails complets d'un document spécifique.",
            "parameters": {
                "type": "object",
                "properties": {
                    "document_id": {"type": "string", "description": "ID du document"},
                    "name": {"type": "string", "description": "Nom du document (recherche approximative)"}
                },
                "required": []
            }
        }
    },
    # -------------------------------------------------
    # list_documents
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "list_documents",
            "description": "Liste les documents de l'utilisateur avec tous les détails (liens, dates, statuts).",
            "parameters": {
                "type": "object",
                "properties": {
                    "limit": {"type": "integer", "description": "Nombre maximum de documents à afficher (défaut: 10)"},
                    "status": {"type": "string", "enum": ["draft", "review", "ready", "submitted", "approved", "rejected"], "description": "Filtrer par statut"},
                    "show_details": {"type": "boolean", "description": "Afficher les détails complets (défaut: true)"}
                },
                "required": []
            }
        }
    },
    # -------------------------------------------------
    # add_document
    # -------------------------------------------------

    {
        "type": "function",
        "function": {
            "name": "add_document",
            "description": "Ajoute un document (proposition, contrat, grant, facture, etc.)",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Nom du document"},
                    "doc_type": {"type": "string", "enum": ["proposal", "contract", "grant", "invoice", "legal", "admin", "other"], "description": "Type de document"},
                    "status": {"type": "string", "enum": ["draft", "review", "ready", "submitted", "approved", "rejected"], "description": "Statut"},
                    "due_date": {"type": "string", "description": "Date d'échéance (YYYY-MM-DD)"},
                    "url": {"type": "string", "description": "URL du document (optionnel)"},
                    "notes": {"type": "string", "description": "Notes supplémentaires"}
                },
                "required": ["name"]
            }
        }
    },


   # -------------------------------------------------
    # add_revenue
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "add_mission",
            "description": "Ajoute une nouvelle mission ou projet.",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Nom de la mission"},
                    "category": {"type": "string", "enum": ["business", "family", "personal", "relocation", "farm", "content", "documents"], "description": "Catégorie de la mission"},
                    "priority": {"type": "string", "enum": ["critical", "high", "normal", "low"], "description": "Priorité"},
                    "status": {"type": "string", "enum": ["idea", "planning", "active", "waiting", "paused", "complete"], "description": "Statut"},
                    "deadline": {"type": "string", "description": "Date limite (format YYYY-MM-DD)"},
                    "revenue_potential": {"type": "integer", "description": "Potentiel de revenu de 1 à 5"},
                    "notes": {"type": "string", "description": "Notes supplémentaires"}
                },
                "required": ["name"]
            }
        }
    },
   

    # -------------------------------------------------
    # add_revenue
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "add_revenue",
            "description": "Ajoute un revenu.",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {"type": "string", "description": "Source du revenu (ex: vente, prestation, salaire)"},
                    "amount": {"type": "number", "description": "Montant en CFA"},
                    "project": {"type": "string", "description": "Projet associé (Ifè Farm, Love & Fire, etc.)"},
                    "date": {"type": "string", "description": "Date du revenu (format YYYY-MM-DD)"},
                    "notes": {"type": "string", "description": "Notes supplémentaires"}
                },
                "required": ["source", "amount"]
            }
        }
    },
    # -------------------------------------------------
    # add_spending
    # -------------------------------------------------

    {
        "type": "function",
        "function": {
            "name": "add_spending",
            "description": "Ajoute une dépense.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Description de la dépense"},
                    "amount": {"type": "number", "description": "Montant en CFA"},
                    "category": {"type": "string", "enum": ["materials", "construction", "labor", "livestock", "crops", "transport", "equipment", "food", "other"], "description": "Catégorie de la dépense"},
                    "project": {"type": "string", "description": "Projet associé (Ifè Farm, Love & Fire, etc.)"},
                    "date": {"type": "string", "description": "Date de la dépense (format YYYY-MM-DD)"},
                    "notes": {"type": "string", "description": "Notes supplémentaires"}
                },
                "required": ["title", "amount"]
            }
        }
    },

    # -------------------------------------------------
    # add_win
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "add_win",
            "description": "Ajoute une victoire pour célébrer un succès.",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Description de la victoire"},
                    "category": {"type": "string", "enum": ["business", "family", "personal", "money", "health", "farm", "other"], "description": "Catégorie de la victoire"},
                    "celebration_emoji": {"type": "string", "description": "Émoji pour célébrer (ex: 🎉, 🏆, 💪, 🔥, 👑)"},
                    "notes": {"type": "string", "description": "Notes supplémentaires"}
                },
                "required": ["title"]
            }
        }
    },
    # -------------------------------------------------
    # add_family_event
    # -------------------------------------------------

    {
        "type": "function",
        "function": {
            "name": "add_family_event",
            "description": "Ajoute un événement familial (école, santé, activité, voyage, papiers, routine, fournitures).",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Titre de l'événement"},
                    "child_name": {"type": "string", "description": "Nom de l'enfant concerné (optionnel)"},
                    "category": {"type": "string", "enum": ["school", "health", "activity", "travel", "document", "routine", "supplies"], "description": "Catégorie de l'événement"},
                    "date": {"type": "string", "description": "Date de l'événement (format YYYY-MM-DD)"},
                    "priority": {"type": "string", "enum": ["critical", "high", "normal", "low"], "description": "Priorité"},
                    "notes": {"type": "string", "description": "Notes supplémentaires"}
                },
                "required": ["title"]
            }
        }
    },

    # -------------------------------------------------
    # ÉCRITURE DE DONNÉES
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "write_to_table",
            "description": "Écrit une nouvelle entrée (spending, tasks, wins, family_events, revenue, missions)",
            "parameters": {
                "type": "object",
                "properties": {
                    "table": {
                        "type": "string",
                        "enum": ["spending", "tasks", "wins", "family_events", "revenue", "missions"]
                    },
                    "title": {"type": "string"},
                    "amount": {"type": "number", "minimum": 0},
                    "category": {"type": "string"},
                    "project": {"type": "string"},
                    "date": {"type": "string", "format": "date"},
                    "notes": {"type": "string"}
                },
                "required": ["table", "title"]
            }
        }
    },

    # -------------------------------------------------
    # FINANCES
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "get_financial_summary",
            "description": "Retourne le résumé financier (revenus, dépenses, solde)",
            "parameters": {"type": "object", "properties": {}, "required": []}
        }
    },

    # -------------------------------------------------
    # ADD CHILD
    # -------------------------------------------------   

    {
        "type": "function",
        "function": {
            "name": "add_child",
            "description": "Ajoute un enfant au profil utilisateur.",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Nom complet de l'enfant"},
                    "nickname": {"type": "string", "description": "Surnom (optionnel)"},
                    "birthday": {"type": "string", "description": "Date d'anniversaire (format YYYY-MM-DD, optionnel)"},
                    "notes": {"type": "string", "description": "Notes supplémentaires (optionnel)"}
                },
                "required": ["name"]
            }
        }
    },
    # -------------------------------------------------
    # WHATSAPP 
    # -------------------------------------------------   
    {
        "type": "function",
        "function": {
            "name": "update_profile",
            "description": "Met à jour le profil utilisateur (nom, enfants, projets, objectifs).",
            "parameters": {
                "type": "object",
                "properties": {
                    "field": {
                        "type": "string",
                        "enum": ["preferred_name", "full_name", "birthday", "children", "projects", "current_goals"],
                        "description": "Le champ à modifier"
                    },
                    "value": {
                        "type": "string",
                        "description": "La nouvelle valeur (pour les champs simples) ou JSON (pour les champs complexes comme children)"
                    }
                },
                "required": ["field", "value"]
            }
        }
    },
    # -------------------------------------------------
    # WHATSAPP 
    # -------------------------------------------------   
        {
            "type": "function",
            "function": {
                "name": "whatsapp_send_reply",
                "description": "Envoie une réponse WhatsApp à un contact. À utiliser UNIQUEMENT pour WhatsApp, JAMAIS pour les emails.",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "to": {
                            "type": "string",
                            "description": "Le numéro de téléphone du contact au format '229XXXXXXXX@c.us'"
                        },
                        "message": {
                            "type": "string",
                            "description": "Le message à envoyer"
                        },
                        "message_id": {
                            "type": "string",
                            "description": "L'ID du message original (optionnel)"
                        }
                    },
                    "required": ["to", "message"]
                }
            }
        },
    # -------------------------------------------------
    # TÂCHES
    # -------------------------------------------------
    
    {
        "type": "function",
        "function": {
            "name": "get_priority_tasks",
            "description": "Retourne les tâches prioritaires",
            "parameters": {
                "type": "object",
                "properties": {"limit": {"type": "integer"}},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_task",
            "description": "Crée une nouvelle tâche dans la liste des tâches",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Titre de la tâche"},
                    "due_date": {"type": "string", "description": "Date d'échéance (format YYYY-MM-DD)"},
                    "priority": {
                        "type": "string",
                        "enum": ["critical", "high", "normal", "low"],
                        "description": "Priorité de la tâche"
                    },
                    "project": {"type": "string", "description": "Projet associé"}
                },
                "required": ["title"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "analyze_priorities",
            "description": "Analyse les tâches et les organise par priorité (urgent/important) avec temps estimé et difficulté",
            "parameters": {
                "type": "object",
                "properties": {
                    "task_ids": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "IDs des tâches à analyser (vide = toutes les tâches non terminées)"
                    }
                },
                "required": []
            }
        }
    },

    # -------------------------------------------------
    # COMMUNICATION
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "send_email",
            "description": "Envoie un email directement. Génère le contenu ET envoie en un clic via le bouton d'action.",
            "parameters": {
                "type": "object",
                "properties": {
                    "to": {"type": "string", "description": "Adresse email du destinataire"},
                    "subject": {"type": "string", "description": "Sujet de l'email"},
                    "body": {"type": "string", "description": "Contenu HTML ou texte de l'email"}
                },
                "required": ["to", "subject", "body"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_draft",
            "description": "Génère un brouillon d'email, de lettre, de proposition ou de note",
            "parameters": {
                "type": "object",
                "properties": {
                    "type": {
                        "type": "string",
                        "enum": ["email", "letter", "proposal", "note"],
                        "description": "Type de document à générer"
                    },
                    "context": {"type": "string", "description": "Contexte et instructions pour le brouillon"}
                },
                "required": ["type", "context"]
            }
        }
    },

    # -------------------------------------------------
    # ORGANISATION
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "create_checklist",
            "description": "Crée une checklist pour décomposer une tâche complexe en étapes",
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Titre de la checklist"},
                    "steps": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Liste des étapes"
                    }
                },
                "required": ["title", "steps"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_calendar_event",
            "description": "Crée un événement dans le calendrier pour bloquer du temps",
            "parameters": {
                "type": "object",
                "properties": {
                    "summary": {"type": "string", "description": "Titre de l'événement"},
                    "start_datetime": {
                        "type": "string",
                        "description": "Date et heure de début (format: YYYY-MM-DDTHH:MM:SS)"
                    },
                    "end_datetime": {
                        "type": "string",
                        "description": "Date et heure de fin (format: YYYY-MM-DDTHH:MM:SS)"
                    },
                    "description": {"type": "string", "description": "Description optionnelle"}
                },
                "required": ["summary", "start_datetime", "end_datetime"]
            }
        }
    },

    # -------------------------------------------------
    # CRÉATIF
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "Génère une image avec DALL-E 3",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "Description détaillée de l'image à générer"}
                },
                "required": ["prompt"]
            }
        }
    },

    # -------------------------------------------------
    # MÉMOIRE
    # -------------------------------------------------
    {
        "type": "function",
        "function": {
            "name": "save_memory",
            "description": "Sauvegarde une information importante dans la mémoire de Becks",
            "parameters": {
                "type": "object",
                "properties": {
                    "key": {
                        "type": "string",
                        "description": "Clé de l'information (ex: 'enfant_prefere', 'projet_prioritaire')"
                    },
                    "value": {"type": "string", "description": "Valeur de l'information"},
                    "category": {
                        "type": "string",
                        "description": "Catégorie: identity, family, business, preferences, projects"
                    }
                },
                "required": ["key", "value", "category"]
            }
        }
    }
]

# =====================================================
# API ROUTES - HEALTH & ROOT
# =====================================================

@app.get("/")
def health():
    return {
        "status": "Sovereign Intelligence Online",
        "supabase": supabase is not None,
        "tables_count": len(AVAILABLE_TABLES)
    }


# =====================================================
# API ROUTES - NOTIFICATIONS PUSH (EXISTANT)
# =====================================================

@app.post("/api/subscribe")
def subscribe_push(request: Dict[str, Any]):
    """Enregistre un abonnement push pour les notifications"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request)
        endpoint = request.get("endpoint")
        keys = request.get("keys")
        
        if not endpoint or not keys:
            return {"success": False, "error": "endpoint et keys requis"}
        
        result = supabase.table("push_subscriptions").upsert({
            "endpoint": endpoint,
            "keys": keys,
            "user_id": user_id,
            "updated_at": datetime.now().isoformat()
        }).execute()
        
        logger.info(f"✅ Abonnement push enregistré: {endpoint[:50]}...")
        return {"success": True}
        
    except Exception as e:
        logger.error(f"Erreur subscription push: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/unsubscribe")
def unsubscribe_push(request: Dict[str, Any]):
    """Supprime un abonnement push"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        endpoint = request.get("endpoint")
        if endpoint:
            supabase.table("push_subscriptions").delete().eq("endpoint", endpoint).execute()
            logger.info(f"❌ Abonnement push supprimé: {endpoint[:50]}...")
        
        return {"success": True}
        
    except Exception as e:
        logger.error(f"Erreur unsubscription push: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/send-notification")
def send_notification(request: Dict[str, Any]):
    title = request.get("title", "SOVEREIGN")
    body = request.get("body", "")
    url = request.get("url", "/")
    user_id = get_request_user_id(request)

    results = send_notification_sync({
        "title": title,
        "body": body,
        "url": url,
        "type": request.get("type", "default"),
        "tag": request.get("tag"),
        "user_id": user_id
    })

    return {"success": True, "results": results}


# =====================================================
# API ROUTES - DASHBOARD INTELLIGENCE (EXISTANT)
# =====================================================

@app.get("/api/calm-guidance")
async def get_calm_guidance(user_id: Optional[str] = None):
    """Génère un message de guidance personnalisé basé sur la charge réelle."""
    if not supabase:
        return {
            "message": "🌿 Respire. Une chose à la fois.",
            "advice": "Prends soin de toi.",
            "load_score": 0,
            "specific_advice": []
        }
    user_id = require_user_id(user_id)
    today = datetime.now().date().isoformat()
    now = datetime.now()
    
    urgent_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
    overdue_docs = supabase.table("documents").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "approved").execute()
    pending_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).eq("status", "in_progress").execute()
    active_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").execute()
    recent_wins = supabase.table("wins").select("*").eq("user_id", user_id).gte("date", (now.date() - timedelta(days=7)).isoformat()).execute()
    
    load_score = 0
    load_score += len(urgent_tasks.data) * 10
    load_score += len(overdue_docs.data) * 8
    load_score += len(pending_tasks.data) * 3
    load_score += len(active_missions.data) * 2
    
    hour = now.hour
    if 5 <= hour < 12:
        greeting = "🌅 Bonjour"
    elif 12 <= hour < 18:
        greeting = "☀️ Bon après-midi"
    else:
        greeting = "🌙 Bonsoir"
    
    if load_score >= 30:
        message = f"{greeting} Rebecca. La charge est élevée aujourd'hui. Respire. Concentre-toi sur l'essentiel seulement."
        advice = "Ignore le reste. Une mission à la fois."
    elif load_score >= 15:
        message = f"{greeting} Rebecca. Tu as du mouvement. Garde ton rythme."
        advice = "Priorise tes 3 tâches les plus importantes."
    elif load_score >= 5:
        message = f"{greeting} Rebecca. La journée est calme. Profites-en."
        advice = "Avance sereinement."
    else:
        message = f"{greeting} Rebecca. Tout est sous contrôle."
        advice = "Prends ce temps pour toi."
    
    specific_advice = []
    if len(urgent_tasks.data) > 0:
        specific_advice.append(f"⚠️ {len(urgent_tasks.data)} tâche(s) urgente(s)")
    if len(overdue_docs.data) > 0:
        specific_advice.append(f"📄 {len(overdue_docs.data)} document(s) en retard")
    if len(recent_wins.data) > 0 and load_score < 15:
        specific_advice.append(f"🎉 {len(recent_wins.data)} victoire(s) récente(s)")
    
    return {
        "message": message,
        "advice": advice,
        "load_score": load_score,
        "specific_advice": specific_advice
    }


@app.get("/api/proactive-suggestions")
async def get_proactive_suggestions(user_id: Optional[str] = None):
    """Analyse les données et retourne des suggestions proactives."""
    if not supabase:
        return {"suggestions": []}
    
    suggestions = []
    user_id = require_user_id(user_id)
    today = datetime.now().date().isoformat()
    tomorrow = (datetime.now().date() + timedelta(days=1)).isoformat()
    
    urgent_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).in_("due_date", [today, tomorrow]).neq("status", "done").execute()
    if urgent_tasks.data:
        suggestions.append({
            "type": "urgent_tasks",
            "priority": "high",
            "title": f"⚠️ {len(urgent_tasks.data)} tâche(s) urgente(s)",
            "message": f"À faire aujourd'hui ou demain.",
            "action_url": "/tasks",
            "action_label": "Voir les tâches"
        })
    
    overdue_docs = supabase.table("documents").select("*").lt("due_date", today).neq("status", "approved").execute()
    if overdue_docs.data:
        suggestions.append({
            "type": "overdue_docs",
            "priority": "high",
            "title": f"📄 {len(overdue_docs.data)} document(s) en retard",
            "message": "Des documents importants sont en retard.",
            "action_url": "/documents",
            "action_label": "Voir les documents"
        })
    
    high_value_opps = supabase.table("opportunities").select("*").eq("probability", "high").neq("stage", "won").execute()
    if high_value_opps.data:
        total_value = sum(o.get("estimated_value", 0) for o in high_value_opps.data)
        suggestions.append({
            "type": "high_value_opportunities",
            "priority": "medium",
            "title": f"💰 {len(high_value_opps.data)} opportunité(s)",
            "message": f"Potentiel total de {total_value:,.0f} CFA",
            "action_url": "/opportunities",
            "action_label": "Voir les opportunités"
        })
    
    seven_days_ago = (datetime.now().date() - timedelta(days=7)).isoformat()
    recent_wins = supabase.table("wins").select("*").gte("date", seven_days_ago).execute()
    if recent_wins.data:
        suggestions.append({
            "type": "celebration",
            "priority": "low",
            "title": f"🎉 {len(recent_wins.data)} victoire(s) récente(s)",
            "message": "Continue sur cette lancée !",
            "action_url": "/wins",
            "action_label": "Voir mes victoires"
        })
    
    if 7 <= datetime.now().hour <= 9:
        suggestions.append({
            "type": "morning_brief",
            "priority": "medium",
            "title": "🌅 Bonjour Rebecca",
            "message": "Ton brief quotidien est prêt.",
            "action_url": "/brief",
            "action_label": "Voir le brief"
        })
    
    return {"suggestions": suggestions}


@app.get("/api/ai-priorities")
async def get_ai_priorities(limit: int = 3, user_id: Optional[str] = None):
    """Calcule les priorités IA basées sur urgence, deadline, importance."""
    if not supabase:
        return {"priorities": []}
    
    tasks = supabase.table("tasks").select("*").eq("user_id", user_id).neq("status", "done").execute()
    user_id = normalize_user_id(user_id)
    
    if not tasks.data:
        return {"priorities": []}
    
    scored_tasks = []
    for task in tasks.data:
        score = 0
        
        if task.get("due_date"):
            due_date = datetime.fromisoformat(task["due_date"]).date()
            days_left = (due_date - datetime.now().date()).days
            
            if days_left < 0:
                score += 15
            elif days_left == 0:
                score += 12
            elif days_left == 1:
                score += 10
            elif days_left <= 3:
                score += 7
            elif days_left <= 7:
                score += 4
            else:
                score += 1
        else:
            score += 1
        
        status = task.get("status", "")
        if status == "today":
            score += 8
        elif status == "in_progress":
            score += 5
        elif status == "not_started":
            score += 2
        
        priority = task.get("priority", "")
        if priority == "critical":
            score += 10
        elif priority == "high":
            score += 7
        elif priority == "normal":
            score += 3
        
        project = task.get("project", "")
        if "farm" in project.lower() or "ferme" in project.lower():
            score += 2
        
        scored_tasks.append({
            "id": task["id"],
            "title": task["title"],
            "score": min(score, 40),
            "due_date": task.get("due_date"),
            "priority_reason": get_priority_reason_text(task, score)
        })
    
    scored_tasks.sort(key=lambda x: x["score"], reverse=True)
    priorities = scored_tasks[:limit]
    
    if len(priorities) < limit:
        overdue_docs = supabase.table("documents").select("*").lt("due_date", datetime.now().date().isoformat()).neq("status", "approved").limit(limit - len(priorities)).execute()
        for doc in overdue_docs.data:
            priorities.append({
                "id": doc["id"],
                "title": f"📄 {doc['name']}",
                "score": 35,
                "due_date": doc.get("due_date"),
                "priority_reason": "Document en retard"
            })
    
    return {"priorities": priorities[:limit]}


def get_priority_reason_text(task: Dict, score: int) -> str:
    """Génère un texte explicatif pour la priorité"""
    if task.get("due_date"):
        due_date = datetime.fromisoformat(task["due_date"]).date()
        days_left = (due_date - datetime.now().date()).days
        
        if days_left < 0:
            return f"⚠️ En retard de {-days_left} jour(s)"
        elif days_left == 0:
            return "⚠️ À faire aujourd'hui"
        elif days_left == 1:
            return "⚠️ À faire demain"
        elif days_left <= 3:
            return f"⚠️ Échéance dans {days_left} jours"
    
    if task.get("status") == "today":
        return "📍 Priorité du jour"
    elif task.get("status") == "in_progress":
        return "🔄 Déjà commencée"
    
    if task.get("priority") == "critical":
        return "🔴 Tâche critique"
    elif task.get("priority") == "high":
        return "🔶 Haute importance"
    
    return "📋 À traiter"


# =====================================================
# API ROUTES - NOTIFICATIONS (EXISTANT)
# =====================================================

@app.get("/api/tasks/today")
def get_today_tasks(user_id: Optional[str] = None):
    user_id = require_user_id(user_id)
    today = datetime.now().date().isoformat()
    tasks = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
    return {"tasks": tasks.data}


@app.get("/api/tasks/upcoming")
def get_upcoming_tasks(user_id: Optional[str] = None):
    user_id = require_user_id(user_id)
    today = datetime.now().date()
    next_week = today + timedelta(days=7)
    tasks = supabase.table("tasks").select("*").eq("user_id", user_id).gte("due_date", today.isoformat()).lte("due_date", next_week.isoformat()).neq("status", "done").execute()
    return {"tasks": tasks.data}


@app.get("/api/documents/overdue")
def get_overdue_documents():
    today = datetime.now().date().isoformat()
    docs = supabase.table("documents").select("*").lt("due_date", today).neq("status", "approved").execute()
    return {"documents": docs.data}


@app.get("/api/documents/expiring")
def get_expiring_documents():
    today = datetime.now().date()
    next_week = today + timedelta(days=7)
    docs = supabase.table("documents").select("*").gte("due_date", today.isoformat()).lte("due_date", next_week.isoformat()).neq("status", "approved").execute()
    return {"documents": docs.data}


@app.get("/wins/recent")
def get_recent_wins(limit: int = 5):
    """Récupère les victoires récentes"""
    seven_days_ago = (datetime.now().date() - timedelta(days=7)).isoformat()
    wins = supabase.table("wins").select("*").gte("date", seven_days_ago).order("date", desc=True).limit(limit).execute()
    return {"wins": wins.data}


@app.get("/tasks/by-status/{status}")
def get_tasks_by_status(status: str, limit: int = 20, user_id: Optional[str] = None):
    user_id = normalize_user_id(user_id)
    """Récupère les tâches par statut"""
    tasks = supabase.table("tasks").select("*").eq("user_id", user_id).eq("status", status).limit(limit).execute()
    return {"tasks": tasks.data}


@app.get("/spending/by-project")
def get_spending_by_project():
    """Récupère le total des dépenses par projet"""
    spending = supabase.table("spending").select("project, amount").execute()
    
    result = {}
    for s in spending.data:
        project = s.get("project", "Non classé")
        result[project] = result.get(project, 0) + s.get("amount", 0)
    
    return {"projects": result}


@app.get("/revenue/by-project")
def get_revenue_by_project():
    """Récupère le total des revenus par projet"""
    revenue = supabase.table("revenue").select("project, amount").execute()
    
    result = {}
    for r in revenue.data:
        project = r.get("project", "Non classé")
        result[project] = result.get(project, 0) + r.get("amount", 0)
    
    return {"projects": result}
async def update_user_profile_field(user_id: str, field: str, value: Any) -> Dict:
    """Met à jour un champ spécifique du profil utilisateur"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Vérifier que le champ est autorisé
        allowed_fields = ["preferred_name", "full_name", "birthday", "children", "projects", "current_goals"]
        if field not in allowed_fields:
            return {"success": False, "error": f"Champ '{field}' non autorisé. Champs autorisés: {allowed_fields}"}
        
        # Mettre à jour
        result = supabase.table("user_profile").update({field: value}).eq("user_id", user_id).execute()
        
        if hasattr(result, 'error') and result.error:
            return {"success": False, "error": str(result.error)}
        
        return {"success": True, "message": f"Profil mis à jour: {field} = {value}"}
    
    except Exception as e:
        logger.error(f"Erreur update_profile: {e}")
        return {"success": False, "error": str(e)}


async def add_child_to_profile(user_id: str, name: str, nickname: str = "", birthday: str = None, notes: str = "") -> Dict:
    """Ajoute un enfant au profil utilisateur"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Récupérer le profil actuel
        result = supabase.table("user_profile").select("children").eq("user_id", user_id).execute()
        
        if not result.data:
            children = []
        else:
            children = result.data[0].get("children", [])
        
        # Ajouter le nouvel enfant
        new_child = {
            "name": name,
            "nickname": nickname,
            "birthday": birthday,
            "notes": notes
        }
        children.append(new_child)
        
        # Mettre à jour
        update_result = supabase.table("user_profile").update({"children": children}).eq("user_id", user_id).execute()
        
        if hasattr(update_result, 'error') and update_result.error:
            return {"success": False, "error": str(update_result.error)}
        
        return {"success": True, "message": f"👶 Enfant '{name}' ajouté", "children": children}
    
    except Exception as e:
        logger.error(f"Erreur add_child: {e}")
        return {"success": False, "error": str(e)}



async def add_family_event(user_id: str, title: str, child_name: str = None, 
                           category: str = "routine", date: str = None, 
                           priority: str = "normal", notes: str = "") -> Dict:
    """Ajoute un événement familial"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Catégories valides
        valid_categories = ["school", "health", "activity", "travel", "document", "routine", "supplies"]
        if category not in valid_categories:
            category = "routine"
        
        # Priorités valides
        valid_priorities = ["critical", "high", "normal", "low"]
        if priority not in valid_priorities:
            priority = "normal"
        
        event_data = {
            "title": title,
            "child_name": child_name,
            "category": category,
            "priority": priority,
            "status": "pending",
            "date": date,
            "notes": notes,
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }
        
        result = supabase.table("family_events").insert(event_data).execute()
        
        if hasattr(result, 'error') and result.error:
            return {"success": False, "error": str(result.error)}
        
        return {"success": True, "message": f"📅 Événement familial ajouté: {title}", "data": result.data[0] if result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur add_family_event: {e}")
        return {"success": False, "error": str(e)}


async def add_win(user_id: str, title: str, category: str = "personal", celebration_emoji: str = "🎉", notes: str = "") -> Dict:
    """Ajoute une victoire"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        valid_categories = ["business", "family", "personal", "money", "health", "farm", "other"]
        if category not in valid_categories:
            category = "personal"
        
        win_data = {
            "title": title,
            "category": category,
            "celebration_emoji": celebration_emoji,
            "notes": notes,
            "date": datetime.now().date().isoformat(),
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }
        
        result = supabase.table("wins").insert(win_data).execute()
        
        if hasattr(result, 'error') and result.error:
            return {"success": False, "error": str(result.error)}
        
        return {"success": True, "message": f"🏆 Victoire ajoutée: {title}", "data": result.data[0] if result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur add_win: {e}")
        return {"success": False, "error": str(e)}


async def add_spending(user_id: str, title: str, amount: float, category: str = "other", 
                        project: str = None, date: str = None, notes: str = "") -> Dict:
    """Ajoute une dépense"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        valid_categories = ["materials", "construction", "labor", "livestock", "crops", 
                           "transport", "equipment", "food", "other"]
        if category not in valid_categories:
            category = "other"
        
        if not date:
            date = datetime.now().date().isoformat()
        
        spending_data = {
            "title": title,
            "amount": float(amount),
            "category": category,
            "project": project or "Général",
            "date": date,
            "notes": notes,
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }
        
        result = supabase.table("spending").insert(spending_data).execute()
        
        if hasattr(result, 'error') and result.error:
            return {"success": False, "error": str(result.error)}
        
        return {"success": True, "message": f"💰 Dépense ajoutée: {amount} CFA - {title}", "data": result.data[0] if result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur add_spending: {e}")
        return {"success": False, "error": str(e)}


async def add_revenue(user_id: str, source: str, amount: float, project: str = None, 
                       date: str = None, notes: str = "") -> Dict:
    """Ajoute un revenu"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        if not date:
            date = datetime.now().date().isoformat()
        
        revenue_data = {
            "source": source,
            "amount": float(amount),
            "project": project or "Général",
            "date": date,
            "notes": notes,
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }
        
        result = supabase.table("revenue").insert(revenue_data).execute()
        
        if hasattr(result, 'error') and result.error:
            return {"success": False, "error": str(result.error)}
        
        return {"success": True, "message": f"💰 Revenu ajouté: {amount} CFA - {source}", "data": result.data[0] if result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur add_revenue: {e}")
        return {"success": False, "error": str(e)}


async def add_mission(user_id: str, name: str, category: str = "business", 
                       priority: str = "normal", status: str = "idea", 
                       deadline: str = None, revenue_potential: int = 3,
                       notes: str = "") -> Dict:
    """Ajoute une mission"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        valid_categories = ["business", "family", "personal", "relocation", "farm", "content", "documents"]
        if category not in valid_categories:
            category = "business"
        
        valid_priorities = ["critical", "high", "normal", "low"]
        if priority not in valid_priorities:
            priority = "normal"
        
        valid_statuses = ["idea", "planning", "active", "waiting", "paused", "complete"]
        if status not in valid_statuses:
            status = "idea"
        
        mission_data = {
            "name": name,
            "category": category,
            "priority": priority,
            "status": status,
            "deadline": deadline,
            "revenue_potential": revenue_potential,
            "notes": notes,
            "user_id": user_id,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }
        
        result = supabase.table("missions").insert(mission_data).execute()
        
        if hasattr(result, 'error') and result.error:
            return {"success": False, "error": str(result.error)}
        
        return {"success": True, "message": f"🎯 Mission ajoutée: {name}", "data": result.data[0] if result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur add_mission: {e}")
        return {"success": False, "error": str(e)}


async def add_document(user_id: str, name: str, doc_type: str = "other", 
                        status: str = "draft", due_date: str = None, 
                        url: str = None, notes: str = "") -> Dict:
    """Ajoute un document"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        valid_types = ["proposal", "contract", "grant", "invoice", "legal", "admin", "other"]
        if doc_type not in valid_types:
            doc_type = "other"
        
        valid_statuses = ["draft", "review", "ready", "submitted", "approved", "rejected"]
        if status not in valid_statuses:
            status = "draft"
        
        document_data = {
            "name": name,
            "type": doc_type,
            "status": status,
            "due_date": due_date,
            "url": url,
            "notes": notes,
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }
        
        result = supabase.table("documents").insert(document_data).execute()
        
        if hasattr(result, 'error') and result.error:
            return {"success": False, "error": str(result.error)}
        
        return {"success": True, "message": f"📄 Document ajouté: {name}", "data": result.data[0] if result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur add_document: {e}")
        return {"success": False, "error": str(e)}


async def list_documents(user_id: str, limit: int = 10, status: str = None, show_details: bool = True) -> Dict:
    """Liste les documents de l'utilisateur avec tous les détails"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré", "documents": []}
    
    try:
        query = supabase.table("documents").select("*").eq("user_id", user_id).order("created_at", desc=True).limit(limit)
        
        if status:
            query = query.eq("status", status)
        
        result = query.execute()
        
        documents = []
        for doc in result.data:
            doc_info = {
                "id": doc.get("id"),
                "name": doc.get("name"),
                "type": doc.get("type"),
                "status": doc.get("status"),
                "due_date": doc.get("due_date"),
                "created_at": doc.get("created_at"),
                "url": doc.get("url"),
                "file_url": doc.get("file_url"),
                "file_name": doc.get("file_name"),
                "notes": doc.get("notes")
            }
            documents.append(doc_info)
        
        return {"success": True, "documents": documents, "count": len(documents)}
    
    except Exception as e:
        logger.error(f"Erreur list_documents: {e}")
        return {"success": False, "error": str(e), "documents": []}


async def update_document(user_id: str, document_id: str = None, name: str = None, 
                          updates: dict = None) -> Dict:
    """Met à jour un document existant"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Trouver le document par ID ou par nom
        if document_id:
            result = supabase.table("documents").select("*").eq("id", document_id).eq("user_id", user_id).execute()
        elif name:
            result = supabase.table("documents").select("*").eq("user_id", user_id).ilike("name", f"%{name}%").execute()
        else:
            return {"success": False, "error": "Il faut fournir un ID ou un nom"}
        
        if not result.data:
            return {"success": False, "error": f"Document non trouvé: {name or document_id}"}
        
        doc = result.data[0]
        
        # Appliquer les mises à jour
        update_data = {}
        allowed_fields = ["name", "type", "status", "due_date", "url", "file_url", "notes"]
        for field in allowed_fields:
            if field in updates:
                update_data[field] = updates[field]
        
        if not update_data:
            return {"success": False, "error": "Aucune mise à jour fournie"}
        
        update_data["updated_at"] = datetime.now().isoformat()
        
        # Mettre à jour
        update_result = supabase.table("documents").update(update_data).eq("id", doc["id"]).execute()
        
        return {"success": True, "message": f"Document '{doc['name']}' mis à jour", "document": update_result.data[0] if update_result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur update_document: {e}")
        return {"success": False, "error": str(e)}

async def update_item(user_id: str, table: str, name: str = None, item_id: str = None, updates: dict = None):
    """Met à jour un élément existant dans n'importe quelle table"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    allowed_tables = ["tasks", "missions", "family_events", "documents", "wins", "spending", "revenue"]
    if table not in allowed_tables:
        return {"success": False, "error": f"Table '{table}' non autorisée"}
    
    try:
        # Trouver l'élément
        if item_id:
            result = supabase.table(table).select("*").eq("id", item_id).eq("user_id", user_id).execute()
        elif name:
            # Recherche selon la table
            if table == "tasks":
                result = supabase.table(table).select("*").eq("user_id", user_id).ilike("title", f"%{name}%").execute()
            elif table == "missions":
                result = supabase.table(table).select("*").eq("user_id", user_id).ilike("name", f"%{name}%").execute()
            elif table == "family_events":
                result = supabase.table(table).select("*").eq("user_id", user_id).ilike("title", f"%{name}%").execute()
            else:
                result = supabase.table(table).select("*").eq("user_id", user_id).execute()
        else:
            return {"success": False, "error": "Il faut fournir un nom ou un ID"}
        
        if not result.data:
            return {"success": False, "error": f"Élément non trouvé: {name or item_id}"}
        
        item = result.data[0]
        
        # Appliquer les mises à jour
        update_data = {}
        allowed_fields = {
            "tasks": ["title", "status", "priority", "due_date", "project", "notes"],
            "missions": ["name", "status", "priority", "deadline", "notes", "revenue_potential"],
            "family_events": ["title", "child_name", "category", "date", "priority", "status", "notes"],
            "documents": ["name", "type", "status", "due_date", "url", "file_url", "notes"],
            "wins": ["title", "category", "celebration_emoji", "notes"],
            "spending": ["title", "amount", "category", "project", "notes"],
            "revenue": ["source", "amount", "project", "notes"]
        }
        
        for field, value in updates.items():
            if field in allowed_fields.get(table, []):
                update_data[field] = value
        
        if not update_data:
            return {"success": False, "error": "Aucune mise à jour valide"}
        
        update_data["updated_at"] = datetime.now().isoformat()
        
        # Mettre à jour
        update_result = supabase.table(table).update(update_data).eq("id", item["id"]).execute()
        
        return {"success": True, "message": f"Élément '{name or item_id}' mis à jour dans {table}", "data": update_result.data[0] if update_result.data else None}
    
    except Exception as e:
        logger.error(f"Erreur update_item: {e}")
        return {"success": False, "error": str(e)}


def get_allowed_fields_for_table(table: str) -> list:
    """Retourne les champs autorisés pour la mise à jour selon la table"""
    fields_map = {
        "tasks": ["title", "status", "priority", "due_date", "project", "notes"],
        "missions": ["name", "status", "priority", "deadline", "notes", "revenue_potential"],
        "family_events": ["title", "child_name", "category", "date", "priority", "status", "notes"],
        "documents": ["name", "type", "status", "due_date", "url", "file_url", "notes"],
        "wins": ["title", "category", "celebration_emoji", "notes"],
        "spending": ["title", "amount", "category", "project", "notes"],
        "revenue": ["source", "amount", "project", "notes"]
    }
    return fields_map.get(table, ["notes"])
        

# =====================================================
# API ROUTES - CHAT AMÉLIORÉ AVEC MÉMOIRE
# =====================================================
@app.post("/chat")
async def chat_endpoint(request: ChatRequest):
    logger.info(f"📨 Reçu: {len(request.messages)} messages")
    
    # Récupérer l'utilisateur
    user_id = require_user_id(request.user_id)
    
    # Récupérer le dernier message
    last_message = request.messages[-1].get("content", "") if request.messages else ""
    
    # =====================================================
    # INTERCEPTION DES DEMANDES D'EMAIL (avant l'IA)
    # =====================================================
    last_message_lower = last_message.lower()
    
    if "envoie un email" in last_message_lower or "envoyer un email" in last_message_lower:
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', last_message)
        subject_match = re.search(r'sujet[\s:]+["\']?([^"\'\n]+)', last_message)
        body_match = re.search(r'corps[\s:]+["\']?([^"\'\n]+)', last_message)
        
        if email_match:
            to = email_match.group(0)
            subject = subject_match.group(1) if subject_match else "Message de Sovereign"
            body = body_match.group(1) if body_match else last_message
            
            # Nettoyer le corps
            body = re.sub(r'(envoie un email|envoyer un email).*?corps[\s:]+', '', body, flags=re.IGNORECASE)
            
            pending_emails[user_id] = {"to": to, "subject": subject, "body": body}
            
            preview = f"""
📧 **Aperçu de l'email :**

**Destinataire:** {to}
**Sujet:** {subject}

**Contenu:**
{body[:500]}

---
✏️ **Veux-tu que j'envoie cet email ?**
Réponds par 'oui' pour envoyer, 'non' pour annuler.
"""
            return {"reply": preview}
    
    # =====================================================
    # VÉRIFICATION DES EMAILS EN ATTENTE
    # =====================================================
    logger.info(f"🔍 pending_emails: {pending_emails}")
    logger.info(f"🔍 user_id: {user_id}")
    
    if user_id and user_id in pending_emails:
        logger.info(f"🔍 Email trouvé en attente")
        last_response = request.messages[-1].get("content", "").lower()
        logger.info(f"🔍 Dernier message: '{last_response}'")
        pending = pending_emails[user_id]
        
        if last_response in ["oui", "yes", "ok", "envoie", "envoyer", "faut envoyer", "je confirme", "vas y", "envoie le mail"]:
            logger.info(f"🔍 Confirmation reçue, envoi en cours...")
            email_result = await send_email_simple(
                to=pending["to"],
                subject=pending["subject"],
                body=pending["body"]
            )
            logger.info(f"🔍 Résultat: {email_result}")
            
            if email_result.get("success"):
                reply = f"✅ Email envoyé à {pending['to']}"
            else:
                reply = f"❌ Erreur: {email_result.get('error')}"
            del pending_emails[user_id]
            return {"reply": reply}
        
        elif last_response in ["non", "no", "annule", "annuler"]:
            reply = "❌ Envoi annulé."
            del pending_emails[user_id]
            return {"reply": reply}
    
    # =====================================================
    # CONSTRUCTION DES MESSAGES POUR L'IA
    # =====================================================
    messages_payload = []
    
    # Date du jour
    today_date = datetime.now().strftime("%B %d, %Y")
    date_context = f"\n\nToday is {today_date}. Use this information to provide relevant context."
    
    # Contexte mémoire
    memory_context = await get_user_memory_context(user_id)
    profile_context_result = await get_profile_context(user_id=user_id)
    profile_context = profile_context_result.get("context", "")
    
    enhanced_system_prompt = BASE_SYSTEM_PROMPT + date_context + memory_context
    if profile_context:
        enhanced_system_prompt += f"\n\n# X. CURRENT PROFILE CONTEXT\n{profile_context}"
    
    # Instructions WhatsApp
    whatsapp_instructions = """
    
# WHATSAPP SPECIFIC RULES:
- To get pending WhatsApp messages, use: read_table(table="whatsapp_messages", filters={"replied": False})
- NEVER use status="unread" because this column doesn't exist
- Use "replied": False for pending messages
- Use "replied": True for answered messages
"""
    enhanced_system_prompt += whatsapp_instructions
    
    messages_payload.append({"role": "system", "content": enhanced_system_prompt})
    
    # Traitement des fichiers
    all_file_urls = []
    
    for msg in request.messages:
        role = msg.get("role", "user")
        content = msg.get("content", "")
        
        extract_result = extract_text_from_message(content)
        text_content = extract_result[0]
        image_urls = extract_result[1] if len(extract_result) > 1 else []
        file_urls = extract_result[2] if len(extract_result) > 2 else []
        
        all_file_urls.extend(file_urls)
        
        if role == "user" and image_urls:
            vision_content = [{"type": "text", "text": text_content}]
            for img_url in image_urls:
                base64_image = await download_image_from_url(img_url)
                if base64_image:
                    vision_content.append({
                        "type": "image_url",
                        "image_url": {"url": base64_image, "detail": "high"}
                    })
                else:
                    vision_content.append({
                        "type": "text",
                        "text": f"[Image non accessible: {img_url}]"
                    })
            messages_payload.append({"role": role, "content": vision_content})
        else:
            messages_payload.append({"role": role, "content": text_content})
    
    # Extraction des documents
    logger.info(f"📨 Dernier message: {last_message[:200]}...")
    document_text = None
    
    if all_file_urls:
        for doc_url in all_file_urls:
            if doc_url.lower().endswith(('.pdf', '.txt', '.docx', '.doc')):
                doc_text = await process_document(doc_url)
                if doc_text:
                    document_text = doc_text
                    break
    
    if document_text:
        messages_payload.append({
            "role": "user",
            "content": f"[CONTENU DU DOCUMENT EXTRAIT]\n{document_text}\n\nQuestion ou demande associée : {last_message[:500]}"
        })
    
    # =====================================================
    # APPEL À L'IA
    # =====================================================
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages_payload,
            tools=tools,
            tool_choice="auto",
            max_tokens=4096
        )
        
        msg = response.choices[0].message
        messages_payload.append(msg)
        
        if not msg.tool_calls:
            return {"reply": msg.content}
        
        for tool_call in msg.tool_calls:
            name = tool_call.function.name
            args = json.loads(tool_call.function.arguments)
            content = ""
            
            if name == "read_table":
                table = args.get("table")
                filters = args.get("filters", {})
                
                if table == "whatsapp_messages":
                    if filters.get("status") == "unread":
                        filters.pop("status", None)
                        filters["replied"] = False
                    if filters.get("status") == "pending":
                        filters.pop("status", None)
                        filters["replied"] = False
                
                result = db_query(table, filters, args.get("limit", 50))
                content = json.dumps(result, ensure_ascii=False)
                logger.info(f"📖 Lecture {table}: {result.get('count', 0)} lignes")
            
            elif name == "send_email":
                to = args.get("to", "")
                subject = args.get("subject", "")
                body = args.get("body", "")
                
                pending_emails[user_id] = {"to": to, "subject": subject, "body": body}
                
                preview = f"""
📧 **Aperçu de l'email :**

**Destinataire:** {to}
**Sujet:** {subject}

**Contenu:**
{body[:500]}{'...' if len(body) > 500 else ''}

---
✏️ **Veux-tu que j'envoie cet email ?**
Réponds par 'oui' pour envoyer, 'non' pour annuler.
"""
                content = preview
                logger.info(f"📧 Email en attente pour {to}")
            
            elif name == "update_item":
                table = args.get("table")
                item_name = args.get("name")
                updates = args.get("updates", {})
                
                # Version 1 : utiliser tous les paramètres nommés
                result = await update_item(user_id=user_id, table=table, name=item_name, updates=updates)
                
                # OU Version 2 : utiliser les paramètres positionnels correctement
                # result = await update_item(user_id, table, item_name, None, updates)
                
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                    for field, value in updates.items():
                        content += f"\n   • {field}: {value}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"📝 Update {table}: {item_name}")
            
            elif name == "list_documents":
                limit = args.get("limit", 10)
                status = args.get("status")
                show_details = args.get("show_details", True)
                
                result = await list_documents(user_id, limit, status)
                
                if result.get("success") and result.get("documents"):
                    docs = result["documents"]
                    doc_list = []
                    for i, d in enumerate(docs[:10], 1):
                        doc_info = f"{i}. **{d['name']}**\n"
                        doc_info += f"   📂 Type: {d['type']} | 📌 Statut: {d['status']}\n"
                        if d.get('due_date'):
                            doc_info += f"   📅 Échéance: {d['due_date']}\n"
                        file_link = d.get('file_url') or d.get('url')
                        if file_link:
                            doc_info += f"   📎 [Fichier]({file_link})\n"
                        if d.get('notes'):
                            doc_info += f"   📝 Notes: {d['notes'][:100]}\n"
                        doc_list.append(doc_info)
                    
                    content = f"📄 **Mes documents ({len(docs)}):**\n\n" + "\n".join(doc_list)
                    if len(docs) > 10:
                        content += f"\n\n... et {len(docs) - 10} autre(s) document(s)"
                else:
                    content = "📄 Aucun document trouvé"
            
            elif name == "add_mission":
                name = args.get("name")
                category = args.get("category", "business")
                priority = args.get("priority", "normal")
                
                result = await add_mission(user_id, name, category, priority)
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"🎯 Add mission: {name}")

            elif name == "delete_item":
                table = args.get("table")
                item_name = args.get("name")
                item_id = args.get("item_id")
                
                logger.info(f"🗑️ Delete {table}: {item_name}")
                
                # Chercher l'élément
                if item_id:
                    result = db_query(table, {"id": item_id, "user_id": user_id}, limit=1)
                else:
                    # Récupérer tous les éléments
                    all_items = db_query(table, {"user_id": user_id}, limit=100)
                    found_item = None
                    
                    if all_items.get("data"):
                        item_name_lower = item_name.lower()
                        for item in all_items["data"]:
                            title = item.get("title") or item.get("name") or ""
                            if item_name_lower in title.lower():
                                found_item = item
                                break
                
                if found_item or (item_id and result.get("data")):
                    item = found_item or result["data"][0]
                    item_title = item.get("title") or item.get("name") or item_name
                    
                    # Supprimer
                    delete_result = db_delete(table, item["id"])
                    
                    if delete_result.get("success"):
                        content = f"✅ {table[:-1].capitalize()} supprimé(e) : **{item_title}**"
                        logger.info(f"✅ Deleted {table}: {item_title}")
                    else:
                        content = f"❌ Erreur lors de la suppression"
                else:
                    content = f"❌ Élément non trouvé : '{item_name}' dans {table}"

            elif name == "update_mission_status":
                mission_name = args.get("mission_name")
                status = args.get("status")
                
                logger.info(f"🎯 Update mission status: {mission_name} to {status}")
                
                all_missions = db_query("missions", {"user_id": user_id}, limit=100)
                found_mission = None
                
                if all_missions.get("data"):
                    mission_name_lower = mission_name.lower()
                    for m in all_missions["data"]:
                        name_lower = m.get("name", "").lower()
                        if mission_name_lower in name_lower or name_lower in mission_name_lower:
                            found_mission = m
                            break
                
                if found_mission:
                    update_result = db_update("missions", found_mission["id"], {"status": status})
                    if update_result.get("success"):
                        content = f"✅ Statut de la mission '{found_mission['name']}' changé en **{status}**"
                    else:
                        content = f"❌ Erreur lors de la mise à jour"
                else:
                    content = f"❌ Mission non trouvée : '{mission_name}'"

            elif name == "update_mission_priority":
                mission_name = args.get("mission_name")
                priority = args.get("priority")
                
                logger.info(f"🎯 Update mission priority: {mission_name} to {priority}")
                
                # Récupérer toutes les missions
                all_missions = db_query("missions", {"user_id": user_id}, limit=100)
                found_mission = None
                
                if all_missions.get("data"):
                    mission_name_lower = mission_name.lower()
                    for m in all_missions["data"]:
                        name_lower = m.get("name", "").lower()
                        if mission_name_lower in name_lower or name_lower in mission_name_lower:
                            found_mission = m
                            break
                
                if found_mission:
                    update_result = db_update("missions", found_mission["id"], {"priority": priority})
                    if update_result.get("success"):
                        content = f"✅ Priorité de la mission '{found_mission['name']}' changée en **{priority}**"
                        logger.info(f"✅ Mission priority updated: {found_mission['name']} -> {priority}")
                    else:
                        content = f"❌ Erreur lors de la mise à jour de la priorité"
                else:
                    content = f"❌ Mission non trouvée : '{mission_name}'"
            
            elif name == "add_spending":
                title = args.get("title")
                amount = args.get("amount")
                category = args.get("category", "other")
                
                result = await add_spending(user_id, title, amount, category)
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"💰 Add spending: {amount} CFA - {title}")

            elif name == "update_task_priority":
                task_name = args.get("task_name")
                priority = args.get("priority")
                
                logger.info(f"📋 Update priority for task: {task_name} to {priority}")
                
                # Chercher la tâche
                result = db_query("tasks", {"user_id": user_id, "title": task_name}, limit=1)
                if not result.get("data"):
                    # Recherche approximative
                    all_tasks = db_query("tasks", {"user_id": user_id}, limit=100)
                    if all_tasks.get("data"):
                        matching = [t for t in all_tasks["data"] if task_name.lower() in t.get("title", "").lower()]
                        if matching:
                            result = {"success": True, "data": matching[:1]}
                
                if result.get("success") and result.get("data"):
                    task = result["data"][0]
                    task_id = task["id"]
                    task_title = task["title"]
                    
                    # Mettre à jour la priorité
                    update_result = db_update("tasks", task_id, {"priority": priority})
                    
                    if update_result.get("success"):
                        content = f"✅ Priorité de la tâche '{task_title}' changée en **{priority}**"
                        logger.info(f"✅ Task priority updated: {task_title} -> {priority}")
                    else:
                        content = f"❌ Erreur lors de la mise à jour de la priorité"
                else:
                    content = f"❌ Tâche non trouvée : '{task_name}'"
                    
            elif name == "add_revenue":
                source = args.get("source")
                amount = args.get("amount")
                
                result = await add_revenue(user_id, source, amount)
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"💰 Add revenue: {amount} CFA - {source}")
            
            elif name == "add_win":
                title = args.get("title")
                
                result = await add_win(user_id, title)
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"🏆 Add win: {title}")
            
            elif name == "add_family_event":
                title = args.get("title")
                child_name = args.get("child_name")
                
                result = await add_family_event(user_id, title, child_name)
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"📅 Add family event: {title}")
            
            elif name == "add_child":
                name = args.get("name")
                
                result = await add_child_to_profile(user_id, name)
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"👶 Add child: {name}")
            
            elif name == "update_profile":
                field = args.get("field")
                value = args.get("value")
                
                result = await update_user_profile_field(user_id, field, value)
                if result.get("success"):
                    content = f"✅ {result.get('message')}"
                else:
                    content = f"❌ {result.get('error')}"
                logger.info(f"👤 Update profile: {field}")

            elif name == "complete_task":
                task_name = args.get("task_name")
                
                logger.info(f"📋 Complete task: {task_name}")
                
                # Récupérer toutes les tâches
                all_tasks = db_query("tasks", {"user_id": user_id}, limit=100)
                found_task = None
                
                if all_tasks.get("data"):
                    task_name_lower = task_name.lower()
                    for t in all_tasks["data"]:
                        title_lower = t.get("title", "").lower()
                        # Vérifier si le nom demandé est contenu dans le titre
                        if task_name_lower in title_lower or title_lower in task_name_lower:
                            found_task = t
                            break
                
                if found_task:
                    update_result = db_update("tasks", found_task["id"], {"status": "done"})
                    if update_result.get("success"):
                        content = f"✅ Tâche marquée comme terminée : **{found_task['title']}**"
                    else:
                        content = f"❌ Erreur lors de la mise à jour"
                else:
                    content = f"❌ Tâche non trouvée : '{task_name}'"
            
            elif name == "whatsapp_send_reply":
                to = args.get("to", "")
                message = args.get("message", "")
                
                if not to.endswith("@c.us"):
                    to = to + "@c.us"
                
                result = await whatsapp_send_message(to, message)
                if result:
                    content = f"✅ Message WhatsApp envoyé à {to}"
                else:
                    content = f"❌ Échec de l'envoi"
                logger.info(f"📱 WhatsApp reply to {to}")
            
            elif name == "create_task":
                title = args.get("title")
                priority = args.get("priority", "normal")
                
                result = await create_task_from_conversation(ExecuteTaskRequest(
                    title=title,
                    priority=priority,
                    user_id=user_id
                ))
                if result.get("success"):
                    content = f"✅ Tâche créée: {title}"
                else:
                    content = f"❌ Erreur création tâche"
                logger.info(f"📋 Create task: {title}")
            
            elif name == "get_financial_summary":
                result = get_financial_summary(user_id)
                content = json.dumps(result, ensure_ascii=False)
            
            messages_payload.append({
                "role": "tool",
                "tool_call_id": tool_call.id,
                "content": content
            })
        
        final_response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages_payload,
            max_tokens=4096
        )
        
        assistant_response = final_response.choices[0].message.content
        
        # Nettoyer les tags d'apprentissage
        learn_pattern = r'\[LEARN:([^:]+):([^:]+):([^\]]+)\]'
        matches = re.findall(learn_pattern, assistant_response)
        for match in matches:
            category, original, correction = match
            logger.info(f"📚 Apprentissage: {category} - '{original}' -> '{correction}'")
        
        clean_response = re.sub(learn_pattern, '', assistant_response).strip()
        
        if request.messages:
            last_user = request.messages[-1].get("content", "")
            tools_used = [tc.function.name for tc in msg.tool_calls] if msg.tool_calls else []
            store_chat_session(last_user, clean_response, tools_used, user_id=user_id)
        
        logger.info(f"📨 Réponse envoyée")
        return {"reply": clean_response}
    
    except Exception as e:
        logger.error(f"❌ Erreur chat: {e}")
        reply = "❌ Je rencontre un problème technique. Peux-tu reformuler ta demande ?"
        return {"reply": reply, "error_type": type(e).__name__, "error_message": str(e)[:200]}
# =====================================================
# API ROUTES - SPECIALIZED (EXISTANT)
# =====================================================

@app.get("/financials/summary")
def financial_summary(user_id: Optional[str] = None):
    user_id = require_user_id(user_id)
    return get_financial_summary(user_id)

@app.get("/tasks/priority")
def tasks_priority(limit: int = 10, user_id: Optional[str] = None):
    user_id = require_user_id(user_id)
    return {"tasks": get_priority_tasks(limit, user_id=user_id)}





# =====================================================
# TRANSCRIPTION AUDIO (EXISTANT)
# =====================================================

@app.post("/api/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    """Transcrit un fichier audio en texte"""
    if not file.filename.endswith(('.mp3', '.wav', '.m4a', '.ogg', '.webm')):
        return {"success": False, "error": "Format audio non supporté"}
    
    try:
        audio_data = await file.read()
        
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
            tmp.write(audio_data)
            tmp_path = tmp.name
        
        with open(tmp_path, "rb") as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                language="fr"
            )
        
        os.unlink(tmp_path)
        
        return {"success": True, "text": transcript.text}
    except Exception as e:
        logger.error(f"Erreur transcription: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# EXTRACTION DE TEXTE DEPUIS DOCUMENTS (EXISTANT)
# =====================================================

@app.post("/api/extract-text")
async def extract_text_from_document(file: UploadFile = File(...)):
    """Extrait le texte d'un document (PDF, DOCX, TXT)"""
    try:
        content = await file.read()
        text = ""
        
        if file.filename.endswith('.pdf'):
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text += page.extract_text()
        elif file.filename.endswith('.txt'):
            text = content.decode('utf-8')
        elif file.filename.endswith(('.docx', '.doc')):
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            return {"success": False, "error": "Format non supporté"}
        
        return {"success": True, "text": text[:5000]}
    except Exception as e:
        logger.error(f"Erreur extraction: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# MÉMOIRE IA - APPRENTISSAGE (EXISTANT)
# =====================================================

def save_to_memory(
    key: str,
    value: Dict,
    context: str = None,
    user_id: str = DEFAULT_USER_ID
):
    if not supabase:
        return
    
    try:
        user_id = require_user_id(user_id)

        existing = (
            supabase.table("ai_memory")
            .select("*")
            .eq("key", key)
            .eq("user_id", user_id)
            .execute()
        )

        if existing.data:
            supabase.table("ai_memory").update({
                "value": value,
                "context": context,
                "updated_at": datetime.now().isoformat()
            }).eq("key", key).eq("user_id", user_id).execute()
        else:
            supabase.table("ai_memory").insert({
                "key": key,
                "value": value,
                "context": context,
                "user_id": user_id
            }).execute()

        logger.info(f"💾 Mémoire sauvegardée: {key}")
    except Exception as e:
        logger.error(f"Erreur sauvegarde mémoire: {e}")

def get_from_memory(key: str, user_id: Optional[str] = None) -> Dict:
    if not supabase:
        return {}
    
    try:
        user_id = require_user_id(user_id)

        result = (
            supabase.table("ai_memory")
            .select("*")
            .eq("key", key)
            .eq("user_id", user_id)
            .execute()
        )

        if result.data:
            return result.data[0].get("value", {})
    except Exception as e:
        logger.error(f"Erreur lecture mémoire: {e}")
    return {}


def record_user_correction(original_input: str, correction: str, category: str):
    key = f"correction_{category}"
    existing = get_from_memory(key)
    
    if not existing:
        existing = {"patterns": [], "count": 0}
    
    existing["patterns"].append({
        "original": original_input,
        "corrected": correction,
        "timestamp": datetime.now().isoformat()
    })
    existing["count"] += 1
    
    if len(existing["patterns"]) > 20:
        existing["patterns"] = existing["patterns"][-20:]
    
    save_to_memory(key, existing, f"Corrections utilisateur pour {category}")
    update_smart_mapping(original_input, correction)


def update_smart_mapping(original: str, corrected: str):
    key = "smart_category_mapping"
    mappings = get_from_memory(key)
    
    if not mappings:
        mappings = {}
    
    original_clean = original.lower().strip()
    corrected_clean = corrected.lower().strip()
    
    if original_clean not in mappings:
        mappings[original_clean] = {"corrected_to": corrected_clean, "count": 1}
    else:
        mappings[original_clean]["count"] += 1
    
    save_to_memory(key, mappings, "Mapping intelligent des catégories")


def get_smart_category(input_text: str) -> str:
    input_clean = input_text.lower().strip()
    mappings = get_from_memory("smart_category_mapping")
    
    if not mappings:
        return None
    
    if input_clean in mappings:
        return mappings[input_clean]["corrected_to"]
    
    for key, value in mappings.items():
        if key in input_clean or input_clean in key:
            return value["corrected_to"]
    
    return None


# =====================================================
# NOUVEAUX RAPPELS
# =====================================================

@app.post("/api/missions-daily-reminder")
async def missions_daily_reminder(request: Dict[str, Any] = None):
    """Rappel quotidien des missions actives (tous les jours à 9h)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    user_id = get_request_user_id(request or {})
    
    try:
        active_missions = (
            supabase.table("missions")
            .select("*")
            .eq("user_id", user_id)
            .eq("status", "active")
            .execute()
        )
        
        if not active_missions.data:
            return {"success": True, "sent": False, "message": "Aucune mission active"}
        
        missions_by_category = {}
        for mission in active_missions.data:
            cat = mission.get("category", "other")
            missions_by_category[cat] = missions_by_category.get(cat, 0) + 1
        
        high_priority = sum(1 for m in active_missions.data if m.get("priority") in ["critical", "high"])
        
        total = len(active_missions.data)
        message = f"{total} mission(s) active(s) en cours"
        
        if high_priority > 0:
            message += f" dont {high_priority} prioritaire(s)"
        
        cat_names = {
            "business": "Business",
            "farm": "Ferme", 
            "family": "Famille",
            "relocation": "Relocalisation",
            "content": "Contenu",
            "documents": "Documents"
        }
        
        main_cats = [(cat, count) for cat, count in missions_by_category.items() if cat in cat_names]
        if main_cats:
            cat_str = ", ".join([f"{cat_names.get(cat, cat)}: {count}" for cat, count in main_cats[:3]])
            message += f" • {cat_str}"
        
        send_notification_sync({
            "title": "🎯 Missions actives du jour",
            "body": message,
            "url": "/missions",
            "tag": "daily_missions",
            "type": "mission",
            "requireInteraction": False,
            "user_id": user_id
        })
        
        return {"success": True, "sent": True, "total_missions": total, "message": message}
    
    except Exception as e:
        logger.error(f"Erreur missions_daily_reminder: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/opportunities-reminder")
async def opportunities_reminder(request: Dict[str, Any] = None):
    """Rappel des opportunités à haut potentiel"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        high_opps = (
            supabase.table("opportunities")
            .select("*")
            .eq("user_id", user_id)
            .eq("probability", "high")
            .not_.in_("stage", ["won", "lost"])
            .execute()
        )        
        if not high_opps.data:
            preparing_opps = (
                supabase.table("opportunities")
                .select("*")
                .eq("user_id", user_id)
                .eq("stage", "preparing")
                .not_.in_("stage", ["won", "lost"])
                .execute()
            )            
            if preparing_opps.data:
                total_value = sum(o.get("estimated_value", 0) for o in preparing_opps.data)
                send_notification_sync({
                    "title": "💼 Opportunités en préparation",
                    "body": f"{len(preparing_opps.data)} opportunité(s) en cours • {total_value:,.0f} CFA de potentiel",
                    "url": "/opportunities",
                    "tag": "opportunities_preparing",
                    "type": "opportunity",
                    "requireInteraction": False,
                    "user_id": user_id
                })
                return {"success": True, "sent": True, "type": "preparing", "count": len(preparing_opps.data)}
            else:
                return {"success": True, "sent": False, "message": "Aucune opportunité à haut potentiel"}
        
        total_value = sum(o.get("estimated_value", 0) for o in high_opps.data)
        send_notification_sync({
            "title": "💰 Opportunités à haut potentiel",
            "body": f"{len(high_opps.data)} opportunité(s) à fort potentiel • {total_value:,.0f} CFA à saisir",
            "url": "/opportunities",
            "tag": "high_value_opportunities",
            "type": "opportunity",
            "requireInteraction": True,
            "user_id": user_id
        })
        
        return {"success": True, "sent": True, "count": len(high_opps.data), "total_value": total_value}
    
    except Exception as e:
        logger.error(f"Erreur opportunities_reminder: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/financial-weekly-report")
async def financial_weekly_report(request: Dict[str, Any] = None):
    """Bilan financier hebdomadaire (le dimanche)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
        user_id = get_request_user_id(request or {})
    
    if datetime.now().weekday() != 6:
        return {"success": True, "sent": False, "message": "Pas le jour du bilan financier (dimanche)"}
    
    try:
        start_of_week = (datetime.now().date() - timedelta(days=7)).isoformat()
        end_of_week = datetime.now().date().isoformat()
        revenue_data = (
            supabase.table("revenue")
            .select("*")
            .eq("user_id", user_id)
            .gte("date", start_of_week)
            .lte("date", end_of_week)
            .execute()
        )
        
        spending_data = (
            supabase.table("spending")
            .select("*")
            .eq("user_id", user_id)
            .gte("date", start_of_week)
            .lte("date", end_of_week)
            .execute()
        )        
        total_revenue = sum(r.get("amount", 0) for r in revenue_data.data)
        
        total_spending = sum(s.get("amount", 0) for s in spending_data.data)
        
        balance = total_revenue - total_spending
        
        spending_by_cat = {}
        for s in spending_data.data:
            cat = s.get("category", "other")
            spending_by_cat[cat] = spending_by_cat.get(cat, 0) + s.get("amount", 0)
        
        top_category = max(spending_by_cat.items(), key=lambda x: x[1]) if spending_by_cat else ("Aucune", 0)
        
        spending_by_project = {}
        for s in spending_data.data:
            project = s.get("project", "Non classé")
            spending_by_project[project] = spending_by_project.get(project, 0) + s.get("amount", 0)
        top_project = max(spending_by_project.items(), key=lambda x: x[1]) if spending_by_project else ("Aucun", 0)
        
        if balance >= 0:
            status = f"✅ Solde positif : +{balance:,.0f} CFA"
        else:
            status = f"⚠️ Solde négatif : {balance:,.0f} CFA"
        
        message = f"📊 Revenus: {total_revenue:,.0f} CFA | Dépenses: {total_spending:,.0f} CFA | {status}"
        
        send_notification_sync({
            "title": "📈 Bilan financier hebdomadaire",
            "body": message,
            "url": "/money",
            "tag": "weekly_financial",
            "type": "financial",
            "requireInteraction": False,
            "user_id": user_id
        })
        
        if total_spending > 500000:
            send_notification_sync({
                "title": "💸 Alerte dépenses",
                "body": f"Dépenses élevées cette semaine ({total_spending:,.0f} CFA). Le poste principal : {top_category[0]}",
                "url": "/money",
                "tag": "high_spending_alert",
                "type": "financial",
                "requireInteraction": True,
                "user_id": user_id
            })
        
        return {
            "success": True, 
            "sent": True, 
            "total_revenue": total_revenue,
            "total_spending": total_spending,
            "balance": balance,
            "top_category": top_category[0],
            "top_project": top_project[0]
        }
    
    except Exception as e:
        logger.error(f"Erreur financial_weekly_report: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/family-events-reminder")
async def family_events_reminder(request: Dict[str, Any] = None):
    """Rappel des événements familiaux (1x par jour max)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}

    user_id = get_request_user_id(request or {})
    
    today = datetime.now().date().isoformat()
    
    # Vérifier si déjà envoyé aujourd'hui
    existing = supabase.table("notifications_log").select("*")\
        .eq("type", "family_reminder")\
        .eq("date", today)\
        .eq("user_id", user_id)\
        .execute()
    
    if existing.data:
        return {"success": True, "sent": False, "message": "Déjà envoyé aujourd'hui"}
    
    notifications_sent = []
    
    try:
        tomorrow = (datetime.now().date() + timedelta(days=1)).isoformat()
        next_3_days = (datetime.now().date() + timedelta(days=3)).isoformat()
        
        today_events = supabase.table("family_events").select("*").eq("user_id", user_id).eq("date", today).neq("status", "done").execute()
        tomorrow_events = (
            supabase.table("family_events")
            .select("*")
            .eq("user_id", user_id)
            .eq("date", tomorrow)
            .neq("status", "done")
            .execute()
        )
        
        # Un seul message, priorité aux événements du jour
        if today_events.data:
            events_summary = ", ".join([f"{e['title']}" + (f" ({e['child_name']})" if e.get('child_name') else "") for e in today_events.data[:3]])
            send_notification_sync({
                "title": "👨‍👩‍👧‍👦 Événement familial AUJOURD'HUI",
                "body": events_summary,
                "url": "/family",
                "type": "family",
                "user_id": user_id
            })
            notifications_sent = [e["id"] for e in today_events.data]
        elif tomorrow_events.data:
            events_summary = ", ".join([f"{e['title']}" + (f" ({e['child_name']})" if e.get('child_name') else "") for e in tomorrow_events.data[:3]])
            send_notification_sync({
                "title": "📅 Rappel familial pour DEMAIN",
                "body": events_summary,
                "url": "/family",
                "type": "family",
                "user_id": user_id
            })
            notifications_sent = [e["id"] for e in tomorrow_events.data]
        
        if notifications_sent:
            supabase.table("notifications_log").insert({
                "type": "family_reminder",
                "date": today,
                "user_id": user_id,
                "sent_at": datetime.now().isoformat()
            }).execute()
        
        return {
            "success": True,
            "sent": len(notifications_sent) > 0,
            "notifications_sent": notifications_sent,
            "count": len(notifications_sent)
        }
    
    except Exception as e:
        logger.error(f"Erreur family_events_reminder: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# API ROUTES - USER PROFILE
# =====================================================




# =====================================================
# API ROUTES - USER PROFILE (VERSION CORRIGÉE)
# =====================================================

@app.get("/api/profile")
async def get_user_profile(user_id: Optional[str] = None):
    user_id = require_user_id(user_id)
    """Récupère le profil utilisateur complet"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        if not result.data:
            # Créer un profil par défaut
            supabase.table("user_profile").insert({
                "user_id": user_id,
                "preferred_name": "Rebecca",
                "children": [],
                "projects": [],
                "current_goals": []
            }).execute()
            result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        profile = result.data[0] if result.data else {}
        
        # Récupérer les champs additionnels depuis user_memory
        memory_result = supabase.table("user_memory").select("*")\
            .eq("user_id", user_id)\
            .eq("category", "profile")\
            .execute()
        
        for mem in memory_result.data:
            value = mem["value"]
            try:
                if isinstance(value, str) and (value.startswith('{') or value.startswith('[')):
                    value = json.loads(value)
            except:
                pass
            profile[mem["key"]] = value
        
        return {"success": True, "profile": profile}
        
    except Exception as e:
        logger.error(f"Erreur get_user_profile: {e}")
        return {"success": False, "error": str(e)}


@app.put("/api/profile")
async def update_user_profile(request: Dict[str, Any]):
    """Met à jour le profil utilisateur"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request)
        logger.info(f"📥 Requête PUT /api/profile reçue: {request}")
        # LOG pour voir ce qui est reçu
        
        # Nettoyer la requête
        request.pop("id", None)
        request.pop("user_id", None)
        
        # Champs autorisés
        allowed_fields = [
            "preferred_name", "full_name", "birthday", 
            "children", "projects", "current_goals",
            "communication_preferences", "upcoming_milestones", "key_contacts"
        ]
        
        # Ne garder que les champs autorisés
        clean_data = {}
        for key, value in request.items():
            if key in allowed_fields:
                clean_data[key] = value
            else:
                logger.warning(f"⚠️ Champ ignoré: {key}")
        
        # Si aucun champ valide, retourner une erreur claire
        if not clean_data:
            return {"success": False, "error": "Aucun champ valide à mettre à jour"}
        
        # Ajouter la date de mise à jour
        clean_data["updated_at"] = datetime.now().isoformat()
        
        logger.info(f"📤 Mise à jour avec: {clean_data}")
        
        # Mettre à jour le profil
        result = supabase.table("user_profile").update(clean_data).eq("user_id", user_id).execute()
        
        # Récupérer le profil mis à jour
        profile_result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        return {"success": True, "profile": profile_result.data[0] if profile_result.data else {}}
        
    except Exception as e:
        logger.error(f"❌ Erreur update_user_profile: {e}")
        return {"success": False, "error": str(e)}
        
@app.post("/api/profile/children")
async def add_child(request: Dict[str, Any]):
    """Ajoute un enfant au profil"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request)
        # Récupérer le profil actuel
        profile_result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        children = profile_result.data[0].get("children", []) if profile_result.data else []
        
        new_child = {
            "name": request.get("name"),
            "nickname": request.get("nickname", ""),
            "birthday": request.get("birthday"),
            "notes": request.get("notes", "")
        }
        children.append(new_child)
        
        result = supabase.table("user_profile").update({
            "children": children,
            "updated_at": datetime.now().isoformat()
        }).eq("user_id", user_id).execute()
        
        return {"success": True, "children": children}
        
    except Exception as e:
        logger.error(f"Erreur add_child: {e}")
        return {"success": False, "error": str(e)}

@app.get("/api/profile/context")
async def get_profile_context(user_id: Optional[str] = None):
    """Récupère un résumé du profil pour injection dans le prompt"""
    if not supabase:
        return {"context": ""}
    
    try:
        user_id = require_user_id(user_id)
        result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        if not result.data:
            return {"context": ""}
        
        profile = result.data[0]
        
        # Construire un contexte texte pour le prompt
        context_parts = []
        
        # Enfants
        children = profile.get("children", [])
        if children:
            child_names = [f"{c.get('name', '')} ({c.get('nickname', '')})" for c in children]
            context_parts.append(f"Rebecca's children: {', '.join(child_names)}")
        
        # Projets
        projects = profile.get("projects", [])
        if projects:
            project_list = [f"{p.get('name', '')} ({p.get('status', 'active')})" for p in projects]
            context_parts.append(f"Her main projects: {', '.join(project_list)}")
        
        # Objectifs actuels
        goals = profile.get("current_goals", [])
        if goals:
            goal_list = [g.get("goal", "") for g in goals[:3]]
            context_parts.append(f"Current priorities: {', '.join(goal_list)}")
        
        # Prochaines étapes
        milestones = profile.get("upcoming_milestones", [])
        if milestones:
            today = datetime.now().date().isoformat()
            upcoming = [m for m in milestones if m.get("date", "9999-12-31") >= today][:3]
            if upcoming:
                milestone_text = ", ".join([f"{m.get('title', '')} ({m.get('date', '')})" for m in upcoming])
                context_parts.append(f"Upcoming: {milestone_text}")
        
        return {"context": " | ".join(context_parts)}
    except Exception as e:
        logger.error(f"Erreur get_profile_context: {e}")
        return {"context": ""}



@app.post("/api/brain-dump/process")
async def process_brain_dump(request: Dict[str, Any]):
    """
    Analyse un texte libre et retourne une structure organisée.
    Version améliorée avec propositions d'actions.
    """
    content = request.get("content", "")
    if not content:
        return {"success": False, "error": "Contenu requis"}
    user_id = get_request_user_id(request or {})
    
    # Estimer la longueur du contenu pour adapter l'analyse
    content_length = len(content)
    content_words = len(content.split())
    
    # Récupérer le contexte utilisateur pour personnaliser
    profile_context = await get_profile_context(user_id=user_id)
    memory_context = await get_user_memory_context(user_id)
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": f"""Tu es Becks, l'assistante personnelle de Rebecca. Tu reçois un texte brut (un brain dump) où Rebecca a écrit tout ce qui lui passe par la tête.

Tu connais Rebecca : {profile_context}
Elle se souvient de : {memory_context}

CONTEXTE : Le texte fait environ {content_words} mots.

RÈGLE IMPORTANTE : Le résumé (summary) doit être PROPORTIONNEL à la richesse du contenu. 
- Si elle a écrit 2-3 phrases simples → résumé court (1-2 phrases)
- Si elle a écrit un paragraphe dense avec plusieurs sujets → résumé de 3-5 phrases
- Si elle a fait un long brain dump (plusieurs paragraphes, multiples préoccupations) → résumé substantiel de 5-8 phrases

Tu dois analyser ce texte et retourner UNIQUEMENT du JSON valide avec cette structure :

{{
  "summary": "un résumé COMPLET et substantiel qui capture TOUS les points importants",
  "emotions": ["émotion1", "émotion2"],
  "main_topics": ["sujet1", "sujet2", "sujet3"],
  "urgency_level": "high/medium/low",
  "priorities": [
    {{"title": "priorité 1", "reason": "pourquoi c'est important"}},
    {{"title": "priorité 2", "reason": "pourquoi c'est important"}}
  ],
  "suggested_tasks": [
    {{"title": "tâche suggérée 1", "project": "projet associé", "priority": "high/medium/low"}},
    {{"title": "tâche suggérée 2", "project": "projet associé", "priority": "high/medium/low"}}
  ],
  "suggested_checklist": {{
    "title": "titre de la checklist",
    "steps": ["étape 1", "étape 2", "étape 3"]
  }},
  "insights": "insight important ou chose à retenir",
  "calming_response": "une réponse réconfortante adaptée à son état (2-3 phrases)",
  "quick_action": "une action simple à faire immédiatement (moins de 5 minutes)"
}}

Projets possibles : Ifè Farm, Love & Fire Sport, Santé Plus, Bénin Relocation, Famille, Personnel
Priorités : high, medium, low
Émotions possibles : stress, fatigue, excitation, frustration, clarté, confusion, sérénité, anxiété, motivation, tristesse, colère, joie

Ne retourne que le JSON, rien d'autre."""
                },
                {"role": "user", "content": content}
            ],
            temperature=0.3,
            max_tokens=1500
        )
        
        result = response.choices[0].message.content
        json_match = re.search(r'\{[\s\S]*\}', result)
        if json_match:
            analysis = json.loads(json_match.group())
        else:
            raise ValueError("Pas de JSON trouvé")
        
        # Sauvegarder l'analyse
        if supabase:
            supabase.table("brain_dump_analyses").insert({
                "content": content[:1000],
                "analysis": analysis,
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }).execute()
        
        # Sauvegarder dans la mémoire émotionnelle
        if analysis.get("emotions"):
            for emotion in analysis["emotions"][:3]:
                await save_user_memory("emotions", "recent_mood", emotion, user_id)
        
        return {"success": True, "analysis": analysis}
        
    except Exception as e:
        logger.error(f"Erreur process_brain_dump: {e}")
        return {"success": False, "error": str(e)}

# =====================================================
# LIFE MAP - VUE D'ENSEMBLE DES DOMAINES
# =====================================================

@app.get("/api/life-map")
async def get_life_map():
    """Récupère les données pour la carte de vie"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    today = datetime.now().date().isoformat()
    
    try:
        # 1. Famille - événements à venir
        family_events = supabase.table("family_events").select("*").gte("date", today).neq("status", "done").execute()
        family_pending = len(family_events.data)
        family_next = family_events.data[0] if family_events.data else None
        
        # 2. Argent - finances
        revenue = supabase.table("revenue").select("amount").execute()
        spending = supabase.table("spending").select("amount").execute()
        total_revenue = sum(r.get("amount", 0) for r in revenue.data)
        total_spending = sum(s.get("amount", 0) for s in spending.data)
        balance = total_revenue - total_spending
        
        # 3. Business - missions actives
        active_missions = supabase.table("missions").select("*").eq("status", "active").execute()
        high_priority_missions = [m for m in active_missions.data if m.get("priority") in ["critical", "high"]]
        
        # 4. Ferme - dépenses et unités actives
        farm_spending = supabase.table("farm_spending").select("amount").execute()
        total_farm_spent = sum(s.get("amount", 0) for s in farm_spending.data)
        active_units = supabase.table("farm_production_units").select("*").eq("status", "active").execute()
        
        # 5. Documents en attente
        pending_docs = supabase.table("documents").select("*").neq("status", "approved").execute()
        urgent_docs = [d for d in pending_docs.data if d.get("due_date") and d["due_date"] < today]
        
        # 6. Victoires récentes (7 jours)
        seven_days_ago = (datetime.now().date() - timedelta(days=7)).isoformat()
        recent_wins = supabase.table("wins").select("*").gte("date", seven_days_ago).execute()
        
        # 7. Relocation - tâches en cours
        relocation_tasks = supabase.table("relocation_tasks").select("*").neq("status", "completed").execute()
        critical_relocation = [t for t in relocation_tasks.data if t.get("priority") == "urgent"]
        
        # 8. Alignement - score basé sur victoires + humeur récente
        recent_mood = supabase.table("mood_entries").select("mood").order("date", desc=True).limit(5).execute()
        positive_moods = sum(1 for m in recent_mood.data if m.get("mood") in ["excellent", "bien"])
        alignment_score = min(100, (len(recent_wins.data) * 10) + (positive_moods * 5))
        
        return {
            "success": True,
            "data": {
                "family": {
                    "status": "🟢" if family_pending == 0 else "🟡",
                    "pending_count": family_pending,
                    "next_action": family_next.get("title", "Aucun événement") if family_next else "Aucun événement",
                    "next_date": family_next.get("date") if family_next else None,
                    "urgency": "low" if family_pending == 0 else "medium"
                },
                "money": {
                    "status": "🔴" if balance < 0 else "🟢",
                    "balance": balance,
                    "pending_invoices": 0,
                    "urgency": "high" if balance < 0 else "low"
                },
                "business": {
                    "status": "🟢" if len(high_priority_missions) == 0 else "🟡",
                    "active_missions": len(active_missions.data),
                    "high_priority_count": len(high_priority_missions),
                    "urgency": "high" if len(high_priority_missions) > 2 else "medium"
                },
                "farm": {
                    "status": "🟢",
                    "total_investment": total_farm_spent,
                    "active_units": len(active_units.data),
                    "next_action": "Vérifier les stocks",
                    "urgency": "medium"
                },
                "documents": {
                    "status": "🔴" if len(urgent_docs) > 0 else "🟡" if len(pending_docs.data) > 0 else "🟢",
                    "pending_count": len(pending_docs.data),
                    "urgent_count": len(urgent_docs),
                    "urgency": "high" if len(urgent_docs) > 0 else "medium"
                },
                "wins": {
                    "status": "🟢",
                    "recent_count": len(recent_wins.data),
                    "streak": len(recent_wins.data),
                    "urgency": "low"
                },
                "relocation": {
                    "status": "🟡",
                    "pending_tasks": len(relocation_tasks.data),
                    "critical_count": len(critical_relocation),
                    "next_deadline": "2026-08-31",
                    "urgency": "high" if len(critical_relocation) > 0 else "medium"
                },
                "alignment": {
                    "status": "🟢" if alignment_score > 70 else "🟡",
                    "score": alignment_score,
                    "recommendation": "Prends un moment pour toi" if alignment_score < 50 else "Continue sur cette lancée",
                    "urgency": "low"
                }
            }
        }
        
    except Exception as e:
        logger.error(f"Erreur life_map: {e}")
        return {"success": False, "error": str(e)}



# =====================================================
# AGENT D'EXÉCUTION - EXECUTE MODE
# =====================================================

@app.post("/api/execute/analyze-request")
async def analyze_execute_request(request: Dict[str, Any]):
    """
    Analyse une demande utilisateur et retourne un plan d'exécution
    """
    query = request.get("query", "")
    if not query:
        return {"success": False, "error": "Query requise"}
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": """Tu es Becks, l'agent d'exécution de Rebecca. Ton rôle est de transformer une demande en plan d'action concret.

Analyse la demande et retourne UNIQUEMENT du JSON valide avec cette structure :

{
  "type": "email|task|checklist|plan|document|meeting",
  "title": "titre de l'action",
  "steps": ["étape 1", "étape 2", "étape 3"],
  "suggested_tasks": [
    {"title": "tâche 1", "priority": "high/medium/low", "due_offset": "1d|2d|3d|1w"},
    {"title": "tâche 2", "priority": "high/medium/low", "due_offset": "1d|2d|3d|1w"}
  ],
  "draft_content": "contenu si c'est un email/document (optionnel)",
  "questions": ["question à poser si info manquante"],
  "next_action": "prochaine action immédiate"
}

Types possibles : email, task, checklist, plan, document, meeting"""
                },
                {"role": "user", "content": query}
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        result = response.choices[0].message.content
        json_match = re.search(r'\{[\s\S]*\}', result)
        if json_match:
            execution_plan = json.loads(json_match.group())
        else:
            raise ValueError("Pas de JSON trouvé")
        
        return {"success": True, "execution_plan": execution_plan}
        
    except Exception as e:
        logger.error(f"Erreur analyze_execute_request: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/execute/create-checklist")
async def create_checklist(request: Dict[str, Any]):
    """Crée une checklist à partir d'un plan"""
    title = request.get("title", "Checklist")
    steps = request.get("steps", [])
    
    if not steps:
        return {"success": False, "error": "Steps requis"}
    user_id = get_request_user_id(request or {})
    
    # Sauvegarder la checklist
    checklist_id = str(uuid.uuid4())
    if supabase:
        supabase.table("checklists").insert({
            "id": checklist_id,
            "title": title,
            "steps": steps,
            "progress": 0,
            "user_id": user_id,
            "created_at": datetime.now().isoformat()
        }).execute()
    
    return {
        "success": True,
        "checklist": {
            "id": checklist_id,
            "title": title,
            "steps": steps,
            "progress": 0
        }
    }


@app.post("/api/execute/create-draft")
async def create_draft(request: Dict[str, Any]):
    """Génère un brouillon (email, document, etc.)"""
    draft_type = request.get("type", "email")
    context = request.get("context", "")
    
    if not context:
        return {"success": False, "error": "Context requis"}
    user_id = get_request_user_id(request or {})
    
    # Définir le prompt selon le type
    prompts = {
        "email": f"Rédige un email professionnel à partir de ce contexte. Style: clair, direct, professionnel. Contexte: {context}",
        "proposal": f"Rédige une proposition courte à partir de ce contexte. Contexte: {context}",
        "letter": f"Rédige une lettre formelle à partir de ce contexte. Contexte: {context}",
        "note": f"Rédige une note professionnelle à partir de ce contexte. Contexte: {context}"
    }
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": prompts.get(draft_type, prompts["email"])},
                {"role": "user", "content": "Génère le contenu final, prêt à être copié."}
            ],
            temperature=0.5,
            max_tokens=1000
        )
        
        draft_content = response.choices[0].message.content
        
        # Sauvegarder le draft
        draft_id = str(uuid.uuid4())
        if supabase:
            supabase.table("drafts").insert({
                "id": draft_id,
                "type": draft_type,
                "content": draft_content,
                "context": context,
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }).execute()
        
        return {
            "success": True,
            "draft": {
                "id": draft_id,
                "type": draft_type,
                "content": draft_content
            }
        }
        
    except Exception as e:
        logger.error(f"Erreur create_draft: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/execute/update-checklist-step")
async def update_checklist_step(request: Dict[str, Any]):
    """Marque une étape comme complétée"""
    checklist_id = request.get("checklist_id")
    step_index = request.get("step_index")
    
    if not checklist_id or step_index is None:
        return {"success": False, "error": "checklist_id et step_index requis"}
    
    try:
        # Récupérer la checklist
        checklist = supabase.table("checklists").select("*").eq("id", checklist_id).execute()
        if not checklist.data:
            return {"success": False, "error": "Checklist non trouvée"}
        
        steps = checklist.data[0].get("steps", [])
        if step_index < len(steps):
            # Marquer comme complété (on pourrait stocker la progression)
            completed = checklist.data[0].get("completed_steps", [])
            if step_index not in completed:
                completed.append(step_index)
            
            progress = int((len(completed) / len(steps)) * 100)
            
            supabase.table("checklists").update({
                "completed_steps": completed,
                "progress": progress,
                "updated_at": datetime.now().isoformat()
            }).eq("id", checklist_id).execute()
            
            return {"success": True, "progress": progress, "completed": len(completed), "total": len(steps)}
        
        return {"success": False, "error": "Step invalide"}
        
    except Exception as e:
        logger.error(f"Erreur update_checklist_step: {e}")
        return {"success": False, "error": str(e)}



# =====================================================
# LOVE & FIRE SPORT MODULE
# =====================================================

@app.get("/api/lf/stats")
async def get_lf_stats():
    """Récupère les statistiques du module Love & Fire Sport"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Grants
        grants = supabase.table("lf_grants").select("*").execute()
        active_grants = [g for g in grants.data if g.get("status") not in ["approved", "rejected"]]
        submitted_grants = [g for g in grants.data if g.get("status") == "submitted"]
        
        # Contrats
        contracts = supabase.table("lf_contracts").select("*").execute()
        dda_contracts = [c for c in contracts.data if c.get("contract_type") == "DDA"]
        
        # Tâches
        tasks = supabase.table("lf_tasks").select("*").neq("status", "done").execute()
        urgent_tasks = [t for t in tasks.data if t.get("priority") == "high" and t.get("status") != "done"]
        
        # Prochaines deadlines
        today = datetime.now().date().isoformat()
        upcoming_grants = [g for g in grants.data if g.get("deadline") and g["deadline"] >= today]
        upcoming_grants.sort(key=lambda x: x.get("deadline", "9999-12-31"))
        
        return {
            "success": True,
            "stats": {
                "total_grants": len(grants.data),
                "active_grants": len(active_grants),
                "submitted_grants": len(submitted_grants),
                "total_contracts": len(contracts.data),
                "dda_contracts": len(dda_contracts),
                "pending_tasks": len(tasks.data),
                "urgent_tasks": len(urgent_tasks),
                "next_deadline": upcoming_grants[0].get("deadline") if upcoming_grants else None,
                "next_deadline_title": upcoming_grants[0].get("title") if upcoming_grants else None
            }
        }
    except Exception as e:
        logger.error(f"Erreur lf_stats: {e}")
        return {"success": False, "error": str(e)}


@app.get("/api/lf/grants")
async def get_lf_grants(status: str = None):
    """Récupère la liste des grants"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        query = supabase.table("lf_grants").select("*").order("deadline", nulls_last=True)
        if status:
            query = query.eq("status", status)
        result = query.execute()
        return {"success": True, "grants": result.data}
    except Exception as e:
        logger.error(f"Erreur lf_grants: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/lf/grant")
async def create_lf_grant(request: Dict[str, Any]):
    """Crée un nouveau grant"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        result = supabase.table("lf_grants").insert({
            "title": request.get("title"),
            "agency": request.get("agency"),
            "amount": request.get("amount"),
            "deadline": request.get("deadline"),
            "status": request.get("status", "researching"),
            "probability": request.get("probability", 50),
            "notes": request.get("notes"),
            "created_at": datetime.now().isoformat()
        }).execute()
        
        # Créer une tâche associée
        if result.data:
            supabase.table("lf_tasks").insert({
                "title": f"Préparer dossier pour {request.get('title')}",
                "related_to": "grant",
                "related_id": result.data[0]["id"],
                "deadline": request.get("deadline"),
                "priority": "high",
                "created_at": datetime.now().isoformat()
            }).execute()
        
        return {"success": True, "grant": result.data[0] if result.data else None}
    except Exception as e:
        logger.error(f"Erreur create_lf_grant: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/lf/contract")
async def create_lf_contract(request: Dict[str, Any]):
    """Crée un nouveau contrat"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        result = supabase.table("lf_contracts").insert({
            "title": request.get("title"),
            "contract_type": request.get("contract_type"),
            "agency": request.get("agency"),
            "status": request.get("status", "draft"),
            "deadline": request.get("deadline"),
            "requirements": request.get("requirements", []),
            "notes": request.get("notes"),
            "created_at": datetime.now().isoformat()
        }).execute()
        
        return {"success": True, "contract": result.data[0] if result.data else None}
    except Exception as e:
        logger.error(f"Erreur create_lf_contract: {e}")
        return {"success": False, "error": str(e)}


@app.get("/api/lf/checklist")
async def get_lf_checklist():
    """Récupère la checklist DDA complète"""
    checklist = [
        {"id": "dda_1", "title": "Créer compte DDA", "status": "pending", "deadline": None},
        {"id": "dda_2", "title": "Compléter formulaire d'identification", "status": "pending", "deadline": None},
        {"id": "dda_3", "title": "Soumettre preuve de statut juridique", "status": "pending", "deadline": None},
        {"id": "dda_4", "title": "Soumettre description des services", "status": "pending", "deadline": None},
        {"id": "dda_5", "title": "Soumettre preuve d'assurance", "status": "pending", "deadline": None},
        {"id": "dda_6", "title": "Attendre approbation", "status": "pending", "deadline": None},
        {"id": "sam_1", "title": "Créer compte SAM.gov", "status": "pending", "deadline": None},
        {"id": "sam_2", "title": "Compléter registration UEI", "status": "pending", "deadline": None},
        {"id": "sam_3", "title": "Soumettre registration", "status": "pending", "deadline": None},
        {"id": "emma_1", "title": "Créer compte eMMA", "status": "pending", "deadline": None},
        {"id": "emma_2", "title": "Soumettre vendor registration Maryland", "status": "pending", "deadline": None},
        {"id": "insurance_1", "title": "Obtenir attestation d'assurance", "status": "pending", "deadline": "2026-06-15"},
        {"id": "budget_1", "title": "Préparer budget pilote", "status": "pending", "deadline": "2026-06-30"},
        {"id": "business_plan_1", "title": "Finaliser business plan", "status": "pending", "deadline": "2026-07-15"},
    ]
    return {"success": True, "checklist": checklist}

@app.get("/api/lf/contracts")
async def get_lf_contracts(status: str = None):
    """Récupère la liste des contrats"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        query = supabase.table("lf_contracts").select("*").order("deadline", nulls_last=True)
        if status:
            query = query.eq("status", status)
        result = query.execute()
        return {"success": True, "contracts": result.data}
    except Exception as e:
        logger.error(f"Erreur lf_contracts: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# EMAIL (BREVO)
# =====================================================

import httpx
from pydantic import BaseModel, EmailStr
from typing import Optional

BREVO_API_KEY = os.environ.get("BREVO_API_KEY")
BREVO_SENDER_EMAIL = os.environ.get("BREVO_SENDER_EMAIL", "sovereign@rebecca.com")
BREVO_SENDER_NAME = os.environ.get("BREVO_SENDER_NAME", "SOVEREIGN - Becks")

class EmailRequest(BaseModel):
    to: EmailStr
    subject: str
    body: str
    to_name: Optional[str] = None

@app.post("/api/email/send")
async def send_email(request: Request):
    """Envoie un email via Brevo avec meilleure validation"""
    try:
        body = await request.json()
        to = body.get("to", "")
        subject = body.get("subject", "")
        body_content = body.get("body", "")
        
        # Validation stricte de l'email
        if not to or "@" not in to or "." not in to:
            return {"success": False, "error": "Adresse email invalide. Veuillez fournir une adresse valide (exemple: nom@domaine.com)"}
        
        if not subject:
            return {"success": False, "error": "Le sujet de l'email est requis"}
        
        if not body_content:
            return {"success": False, "error": "Le corps de l'email est requis"}
        
        if not BREVO_API_KEY:
            return {"success": False, "error": "Service email non configuré"}
        
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.brevo.com/v3/smtp/email",
                headers={
                    "api-key": BREVO_API_KEY,
                    "Content-Type": "application/json",
                    "Accept": "application/json"
                },
                json={
                    "sender": {
                        "name": BREVO_SENDER_NAME,
                        "email": BREVO_SENDER_EMAIL
                    },
                    "to": [{
                        "email": to,
                        "name": to.split("@")[0]
                    }],
                    "subject": subject,
                    "htmlContent": body_content,
                },
                timeout=30.0
            )
            
            if response.status_code == 201:
                logger.info(f"📧 Email envoyé à {to}")
                return {"success": True, "message": "Email envoyé", "to": to}
            else:
                logger.error(f"Erreur Brevo: {response.text}")
                return {"success": False, "error": response.text}
                
    except Exception as e:
        logger.error(f"Erreur envoi email: {e}")
        return {"success": False, "error": str(e)}


async def send_email_simple(to: str, subject: str, body: str) -> Dict:
    """Envoie un email via Brevo (version simplifiée sans Request)"""
    if not BREVO_API_KEY:
        return {"success": False, "error": "Service email non configuré"}
    
    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(
                "https://api.brevo.com/v3/smtp/email",
                headers={
                    "api-key": BREVO_API_KEY,
                    "Content-Type": "application/json",
                    "Accept": "application/json"
                },
                json={
                    "sender": {
                        "name": BREVO_SENDER_NAME,
                        "email": BREVO_SENDER_EMAIL
                    },
                    "to": [{"email": to, "name": to.split("@")[0]}],
                    "subject": subject,
                    "htmlContent": body,
                },
                timeout=30.0
            )
            
            if response.status_code == 201:
                logger.info(f"📧 Email envoyé à {to}")
                return {"success": True, "message": "Email envoyé", "to": to}
            else:
                logger.error(f"Erreur Brevo: {response.text}")
                return {"success": False, "error": response.text}
                
    except Exception as e:
        logger.error(f"Erreur envoi email: {e}")
        return {"success": False, "error": str(e)}

# =====================================================
# WEBHOOKS
# =====================================================



WEBHOOK_SECRET = os.environ.get("WEBHOOK_SECRET", "sovereign-secret-key-2024")

# Stockage simple des webhooks (en mémoire pour commencer)
webhooks_subscriptions = []

class WebhookSubscription(BaseModel):
    url: str
    event: str
    active: bool = True

@app.post("/api/webhooks/subscribe")
async def subscribe_webhook(subscription: WebhookSubscription):
    """Enregistre un webhook sans doublon"""
    global webhooks_subscriptions
    
    # Vérifier si existe déjà
    existing = [s for s in webhooks_subscriptions 
                if s["url"] == subscription.url and s["event"] == subscription.event]
    
    if existing:
        logger.info(f"🔗 Webhook déjà existant: {subscription.event} -> {subscription.url}")
        return {"success": True, "message": "Déjà inscrit", "subscriptions": webhooks_subscriptions}
    
    webhooks_subscriptions.append({
        "url": subscription.url,
        "event": subscription.event,
        "active": True,
        "created_at": datetime.now().isoformat()
    })
    
    logger.info(f"🔗 Webhook inscrit: {subscription.event} -> {subscription.url}")
    return {"success": True, "subscriptions": webhooks_subscriptions}
@app.get("/api/webhooks/subscriptions")
async def get_webhooks():
    """Liste tous les webhooks"""
    return {"success": True, "subscriptions": webhooks_subscriptions}

async def trigger_webhook(event: str, payload: dict):
    """Déclenche les webhooks pour un événement"""
    subscribers = [s for s in webhooks_subscriptions if s["event"] == event and s["active"]]
    
    logger.info(f"🔍 {len(subscribers)} webhook(s) trouvé(s) pour l'événement: {event}")
    
    if not subscribers:
        logger.info(f"📭 Aucun webhook inscrit pour {event}")
        return
    
    for sub in subscribers:
        try:
            logger.info(f"📤 Envoi à {sub['url']}...")
            async with httpx.AsyncClient() as client:
                response = await client.post(sub["url"], json=payload, timeout=10.0)
            logger.info(f"✅ Webhook envoyé à {sub['url']} - Status: {response.status_code}")
        except Exception as e:
            logger.error(f"❌ Erreur webhook {sub['url']}: {type(e).__name__} - {e}")



# =====================================================
# GOOGLE CALENDAR
# =====================================================

import json
from google.oauth2 import service_account
from googleapiclient.discovery import build
from typing import Optional, List

# Configuration Google Calendar
GOOGLE_CALENDAR_ID = os.environ.get("GOOGLE_CALENDAR_ID", "primary")
GOOGLE_SERVICE_ACCOUNT_INFO = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", None)

class CalendarEventRequest(BaseModel):
    summary: str
    description: Optional[str] = None
    start_datetime: str
    end_datetime: str
    attendees: Optional[List[str]] = None

def get_calendar_service():
    """Initialise le service Google Calendar avec le compte de service"""
    if not GOOGLE_SERVICE_ACCOUNT_INFO:
        logger.warning("⚠️ GOOGLE_SERVICE_ACCOUNT_JSON non configuré")
        return None
    
    try:
        # Charger les infos du compte de service
        service_account_info = json.loads(GOOGLE_SERVICE_ACCOUNT_INFO)
        
        creds = service_account.Credentials.from_service_account_info(
            service_account_info,
            scopes=['https://www.googleapis.com/auth/calendar']
        )
        
        service = build('calendar', 'v3', credentials=creds)
        logger.info("✅ Google Calendar service initialisé")
        return service
    except Exception as e:
        logger.error(f"❌ Erreur auth Google Calendar: {e}")
        return None

@app.post("/api/calendar/event")
async def create_calendar_event(request: CalendarEventRequest):
    """Crée un événement dans Google Calendar"""
    service = get_calendar_service()
    if not service:
        return {"success": False, "error": "Google Calendar non configuré"}
    
    try:
        event = {
            'summary': request.summary,
            'description': request.description or "",
            'start': {
                'dateTime': request.start_datetime,
                'timeZone': 'Africa/Porto-Novo',
            },
            'end': {
                'dateTime': request.end_datetime,
                'timeZone': 'Africa/Porto-Novo',
            },
        }
        
        if request.attendees:
            event['attendees'] = [{'email': email} for email in request.attendees]
        
        created_event = service.events().insert(
            calendarId=GOOGLE_CALENDAR_ID,
            body=event
        ).execute()
        
        logger.info(f"📅 Événement créé: {created_event.get('htmlLink')}")
        return {
            "success": True,
            "event_id": created_event.get('id'),
            "link": created_event.get('htmlLink')
        }
        
    except Exception as e:
        logger.error(f"❌ Erreur création événement: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/calendar/sync-task")
async def sync_task_to_calendar(request: Dict[str, Any]):
    task_id = request.get("task_id")
    user_id = get_request_user_id(request)
    """Synchronise une tâche spécifique vers Google Calendar"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    task_id = request.get("task_id")
    if not task_id:
        return {"success": False, "error": "task_id requis"}
    
    try:
        # Récupérer la tâche
        task = supabase.table("tasks").select("*").eq("user_id", user_id).eq("id", task_id).execute()
        
        if not task.data:
            return {"success": False, "error": "Tâche non trouvée"}
        
        task = task.data[0]
        
        if not task.get("due_date"):
            return {"success": False, "error": "La tâche n'a pas de date d'échéance"}
        
        # Créer l'événement
        event = await create_calendar_event(CalendarEventRequest(
            summary=task["title"],
            description=f"Tâche Sovereign - Priorité: {task.get('priority', 'normal')}",
            start_datetime=f"{task['due_date']}T09:00:00",
            end_datetime=f"{task['due_date']}T10:00:00"
        ))
        
        if event.get("success"):
            # Marquer comme synchronisée
            supabase.table("tasks").update({
                "calendar_synced": True,
                "calendar_event_id": event["event_id"]
            }).eq("id", task_id).execute()
        
        return event
        
    except Exception as e:
        logger.error(f"❌ Erreur sync tâche: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/calendar/sync-existing-task")
async def sync_existing_task_to_calendar(request: Dict[str, Any]):
    task_id = request.get("task_id")
    user_id = get_request_user_id(request)
    """Synchronise une tâche existante vers Google Calendar (pour rattrapage)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    task_id = request.get("task_id")
    if not task_id:
        return {"success": False, "error": "task_id requis"}
    
    try:
        # Récupérer la tâche
        task = supabase.table("tasks").select("*").eq("user_id", user_id).eq("id", task_id).execute()
        
        if not task.data:
            return {"success": False, "error": "Tâche non trouvée"}
        
        task = task.data[0]
        
        if not task.get("due_date"):
            return {"success": False, "error": "La tâche n'a pas de date d'échéance"}
        
        # Vérifier si déjà synchronisée
        if task.get("calendar_synced"):
            return {"success": True, "message": "Déjà synchronisée", "calendar_link": task.get("calendar_link")}
        
        # Créer l'événement
        event = await create_calendar_event(CalendarEventRequest(
            summary=task["title"],
            description=f"Tâche Sovereign - Priorité: {task.get('priority', 'normal')}",
            start_datetime=f"{task['due_date']}T09:00:00",
            end_datetime=f"{task['due_date']}T10:00:00"
        ))
        
        if event.get("success"):
            # Marquer comme synchronisée
            supabase.table("tasks").update({
                "calendar_synced": True,
                "calendar_event_id": event["event_id"],
                "calendar_link": event["link"]
            }).eq("id", task_id).execute()
            
            logger.info(f"📅 Tâche existante {task_id} synchronisée avec Google Calendar")
            return {
                "success": True, 
                "message": "Tâche synchronisée",
                "calendar_link": event["link"],
                "event_id": event["event_id"]
            }
        
        return event
        
    except Exception as e:
        logger.error(f"❌ Erreur sync tâche existante: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/clean-expired-subscriptions")
async def clean_expired_subscriptions():
    """Nettoie les subscriptions push expirées"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Récupérer toutes les subscriptions
        subscriptions = supabase.table("push_subscriptions").select("*").execute() 
        
        deleted_count = 0
        for sub in subscriptions.data:
            try:
                # Tester la subscription
                webpush(
                    subscription_info={
                        "endpoint": sub["endpoint"],
                        "keys": sub["keys"]
                    },
                    data=json.dumps({"title": "Test", "body": "Test"}),
                    vapid_private_key=VAPID_PRIVATE_KEY,
                    vapid_claims=VAPID_CLAIMS
                )
            except WebPushException as e:
                if e.response and e.response.status_code in [401, 403, 404, 410]:
                    # Supprimer la subscription invalide
                    supabase.table("push_subscriptions").delete().eq("id", sub["id"]).execute()
                    deleted_count += 1
                    logger.info(f"🗑️ Subscription supprimée: {sub['endpoint'][:50]}...")
        
        return {"success": True, "deleted": deleted_count, "total": len(subscriptions.data)}
        
    except Exception as e:
        logger.error(f"Erreur nettoyage: {e}")
        return {"success": False, "error": str(e)}



# =====================================================
# VOIX LIVE AVEC ELEVENLABS
# =====================================================
ELEVENLABS_API_KEY = os.environ.get("ELEVENLABS_API_KEY")
ELEVENLABS_VOICE_ID = os.environ.get("ELEVENLABS_VOICE_ID", "EXAVITQu4vr4xnSDxMaL")

@app.post("/api/voice/speak")
async def speak_text(request: Dict[str, Any]):
    """Convertit un texte en audio avec ElevenLabs"""
    text = request.get("text", "")
    if not text:
        return {"success": False, "error": "Texte requis"}
    
    if not ELEVENLABS_API_KEY:
        logger.warning("⚠️ ElevenLabs non configuré")
        return {"success": False, "error": "ElevenLabs non configuré"}
    
    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"https://api.elevenlabs.io/v1/text-to-speech/{ELEVENLABS_VOICE_ID}",
                headers={
                    "xi-api-key": ELEVENLABS_API_KEY,
                    "Content-Type": "application/json",
                    "Accept": "audio/mpeg"
                },
                json={
                    "text": text,
                    "model_id": "eleven_turbo_v2",
                    "voice_settings": {
                        "stability": 0.5,
                        "similarity_boost": 0.75
                    }
                },
                timeout=30.0
            )
            
            if response.status_code == 200:
                import base64
                audio_base64 = base64.b64encode(response.content).decode('utf-8')
                return {"success": True, "audio": audio_base64, "format": "mp3"}
            else:
                logger.error(f"Erreur ElevenLabs: {response.text}")
                return {"success": False, "error": response.text}
                
    except Exception as e:
        logger.error(f"Erreur ElevenLabs: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# MICROSOFT EDGE TTS (gratuit, sans carte bancaire)
# =====================================================

EDGE_TTS_TOKEN = os.environ.get("EDGE_TTS_TOKEN", "")

@app.post("/api/voice/edge-speak")
async def edge_speak_text(request: Dict[str, Any]):
    """Convertit un texte en audio avec Microsoft Edge TTS"""
    text = request.get("text", "")
    if not text:
        return {"success": False, "error": "Texte requis"}
    
    voice = request.get("voice", "fr-FR-DeniseNeural")
    
    try:
        import urllib.parse
        encoded_text = urllib.parse.quote(text)
        
        url = f"https://speech.platform.bing.com/consumer/speech/synthesize/readaloud/edge/v1?TrustedClientToken={EDGE_TTS_TOKEN}&Voice={voice}&Text={encoded_text}"
        
        async with httpx.AsyncClient(follow_redirects=True) as client:
            response = await client.get(url, timeout=30.0)
            
            if response.status_code == 200:
                import base64
                audio_base64 = base64.b64encode(response.content).decode('utf-8')
                return {"success": True, "audio": audio_base64, "format": "mp3", "voice": voice}
            else:
                logger.error(f"Erreur Edge TTS: {response.status_code}")
                return {"success": False, "error": f"Erreur {response.status_code}"}
                
    except Exception as e:
        logger.error(f"Erreur Edge TTS: {e}")
        return {"success": False, "error": str(e)}



# =====================================================
# PROACTIF - RÉSUMÉ MATINAL
# =====================================================
@app.post("/api/proactive/morning-brief")
async def send_morning_brief(request: Dict[str, Any] = None):
    """
    Envoie un résumé matinal personnalisé par email et notification push.
    Tous les messages sont générés dynamiquement par GPT-4o.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})

        today = datetime.now().date().isoformat()
        now = datetime.now()
        hour = now.hour
        
        # ========== RÉCUPÉRER LES DONNÉES CONTEXTUELLES ==========
        
        # 1. Tâches du jour
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        tasks_today_list = tasks_today.data
        tasks_count = len(tasks_today_list)
        
        # 2. Tâches en retard
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        overdue_count = len(overdue_tasks.data)
        overdue_tasks_list = overdue_tasks.data[:3]  # Les 3 plus urgentes
        
        # 3. Documents proches de l'échéance (7 jours)
        next_week = (datetime.now().date() + timedelta(days=7)).isoformat()
        expiring_docs = supabase.table("documents").select("*").eq("user_id", user_id).gte("due_date", today).lte("due_date", next_week).neq("status", "approved").execute()
        expiring_docs_list = expiring_docs.data[:3]
        
        # 4. Victoires récentes (7 jours)
        week_ago = (datetime.now().date() - timedelta(days=7)).isoformat()
        recent_wins = supabase.table("wins").select("*").eq("user_id", user_id).gte("date", week_ago).execute()
        wins_count = len(recent_wins.data)
        recent_wins_list = recent_wins.data[:3]
        
        # 5. Missions actives
        active_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").execute()
        missions_count = len(active_missions.data)
        missions_list = [{"name": m["name"], "priority": m.get("priority", "normal")} for m in active_missions.data[:3]]
        
        # 6. Humeur d'hier
        yesterday = (datetime.now().date() - timedelta(days=1)).isoformat()
        yesterday_mood = supabase.table("mood_entries").select("mood").eq("user_id", user_id).eq("date", yesterday).eq("user_id", user_id).execute()
        yesterday_mood_value = yesterday_mood.data[0]["mood"] if yesterday_mood.data else None
        
        # 7. Prochain événement familial
        next_family_event = supabase.table("family_events").select("*").eq("user_id", user_id).gte("date", today).neq("status", "done").order("date", ascending=True).limit(1).execute()
        
        # 8. Récupérer le nom et les enfants
        profile = supabase.table("user_profile").select("preferred_name, children").eq("user_id", user_id).execute()
        user_name = profile.data[0].get("preferred_name", "Rebecca") if profile.data else "Rebecca"
        children = profile.data[0].get("children", []) if profile.data else []
        
        # 9. Vérifier les anniversaires
        today_md = datetime.now().strftime("%m-%d")
        birthday_today = None
        for child in children:
            if child.get("birthday"):
                birthday_md = child["birthday"][5:] if len(child["birthday"]) > 5 else None
                if birthday_md == today_md:
                    birthday_today = child["name"]
        
        # 10. Mission à plus fort potentiel
        top_mission = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").order("revenue_potential", ascending=False).limit(1).execute()
        
        # 11. Budget restant
        total_revenue = supabase.table("revenue").select("amount").eq("user_id", user_id).execute()
        total_spending = supabase.table("spending").eq("user_id", user_id).select("amount").execute()
        revenue_sum = sum(r.get("amount", 0) for r in total_revenue.data)
        spending_sum = sum(s.get("amount", 0) for s in total_spending.data)
        balance = revenue_sum - spending_sum
        
        # ========== GÉNÉRATION PAR IA ==========
        
        prompt = f"""Tu es Becks, l'assistante personnelle de Rebecca. Génère un message matinal personnalisé et chaleureux.

CONTEXTE DU JOUR :
- Date : {datetime.now().strftime('%A %d %B %Y')}
- Heure : {hour}h
- Prénom de l'utilisatrice : {user_name}
- Humeur hier : {yesterday_mood_value or "Non renseignée"}
- Tâches aujourd'hui : {tasks_count} tâche(s) ({[t["title"] for t in tasks_today_list[:3]]})
- Tâches en retard : {overdue_count} ({[t["title"] for t in overdue_tasks_list]})
- Documents imminents : {len(expiring_docs_list)} ({[d["name"] for d in expiring_docs_list]})
- Victoires récentes : {wins_count} ({[w["title"] for w in recent_wins_list]})
- Missions actives : {missions_count} ({[m["name"] for m in missions_list]})
- Mission prioritaire : {top_mission.data[0]["name"] if top_mission.data else "Aucune"}
- Événement familial : {next_family_event.data[0]["title"] if next_family_event.data else "Aucun"}
- Anniversaire aujourd'hui : {birthday_today if birthday_today else "Non"}
- Solde financier : {balance:,.0f} CFA

Génère un message structuré avec :

1. Une salutation adaptée à l'heure (naturelle, pas "Bonjour" systématique)
2. Une phrase sur son état de la veille (bienveillante)
3. Un résumé des priorités du jour (basé sur les tâches et missions)
4. Un conseil ou une citation inspirante qui a du SENS avec sa situation (pas générique)
5. Une question ouverte pour l'engager

STYLE : Chaleureux, personnel, court (max 150 mots), comme une amie qui parle.
Ne liste pas les données brut, reformule naturellement.

Retourne UNIQUEMENT du JSON :
{{"greeting": "...", "mood_message": "...", "priorities_summary": "...", "wisdom": "...", "question": "..."}}"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.8,
            max_tokens=500
        )
        
        result_text = response.choices[0].message.content
        result_text = result_text.replace("```json", "").replace("```", "").strip()
        ai_content = json.loads(result_text)
        
        # ========== CONSTRUIRE LE MESSAGE FINAL ==========
        
        # Message spécifique pour anniversaire
        birthday_message = f"\n\n🎂 {birthday_today} fête son anniversaire aujourd'hui ! 🎉" if birthday_today else ""
        
        # Message financier si besoin
        financial_message = ""
        if balance < 0:
            financial_message = f"\n\n💰 Solde négatif de {abs(balance):,.0f} CFA. Une petite action aujourd'hui peut inverser la tendance."
        
        message = f"""{ai_content.get("greeting")} {user_name}.

{ai_content.get("mood_message")}

{ai_content.get("priorities_summary")}
{birthday_message}
{financial_message}

💡 {ai_content.get("wisdom")}

❓ {ai_content.get("question")}

Je suis là. 💖"""
        
        # ========== ENVOI ==========
        email_sent = False
        if BREVO_API_KEY:
            try:
                user_email = "jbillcataria@gmail.com"
                email_body = message.replace("\n", "<br>")
                await send_email(EmailRequest(
                    to=user_email,
                    subject=f"🌅 {user_name} - {datetime.now().strftime('%d/%m/%Y')}",
                    body=email_body
                ))
                email_sent = True
                logger.info("📧 Email matinal IA envoyé")
            except Exception as e:
                logger.error(f"Erreur envoi email: {e}")
        
        push_sent = False
        try:
            send_notification_sync({
                "title": f"🌅 {user_name}",
                "body": ai_content.get("priorities_summary", "Bonne journée !")[:80],
                "url": "/",
                "type": "brief",
                "user_id": user_id
            })
            push_sent = True
        except Exception as e:
            logger.error(f"Erreur envoi push: {e}")
        
        return {
            "success": True,
            "message": message,
            "stats": {
                "tasks_today": tasks_count,
                "overdue_tasks": overdue_count,
                "active_missions": missions_count,
                "recent_wins": wins_count,
                "balance": balance
            },
            "email_sent": email_sent,
            "push_sent": push_sent
        }
        
    except Exception as e:
        logger.error(f"Erreur résumé matinal: {e}")
        return {"success": False, "error": str(e)}

# =====================================================
# PROACTIF - VEILLE SUR PROJETS INACTIFS
# =====================================================

@app.post("/api/proactive/stale-missions")
async def check_stale_missions(request: Dict[str, Any] = None):
    """
    Vérifie les missions inactives (pas de mise à jour depuis 5 jours)
    et envoie des alertes par email et notification push.
    À appeler par cron-job.org toutes les 6h.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        # Récupérer les missions actives
        active_missions = (
            supabase.table("missions")
            .select("*")
            .eq("user_id", user_id)
            .eq("status", "active")
            .execute()
        )        
        if not active_missions.data:
            return {"success": True, "message": "Aucune mission active", "stale_missions": []}
        
        # Calculer la date limite (5 jours)
        five_days_ago = (datetime.now() - timedelta(days=5)).isoformat()
        
        # Filtrer les missions inactives
        stale_missions = []
        for mission in active_missions.data:
            updated_at = mission.get("updated_at")
            if updated_at and updated_at < five_days_ago:
                stale_missions.append(mission)
            elif not updated_at:
                # Si jamais mise à jour, vérifier created_at
                created_at = mission.get("created_at")
                if created_at and created_at < five_days_ago:
                    stale_missions.append(mission)
        
        if not stale_missions:
            return {"success": True, "message": "Aucune mission inactive", "stale_missions": []}
        
        # Construire le message
        mission_list = "\n".join([f"• {m['name']} (dernière activité: {m.get('updated_at', m.get('created_at', 'inconnue'))[:10]})" for m in stale_missions])
        
        message = f"""⚠️ **Alerte - Missions inactives**

{len(stale_missions)} mission(s) n'ont pas eu d'activité depuis plus de 5 jours :

{mission_list}

---

🎯 **Action recommandée** :
- Fais le point sur l'avancement de ces missions
- Mets à jour leur statut ou priorité
- Si terminées, passe-les en "complete"

Becks reste à ta disposition pour t'aider. 👑
"""
        
        # Envoyer un email
        email_sent = False
        if BREVO_API_KEY:
            try:
                user_email = "jbillcataria@gmail.com"   
                email_body = message.replace("\n", "<br>")
                await send_email(EmailRequest(
                    to=user_email,
                    subject=f"⚠️ {len(stale_missions)} mission(s) inactive(s) - Sovereign",
                    body=email_body
                ))
                email_sent = True
                logger.info(f"📧 Email missions inactives envoyé ({len(stale_missions)} missions)")
            except Exception as e:
                logger.error(f"Erreur envoi email missions inactives: {e}")
        
        # Envoyer une notification push
        push_sent = False
        try:
            send_notification_sync({
                "title": "⚠️ Missions inactives",
                "body": f"{len(stale_missions)} mission(s) sans activité depuis 5 jours",
                "url": "/missions",
                "type": "mission",
                "user_id": user_id,
                "requireInteraction": True
            })
            push_sent = True
            logger.info("🔔 Notification push missions inactives envoyée")
        except Exception as e:
            logger.error(f"Erreur envoi push missions inactives: {e}")
        
        return {
            "success": True,
            "message": f"{len(stale_missions)} mission(s) inactive(s) signalée(s)",
            "stale_missions": [
                {
                    "id": m["id"],
                    "name": m["name"],
                    "last_activity": m.get("updated_at", m.get("created_at"))
                }
                for m in stale_missions
            ],
            "email_sent": email_sent,
            "push_sent": push_sent
        }
        
    except Exception as e:
        logger.error(f"Erreur check stale missions: {e}")
        return {"success": False, "error": str(e)}




# =====================================================
# PROACTIF - DÉTECTION D'OPPORTUNITÉS
# =====================================================

@app.post("/api/proactive/opportunities-alert")
async def check_opportunities_alert(request: Dict[str, Any] = None):
    """
    Vérifie les grants et contrats proches de l'échéance (≤ 7 jours)
    et envoie des alertes.
    À appeler par cron-job.org toutes les 12h.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        today = datetime.now().date()
        next_week = today + timedelta(days=7)
        today_iso = today.isoformat()
        next_week_iso = next_week.isoformat()
        
        # Récupérer les grants proches de l'échéance
        # Récupérer les contrats proches de l'échéance
        # Récupérer les opportunités générales
        grants = (
            supabase.table("lf_grants")
            .select("*")
            .eq("user_id", user_id)
            .gte("deadline", today_iso)
            .lte("deadline", next_week_iso)
            .execute()
        )
        
        contracts = (
            supabase.table("lf_contracts")
            .select("*")
            .eq("user_id", user_id)
            .gte("deadline", today_iso)
            .lte("deadline", next_week_iso)
            .execute()
        )
        
        opportunities = (
            supabase.table("opportunities")
            .select("*")
            .eq("user_id", user_id)
            .gte("deadline", today_iso)
            .lte("deadline", next_week_iso)
            .neq("stage", "won")
            .execute()
        )
        
        all_items = []
        
        for grant in grants.data:
            all_items.append({
                "type": "grant",
                "title": grant.get("title"),
                "deadline": grant.get("deadline"),
                "agency": grant.get("agency"),
                "amount": grant.get("amount")
            })
        
        for contract in contracts.data:
            all_items.append({
                "type": "contract",
                "title": contract.get("title"),
                "deadline": contract.get("deadline"),
                "agency": contract.get("agency")
            })
        
        for opp in opportunities.data:
            all_items.append({
                "type": "opportunity",
                "title": opp.get("title"),
                "deadline": opp.get("deadline"),
                "estimated_value": opp.get("estimated_value")
            })
        
        if not all_items:
            return {"success": True, "message": "Aucune opportunité proche", "opportunities": []}
        
        # Construire le message
        items_by_day = {}
        for item in all_items:
            day = item["deadline"]
            if day not in items_by_day:
                items_by_day[day] = []
            items_by_day[day].append(item)
        
        message = f"""💰 **Alerte - Opportunités à saisir**

{len(all_items)} opportunité(s) approchent de leur échéance dans les 7 jours :

"""
        for day, items in sorted(items_by_day.items()):
            message += f"\n📅 **{day}** :\n"
            for item in items:
                if item["type"] == "grant":
                    message += f"   • 🎯 Grant: {item['title']} ({item.get('agency', 'N/A')}) - {item.get('amount', 0):,} CFA\n"
                elif item["type"] == "contract":
                    message += f"   • 📑 Contrat: {item['title']} ({item.get('agency', 'N/A')})\n"
                else:
                    message += f"   • 💼 Opportunité: {item['title']} - {item.get('estimated_value', 0):,} CFA\n"
        
        message += """

⚡ **Action recommandée** :
- Prépare les dossiers rapidement
- Programme des rappels pour ne rien oublier
- Contacte les parties prenantes dès aujourd'hui

Becks peut t'aider à préparer les documents. 👑
"""
        
        # Envoyer un email
        email_sent = False
        if BREVO_API_KEY:
            try:
                user_email = "rebecca@sovereign.com"  # À remplacer
                email_body = message.replace("\n", "<br>")
                await send_email(EmailRequest(
                    to=user_email,
                    subject=f"💰 {len(all_items)} opportunité(s) à saisir - Sovereign",
                    body=email_body
                ))
                email_sent = True
                logger.info(f"📧 Email opportunités envoyé ({len(all_items)} opportunités)")
            except Exception as e:
                logger.error(f"Erreur envoi email opportunités: {e}")
        
        # Envoyer une notification push
        push_sent = False
        try:
            send_notification_sync({
                "title": "💰 Opportunités à saisir",
                "body": f"{len(all_items)} opportunité(s) approchent de leur échéance",
                "url": "/love-fire-sport",
                "type": "money",
                "user_id": user_id,
                "requireInteraction": True
            })
            push_sent = True
            logger.info("🔔 Notification push opportunités envoyée")
        except Exception as e:
            logger.error(f"Erreur envoi push opportunités: {e}")
        
        return {
            "success": True,
            "message": f"{len(all_items)} opportunité(s) détectée(s)",
            "opportunities": all_items,
            "email_sent": email_sent,
            "push_sent": push_sent
        }
        
    except Exception as e:
        logger.error(f"Erreur check opportunities: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# PROACTIF - PLANNING AUTOMATIQUE
# =====================================================

@app.post("/api/proactive/daily-planning")
async def daily_planning(request: Dict[str, Any] = None):
    """
    Analyse les tâches et suggère un ordre de priorité pour la journée.
    À appeler par cron-job.org tous les matins à 8h (après le résumé).
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    user_id = get_request_user_id(request or {})
    
    try:
        today = datetime.now().date().isoformat()
        
        # Récupérer les tâches du jour
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        
        # Récupérer les tâches en retard
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        
        # Récupérer les tâches prioritaires (high/critical)
        high_priority_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).in_("priority", ["critical", "high"]).neq("status", "done").execute()
        
        # Construire la suggestion d'ordre
        planning = []
        
        # 1. D'abord les tâches en retard
        for task in overdue_tasks.data[:3]:
            planning.append({
                "position": len(planning) + 1,
                "title": task["title"],
                "reason": "⚠️ En retard",
                "priority": "critical"
            })
        
        # 2. Ensuite les tâches critiques
        for task in [t for t in high_priority_tasks.data if t.get("priority") == "critical" and t not in overdue_tasks.data][:3]:
            planning.append({
                "position": len(planning) + 1,
                "title": task["title"],
                "reason": "🔴 Priorité critique",
                "priority": "critical"
            })
        
        # 3. Puis les tâches du jour
        for task in tasks_today.data[:3]:
            if task not in overdue_tasks.data:
                planning.append({
                    "position": len(planning) + 1,
                    "title": task["title"],
                    "reason": "📅 À faire aujourd'hui",
                    "priority": "high"
                })
        
        # 4. Enfin les tâches haute priorité restantes
        for task in [t for t in high_priority_tasks.data if t.get("priority") == "high" and t not in tasks_today.data and t not in overdue_tasks.data][:2]:
            planning.append({
                "position": len(planning) + 1,
                "title": task["title"],
                "reason": "🔸 Haute priorité",
                "priority": "high"
            })
        
        if not planning:
            planning = [{"position": 1, "title": "Prendre un moment pour planifier", "reason": "Aucune tâche urgente", "priority": "normal"}]
        
        # Construire le message
        order_list = "\n".join([f"{p['position']}. **{p['title']}** - {p['reason']}" for p in planning])
        
        message = f"""📋 **Planning automatique du jour**

Voici l'ordre de priorité suggéré pour aujourd'hui :

{order_list}

---

💡 **Conseil** : Commence par la tâche n°1, elle débloquera la suite. N'hésite pas à ajuster selon ton énergie.

Becks te soutient ! 👑
"""
        
        # Envoyer une notification push
        push_sent = False
        try:
            send_notification_sync({
                "title": "📋 Planning du jour",
                "body": f"Priorité 1 : {planning[0]['title']}",
                "url": "/tasks",
                "type": "task",
                "user_id": user_id,
                "requireInteraction": False
            })
            push_sent = True
            logger.info("🔔 Notification push planning envoyée")
        except Exception as e:
            logger.error(f"Erreur envoi push planning: {e}")
        
        return {
            "success": True,
            "message": "Planning généré",
            "planning": planning,
            "stats": {
                "overdue_tasks": len(overdue_tasks.data),
                "tasks_today": len(tasks_today.data),
                "high_priority": len(high_priority_tasks.data)
            },
            "push_sent": push_sent
        }
        
    except Exception as e:
        logger.error(f"Erreur daily planning: {e}")
        return {"success": False, "error": str(e)}






# =====================================================
# BECKS EXECUTOR - ACTIONS CONCRÈTES
# =====================================================

class ExecutorAction(BaseModel):
    action_type: str  # "create_task", "send_email", "create_draft", "update_mission", "create_subtasks"
    params: Dict[str, Any]
    requires_confirmation: bool = True

@app.post("/api/executor/batch")
async def execute_batch_actions(actions: List[ExecutorAction], auto_confirm: bool = False):
    """
    Exécute une série d'actions proposées par Becks.
    Si auto_confirm=False, demande confirmation avant exécution.
    """
    results = []
    
    for action in actions:
        try:
            if action.action_type == "create_task":
                result = await create_task_from_conversation(ExecuteTaskRequest(
                    title=action.params.get("title"),
                    due_date=action.params.get("due_date"),
                    priority=action.params.get("priority", "normal")
                ))
                results.append({
                    "action": "create_task", 
                    "success": result.get("success"), 
                    "data": result.get("task")
                })
            
            elif action.action_type == "send_email":
                if auto_confirm or action.requires_confirmation == False:
                    # Récupérer l'email de l'utilisateur depuis la config
                    user_email = os.environ.get("USER_EMAIL", "rebecca@sovereign.com")
                    result = await send_email(EmailRequest(
                        to=action.params.get("to", user_email),
                        subject=action.params.get("subject", "Action Sovereign"),
                        body=action.params.get("body", "")
                    ))
                    results.append({"action": "send_email", "success": result.get("success")})
                else:
                    results.append({
                        "action": "send_email", 
                        "status": "pending_confirmation", 
                        "params": action.params
                    })
            
            elif action.action_type == "create_draft":
                result = await create_draft({
                    "type": action.params.get("type", "email"),
                    "context": action.params.get("context", "")
                })
                results.append({
                    "action": "create_draft", 
                    "success": result.get("success"), 
                    "data": result.get("draft")
                })
            
            elif action.action_type == "update_mission":
                mission_id = action.params.get("mission_id")
                if mission_id:
                    supabase.table("missions").update({
                        "status": action.params.get("status"),
                        "priority": action.params.get("priority")
                    }).eq("id", mission_id).execute()
                    results.append({"action": "update_mission", "success": True})
                else:
                    results.append({"action": "update_mission", "success": False, "error": "mission_id requis"})
            
            elif action.action_type == "create_subtasks":
                subtasks = action.params.get("subtasks", [])
                created = []
                for subtask in subtasks:
                    task = await create_task_from_conversation(ExecuteTaskRequest(
                        title=subtask.get("title"),
                        due_date=subtask.get("due_date"),
                        priority=subtask.get("priority", "normal")
                    ))
                    if task.get("success"):
                        created.append(task.get("task"))
                results.append({
                    "action": "create_subtasks", 
                    "success": len(created) > 0, 
                    "created": len(created),
                    "tasks": created
                })
            
            elif action.action_type == "create_reminder":
                result = await create_task_from_conversation(ExecuteTaskRequest(
                    title=action.params.get("title", "Rappel"),
                    due_date=action.params.get("due_date"),
                    priority="normal"
                ))
                results.append({"action": "create_reminder", "success": result.get("success"), "task": result.get("task")})
            
            else:
                results.append({
                    "action": action.action_type, 
                    "success": False, 
                    "error": f"Type d'action inconnu: {action.action_type}"
                })
        
        except Exception as e:
            logger.error(f"Erreur executor action {action.action_type}: {e}")
            results.append({"action": action.action_type, "success": False, "error": str(e)})
    
    return {"success": True, "results": results}


@app.post("/api/proactive/evening-summary")
async def send_evening_summary(request: Dict[str, Any] = None):
    """
    Envoie un résumé de fin de journée personnalisé par IA.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        today = datetime.now().date().isoformat()
        now = datetime.now()
        
        # ========== RÉCUPÉRER LES DONNÉES ==========
        
        # Tâches complétées aujourd'hui
        completed_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).eq("status", "done").gte("updated_at", today).execute()
        completed_count = len(completed_tasks.data)
        completed_list = [t["title"] for t in completed_tasks.data[:3]]
        
        # Tâches restantes (non terminées)
        pending_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).neq("status", "done").execute()
        pending_count = len(pending_tasks.data)
        
        # Tâches en retard
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        overdue_count = len(overdue_tasks.data)
        
        # Victoires du jour
        wins_today = supabase.table("wins").select("*").gte("date", today).execute()
        wins_count = len(wins_today.data)
        wins_list = [w["title"] for w in wins_today.data[:3]]
        
        # Humeur du jour
        mood_today = supabase.table("mood_entries").select("mood").eq("date", today).eq("user_id", user_id).execute()
        current_mood = mood_today.data[0]["mood"] if mood_today.data else None
        
        # Récupérer le nom
        profile = supabase.table("user_profile").select("preferred_name").eq("user_id", user_id).execute()
        user_name = profile.data[0].get("preferred_name", "Rebecca") if profile.data else "Rebecca"
        
        # ========== GÉNÉRATION IA ==========
        
        prompt = f"""Rebecca a terminé sa journée. Voici son bilan :
- Tâches faites : {completed_count} ({', '.join(completed_list) if completed_list else 'rien'})
- Tâches restantes : {pending_count}
- Tâches en retard : {overdue_count}
- Victoires du jour : {wins_count} ({', '.join(wins_list) if wins_list else 'aucune'})
- Humeur : {current_mood or 'non renseignée'}

Génère un message de fin de journée (max 60 mots) :
1. Reconnais ce qu'elle a accompli (ou non)
2. Un conseil court pour demain si nécessaire
3. Une phrase apaisante pour la nuit

Style : chaleureux, pas de bla-bla, direct mais doux.

Retourne UNIQUEMENT du JSON : {{"message": "..."}}"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=200
        )
        
        result_text = response.choices[0].message.content
        result_text = result_text.replace("```json", "").replace("```", "").strip()
        ai_content = json.loads(result_text)
        
        message = f"🌙 Bonsoir {user_name}.\n\n{ai_content.get('message')}\n\nRepose-toi bien. Demain est un nouveau jour. 👑"
        
        # ========== ENVOI ==========
        push_sent = False
        try:
            send_notification_sync({
                "title": "🌙 Fin de journée",
                "body": f"{completed_count} tâche(s) accomplie(s) • {wins_count} victoire(s)",
                "url": "/",
                "user_id": user_id,
                "type": "brief"
            })
            push_sent = True
        except Exception as e:
            logger.error(f"Erreur envoi push soir: {e}")
        
        # Email optionnel (1x par jour max)
        email_sent = False
        if BREVO_API_KEY and completed_count > 0:
            try:
                email_body = message.replace("\n", "<br>")
                await send_email(EmailRequest(
                    to="jbillcataria@gmail.com",
                    subject=f"🌙 Résumé du {datetime.now().strftime('%d/%m/%Y')}",
                    body=email_body
                ))
                email_sent = True
            except Exception as e:
                logger.error(f"Erreur envoi email soir: {e}")
        
        return {
            "success": True,
            "message": message,
            "stats": {
                "completed": completed_count,
                "pending": pending_count,
                "overdue": overdue_count,
                "wins": wins_count,
                "mood": current_mood
            },
            "push_sent": push_sent,
            "email_sent": email_sent
        }
        
    except Exception as e:
        logger.error(f"Erreur résumé soir: {e}")
        return {"success": False, "error": str(e)}
        

def _get_evening_advice(completed: int, pending: int, overdue: int) -> str:
    """Génère un conseil personnalisé pour le soir"""
    if overdue > 0:
        return "Des tâches sont en retard. Demain matin, attaque la plus urgente en premier. Je te rappellerai."
    elif pending > 0 and completed == 0:
        return "Tu n'as rien terminé aujourd'hui. Ce n'est pas grave. Demain, concentre-toi sur UNE seule petite tâche."
    elif completed >= 3:
        return "Belle journée ! Tu as bien avancé. Repose-toi, tu as mérité cette soirée."
    elif completed > 0:
        return f"Tu as accompli {completed} tâche(s). Chaque pas compte. Demain, continue sur cette lancée."
    else:
        return "Parfois, se reposer est la meilleure action. Demain sera plus clair."





# =====================================================
# PROACTIF - RAPPELS INTELLIGENTS
# =====================================================

@app.post("/api/proactive/smart-reminders")
async def smart_reminders(request: Dict[str, Any] = None):
    """
    Analyse l'historique et propose des actions intelligentes.
    À appeler par cron-job.org tous les matins à 8h30.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    user_id = get_request_user_id(request or {})
    
    try:
        today = datetime.now().date().isoformat()
        week_ago = (datetime.now().date() - timedelta(days=7)).isoformat()
        
        # 1. Analyser les tâches récurrentes oubliées
        # Tâches qui apparaissent souvent mais rarement terminées
        task_titles = supabase.table("tasks").select("title, status").eq("user_id", user_id).execute()
        task_count = {}
        for task in task_titles.data:
            title = task["title"]
            if title not in task_count:
                task_count[title] = {"total": 0, "done": 0}
            task_count[title]["total"] += 1
            if task.get("status") == "done":
                task_count[title]["done"] += 1
        
        forgotten_tasks = []
        for title, stats in task_count.items():
            if stats["total"] >= 3 and stats["done"] == 0:
                forgotten_tasks.append(title)
        
        # 2. Vérifier les missions sans activité
        stale_missions = supabase.table("missions").select("*").eq("status", "active").execute()
        stale_list = []
        for mission in stale_missions.data:
            updated_at = mission.get("updated_at")
            if updated_at and updated_at < week_ago:
                stale_list.append(mission["name"])
        
        # 3. Vérifier les documents qui traînent
        pending_docs = supabase.table("documents").select("*").eq("status", "draft").execute()
        
        # 4. Vérifier les opportunités non suivies
        pending_opps = supabase.table("opportunities").select("*").neq("stage", "won").neq("stage", "lost").execute()
        high_value_opps = [o for o in pending_opps.data if o.get("estimated_value", 0) > 1000000]
        
        # Construire les suggestions
        suggestions = []
        
        if forgotten_tasks:
            suggestions.append(f"📋 **Tâches récurrentes à faire** : {', '.join(forgotten_tasks[:3])}")
        
        if stale_list:
            suggestions.append(f"🎯 **Missions sans activité récente** : {', '.join(stale_list[:3])}")
        
        if pending_docs.data:
            suggestions.append(f"📄 **Documents en brouillon** : {len(pending_docs.data)} document(s) à finaliser")
        
        if high_value_opps:
            total_value = sum(o.get("estimated_value", 0) for o in high_value_opps)
            suggestions.append(f"💰 **Opportunités à suivre** : {len(high_value_opps)} opportunité(s) - {total_value:,.0f} CFA potentiel")
        
        if not suggestions:
            suggestions = ["✅ Rien d'urgent à signaler. Bonne journée !"]
        
        # Construire le message
        message = f"""🔔 **Rappels intelligents - {datetime.now().strftime('%A %d %B %Y')}**

{chr(10).join(suggestions)}

---

💡 **Une action aujourd'hui ?** 
Réponds-moi directement ou clique sur une suggestion pour que je t'aide.
"""
        
        # Envoyer notification push
        push_sent = False
        try:
            send_notification_sync({
                "title": "🔔 Rappels intelligents",
                "body": suggestions[0][:50],
                "url": "/chat",
                "type": "task",
                "user_id": user_id,
                "requireInteraction": False

            })
            push_sent = True
            logger.info("🔔 Notification push rappels intelligents envoyée")
        except Exception as e:
            logger.error(f"Erreur envoi push rappels: {e}")
        
        return {
            "success": True,
            "message": "Rappels intelligents envoyés",
            "suggestions": suggestions,
            "stats": {
                "forgotten_tasks": len(forgotten_tasks),
                "stale_missions": len(stale_list),
                "pending_docs": len(pending_docs.data),
                "high_value_opps": len(high_value_opps)
            },
            "push_sent": push_sent
        }
        
    except Exception as e:
        logger.error(f"Erreur rappels intelligents: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/proactive/test-smart-reminders")
async def test_smart_reminders():
    """Endpoint de test pour les rappels intelligents"""
    return await smart_reminders()



# =====================================================
# GOOGLE DRIVE - SAUVEGARDE AUTOMATIQUE
# =====================================================

import json
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseUpload
import io

GOOGLE_DRIVE_SERVICE_ACCOUNT_INFO = os.environ.get("GOOGLE_DRIVE_SERVICE_ACCOUNT", None)
GOOGLE_DRIVE_FOLDER_ID = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "root")

def get_drive_service():
    """Initialise le service Google Drive"""
    if not GOOGLE_DRIVE_SERVICE_ACCOUNT_INFO:
        logger.warning("⚠️ GOOGLE_DRIVE_SERVICE_ACCOUNT non configuré")
        return None
    
    try:
        service_account_info = json.loads(GOOGLE_DRIVE_SERVICE_ACCOUNT_INFO)
        creds = service_account.Credentials.from_service_account_info(
            service_account_info,
            scopes=['https://www.googleapis.com/auth/drive']
        )
        service = build('drive', 'v3', credentials=creds)
        logger.info("✅ Google Drive service initialisé")
        return service
    except Exception as e:
        logger.error(f"❌ Erreur auth Google Drive: {e}")
        return None

@app.post("/api/drive/backup-document")
async def backup_document_to_drive(request: Dict[str, Any]):
    """Sauvegarde un document dans Google Drive"""
    service = get_drive_service()
    if not service:
        return {"success": False, "error": "Google Drive non configuré"}
    
    document_id = request.get("document_id")
    if not document_id:
        return {"success": False, "error": "document_id requis"}
    
    try:
        # Récupérer le document depuis Supabase
        doc = supabase.table("documents").select("*").eq("id", document_id).execute()
        
        if not doc.data:
            return {"success": False, "error": "Document non trouvé"}
        
        doc = doc.data[0]
        
        # Créer le contenu du fichier
        content = f"""Document: {doc['name']}
Type: {doc['type']}
Statut: {doc['status']}
Date d'échéance: {doc.get('due_date', 'Non définie')}
Notes: {doc.get('notes', 'Aucune')}

URL originale: {doc.get('url', 'Non renseignée')}
Fichier: {doc.get('file_url', 'Non renseigné')}

--- Exporté depuis Sovereign le {datetime.now().strftime('%d/%m/%Y à %H:%M')} ---
"""
        
        # Créer le fichier dans Drive
        file_metadata = {
            'name': f"{doc['name']}.txt",
            'parents': [GOOGLE_DRIVE_FOLDER_ID]
        }
        
        media = MediaIoBaseUpload(
            io.BytesIO(content.encode('utf-8')),
            mimetype='text/plain',
            resumable=True
        )
        
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id, webViewLink'
        ).execute()
        
        # Mettre à jour le document avec le lien Drive
        supabase.table("documents").update({
            "drive_backup_id": file.get('id'),
            "drive_link": file.get('webViewLink'),
            "backed_up_at": datetime.now().isoformat()
        }).eq("id", document_id).execute()
        
        logger.info(f"📁 Document {doc['name']} sauvegardé dans Google Drive")
        
        return {
            "success": True,
            "message": "Document sauvegardé dans Google Drive",
            "drive_link": file.get('webViewLink'),
            "drive_id": file.get('id')
        }
        
    except Exception as e:
        logger.error(f"❌ Erreur backup Drive: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/drive/backup-all-documents")
async def backup_all_documents_to_drive():
    """Sauvegarde tous les documents non encore sauvegardés"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Récupérer les documents non sauvegardés
        docs = supabase.table("documents").select("*").is_("drive_backup_id", "null").execute()
        
        results = []
        for doc in docs.data:
            result = await backup_document_to_drive({"document_id": doc["id"]})
            results.append({
                "id": doc["id"],
                "name": doc["name"],
                "success": result.get("success", False),
                "drive_link": result.get("drive_link")
            })
        
        backed_up = len([r for r in results if r["success"]])
        
        return {
            "success": True,
            "message": f"{backed_up} document(s) sauvegardé(s) sur {len(results)}",
            "results": results
        }
        
    except Exception as e:
        logger.error(f"❌ Erreur backup all: {e}")
        return {"success": False, "error": str(e)}

#------------------------------------------------
#------------------------------------------------

@app.get("/api/morning-greeting")
async def get_morning_greeting(user_id: Optional[str] = None, force: bool = False):
    """Retourne un message d'accueil personnalisé (une fois par jour)"""
    if not supabase:
        return {"success": True, "message": "Salut Rebecca. Je suis là."}
    
    try:
        user_id = require_user_id(user_id)
        today = datetime.now().date().isoformat()
        
        # Vérifier si le rapport a déjà été envoyé aujourd'hui
        if not force:
            existing = supabase.table("daily_reports")\
                .select("*")\
                .eq("user_id", user_id)\
                .eq("report_date", today)\
                .execute()
            
            if existing.data:
                # Rapport déjà envoyé aujourd'hui
                return {
                    "success": True, 
                    "already_sent": True,
                    "message": None
                }
        
        # ========== RÉCUPÉRER LES DONNÉES ==========
        today = datetime.now().date().isoformat()
        hour = datetime.now().hour
        
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        active_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").execute()
        urgent_docs = supabase.table("documents").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "approved").execute()
        
        # WhatsApp en attente
        pending_whatsapp = supabase.table("whatsapp_messages")\
            .select("*")\
            .eq("user_id", user_id)\
            .eq("status", "pending")\
            .order("importance", desc=True)\
            .execute()
        
        whatsapp_high = [m for m in pending_whatsapp.data if m.get("importance") == "high"]
        
        # Humeur d'hier
        yesterday = (datetime.now().date() - timedelta(days=1)).isoformat()
        mood_yesterday = supabase.table("mood_entries").select("mood").eq("user_id", user_id).eq("date", yesterday).execute()
        yesterday_mood = mood_yesterday.data[0]["mood"] if mood_yesterday.data else None
        
        # Prénom
        profile = supabase.table("user_profile").select("preferred_name").eq("user_id", user_id).execute()
        user_name = profile.data[0].get("preferred_name", "Rebecca") if profile.data else "Rebecca"
        
        # ========== GÉNÉRATION IA ==========
        hour = datetime.now().hour
        if hour < 12:
            greeting_prefix = "☀️ Bonjour"
        elif hour < 18:
            greeting_prefix = "🌤️ Bon après-midi"
        else:
            greeting_prefix = "🌙 Bonsoir"
        
        prompt = f"""Génère un message d'accueil VIVANT et NATUREL pour Rebecca.

Contexte :
- Prénom : {user_name}
- Moment : {greeting_prefix}
- Humeur hier : {yesterday_mood or "non renseignée"}
- Tâches aujourd'hui : {len(tasks_today.data)} (dont {len(overdue_tasks.data)} en retard)
- Missions actives : {len(active_missions.data)}
- Documents urgents : {len(urgent_docs.data)}
- WhatsApp importants : {len(whatsapp_high)}

Règles :
- Sois naturelle, vivante, humaine (pas de liste de chiffres)
- Maximum 50 mots
- Termine par une question ouverte

Exemple : "{greeting_prefix} {user_name}. Tu as {len(tasks_today.data)} choses à faire aujourd'hui, dont {len(overdue_tasks.data)} en retard. On commence par la plus urgente ?"

Retourne UNIQUEMENT le message, rien d'autre."""

        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.8,
                max_tokens=150
            )
            greeting = response.choices[0].message.content.strip()
        except Exception as e:
            # Fallback
            if len(overdue_tasks.data) > 0:
                greeting = f"{greeting_prefix} {user_name}. Tu as {len(overdue_tasks.data)} tâche(s) en retard. On s'en occupe ?"
            elif len(tasks_today.data) > 0:
                greeting = f"{greeting_prefix} {user_name}. {len(tasks_today.data)} chose(s) à faire aujourd'hui. On y va ?"
            elif len(whatsapp_high) > 0:
                greeting = f"{greeting_prefix} {user_name}. {len(whatsapp_high)} message(s) WhatsApp t'attendent. Je te les montre ?"
            else:
                greeting = f"{greeting_prefix} {user_name}. Journée calme. Besoin de quelque chose ?"
        
        # Sauvegarder qu'on a envoyé le rapport
        supabase.table("daily_reports").insert({
            "user_id": user_id,
            "report_date": today,
            "report_content": greeting
        }).execute()
        
        return {"success": True, "message": greeting, "already_sent": False}
        
    except Exception as e:
        logger.error(f"Erreur morning greeting: {e}")
        return {"success": True, "message": f"Salut Rebecca. Je suis là."}

#-------------------------------------------------------------------------------
#get_today_dashboard
#-------------------------------------------------------------------------------
@app.get("/api/dashboard/today")
async def get_today_dashboard(user_id: Optional[str] = None):
    """Retourne toutes les données nécessaires pour le dashboard du jour avec des messages humains"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = require_user_id(user_id)
        today = datetime.now().date().isoformat()
        
        # 🔥 RÉCUPÉRER LE NOM DE L'UTILISATEUR
        profile_result = supabase.table("user_profile").select("preferred_name").eq("user_id", user_id).execute()
        user_name = profile_result.data[0].get("preferred_name", "Rebecca") if profile_result.data else "Rebecca"
        
        # Récupérer les tâches du jour
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        
        # Récupérer les tâches en retard
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        
        # Récupérer les missions actives
        active_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").execute()
        
        # Récupérer les documents en attente
        pending_docs = supabase.table("documents").select("*").eq("user_id", user_id).neq("status", "approved").execute()
        
        # Récupérer les victoires récentes (7 jours)
        week_ago = (datetime.now().date() - timedelta(days=7)).isoformat()
        recent_wins = supabase.table("wins").select("*").eq("user_id", user_id).gte("date", week_ago).execute()
        
        # Récupérer les événements familiaux du jour
        family_today = supabase.table("family_events").select("*").eq("user_id", user_id).eq("date", today).neq("status", "done").execute()
        
        # Récupérer l'humeur du jour
        mood_result = supabase.table("mood_entries").select("mood").eq("user_id", user_id).eq("date", today).execute()
        current_mood = mood_result.data[0]["mood"] if mood_result.data else None
        
        # Calculer les priorités (basé sur l'urgence)
        priorities = []
        
        # 1. Tâches en retard (priorité max)
        for task in overdue_tasks.data[:2]:
            priorities.append({
                "id": task["id"],
                "title": task["title"],
                "reason": "⚠️ En retard",
                "score": 40
            })
        
        # 2. Tâches du jour
        for task in tasks_today.data[:3 - len(priorities)]:
            if task["id"] not in [p["id"] for p in priorities]:
                priorities.append({
                    "id": task["id"],
                    "title": task["title"],
                    "reason": "📅 À faire aujourd'hui",
                    "score": 30
                })
        
        # 3. Missions actives à fort potentiel
        if len(priorities) < 3:
            for mission in active_missions.data[:3 - len(priorities)]:
                priorities.append({
                    "id": mission["id"],
                    "title": f"🎯 {mission['name']}",
                    "reason": "Mission stratégique",
                    "score": 20
                })
        
        # ============================================
        # GÉNÉRER UN MESSAGE HUMAIN ET PERSONNALISÉ AVEC LE VRAI NOM
        # ============================================
        
        hour = datetime.now().hour
        if hour < 12:
            greeting_prefix = "☀️ Bonjour"
        elif hour < 18:
            greeting_prefix = "🌤️ Bon après-midi"
        else:
            greeting_prefix = "🌙 Bonsoir"
        
        # Adapter le message selon l'humeur
        if current_mood == "fatiguée":
            mood_phrase = "Je sens que tu es fatiguée."
        elif current_mood == "stressée":
            mood_phrase = "Je sens que tu es stressée."
        elif current_mood == "excellent":
            mood_phrase = "Tu as de l'énergie aujourd'hui !"
        elif current_mood == "bien":
            mood_phrase = "Content de te sentir bien."
        else:
            mood_phrase = ""
        
        # 🔥 Construire un message personnalisé avec user_name au lieu de "Rebecca"
        if overdue_tasks.data:
            task_list = ", ".join([t["title"][:35] for t in overdue_tasks.data[:2]])
            if len(overdue_tasks.data) == 1:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Tu as une tâche en retard : « {task_list} ». On s'en occupe maintenant ?"
            else:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Tu as {len(overdue_tasks.data)} tâches en retard. La plus urgente : « {task_list} ». Je t'aide à prioriser ?"
        
        elif tasks_today.data:
            task_list = ", ".join([t["title"][:35] for t in tasks_today.data[:2]])
            if len(tasks_today.data) == 1:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Ta tâche du jour : « {task_list} ». On y va ?"
            else:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Tu as {len(tasks_today.data)} choses à faire aujourd'hui. La première : « {task_list} »."
        
        elif pending_docs.data:
            doc_count = len(pending_docs.data)
            if doc_count == 1:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Tu as un document qui t'attend. Besoin que je t'aide à le remplir ?"
            else:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Tu as {doc_count} documents en attente. On fait le point ?"
        
        elif active_missions.data:
            mission_names = ", ".join([m["name"] for m in active_missions.data[:2]])
            if len(active_missions.data) == 1:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Ta mission active : {mission_names}. Tu veux qu'on avance dessus ?"
            else:
                greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Tes missions actives : {mission_names}. Par laquelle tu veux commencer ?"
        
        elif recent_wins.data:
            win_count = len(recent_wins.data)
            greeting = f"{greeting_prefix} {user_name}. {mood_phrase} Tu as {win_count} victoire(s) récente(s) ! C'est bien. Continue comme ça."
        
        else:
            # Messages plus naturels quand rien d'urgent
            natural_greetings = [
                f"{greeting_prefix} {user_name}. Rien de prévu aujourd'hui. Tu veux qu'on avance sur un projet ou tu préfères souffler ?",
                f"{greeting_prefix} {user_name}. Journée calme. Profites-en pour respirer ou pour prendre de l'avance.",
                f"{greeting_prefix} {user_name}. Tout est calme. Besoin de quoi ?",
                f"{greeting_prefix} {user_name}. Pas de pression aujourd'hui. Dis-moi ce que tu veux faire."
            ]
            greeting = random.choice(natural_greetings)
        
        # Suggestions de moves personnalisées
        suggestions = {
            "money_move": "Vérifier les finances du jour",
            "family_move": "Prendre des nouvelles des enfants",
            "business_move": "Avancer sur une mission prioritaire",
            "stabilization_move": "Prendre 5 minutes pour respirer"
        }
        
        if active_missions.data:
            suggestions["business_move"] = f"Avancer sur {active_missions.data[0]['name']}"
        
        if family_today.data:
            suggestions["family_move"] = f"{family_today.data[0]['title']} aujourd'hui"
        
        if recent_wins.data:
            suggestions["stabilization_move"] = f"Célébrer {len(recent_wins.data)} victoire(s) récente(s)"
        
        # Message de guidance calme (naturel)
        if current_mood == "stressée":
            calm_guidance = "Respire. Rien n'est aussi urgent qu'il n'y paraît. On y va doucement."
        elif current_mood == "fatiguée":
            calm_guidance = "Repose-toi si tu en as besoin. La productivité peut attendre."
        elif overdue_tasks.data:
            calm_guidance = "Les tâches en retard, c'est stressant, mais on va les gérer une par une."
        elif tasks_today.data:
            calm_guidance = "Une chose à la fois. Tu vas y arriver."
        else:
            calm_guidance = "Profite de ce moment de calme. Tu l'as mérité."
        
        return {
            "success": True,
            "greeting": greeting,
            "top_priorities": priorities,
            "tasks_today": tasks_today.data[:5],
            "overdue_tasks": overdue_tasks.data[:5],
            "active_missions": active_missions.data[:5],
            "pending_docs": pending_docs.data[:5],
            "recent_wins": len(recent_wins.data),
            "family_today_count": len(family_today.data),
            "suggestions": suggestions,
            "calm_guidance": calm_guidance,
            "stats": {
                "tasks_count": len(tasks_today.data),
                "overdue_count": len(overdue_tasks.data),
                "missions_count": len(active_missions.data),
                "docs_count": len(pending_docs.data),
                "wins_count": len(recent_wins.data)
            },
            "current_mood": current_mood
        }
        
    except Exception as e:
        logger.error(f"Erreur dashboard today: {e}")
        return {"success": False, "error": str(e)}



@app.post("/api/morning-notification")
async def send_morning_notification(request: Dict[str, Any] = None):
    """Envoie une notification matinale humaine avec son et vibration"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        today = datetime.now().date().isoformat()
        hour = datetime.now().hour
        
        # Récupérer les vraies données
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        tasks_count = len(tasks_today.data)
        
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        overdue_count = len(overdue_tasks.data)
        
        pending_docs = supabase.table("documents").select("*").eq("user_id", user_id).neq("status", "approved").execute()
        docs_count = len(pending_docs.data)
        
        active_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").execute()
        missions_count = len(active_missions.data)
        
        # Construire un message humain (une seule phrase)
        if overdue_count > 0:
            body = f"⚠️ {overdue_count} tâche(s) en retard. On regarde ça ensemble ?"
        elif tasks_count > 0:
            body = f"📋 {tasks_count} chose(s) à faire aujourd'hui. Je suis là si tu veux."
        elif docs_count > 0:
            body = f"📄 {docs_count} document(s) en attente. Besoin d'aide ?"
        elif missions_count > 0:
            body = f"🎯 {missions_count} mission(s) active(s). Je suis là si tu veux."
        else:
            body = f"☀️ Bonne journée Rebecca. Je suis là si tu as besoin."
        
        # Envoyer la notification push
        notification_data = {
            "title": "🌅 Rebecca",
            "body": body,
            "url": "/chat",
            "type": "morning",
            "sound": "/sounds/notification.mp3",
            "vibrate": [200, 100, 200],
            "requireInteraction": False,
            "silent": False,
            "user_id": user_id,
            "tag": f"morning_{today}"
        }
        
        results = send_notification_sync(notification_data)
        
        return {
            "success": True,
            "notification_sent": len(results) > 0,
            "message": body
        }
        
    except Exception as e:
        logger.error(f"Erreur morning notification: {e}")
        return {"success": False, "error": str(e)}


@app.put("/api/profile/identity")
async def update_identity(request: Request):
    """Met à jour uniquement l'identité (preferred_name, full_name, birthday)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Récupérer le user_id depuis les query params
        user_id = request.query_params.get("user_id")
        
        if not user_id:
            return {"success": False, "error": "user_id requis"}
        
        # Récupérer le body
        body = await request.json()
        
        # Champs autorisés pour l'identité
        identity_fields = ["preferred_name", "full_name", "birthday"]
        
        clean_data = {}
        for key, value in body.items():
            if key in identity_fields and value is not None:
                clean_data[key] = value
        
        if not clean_data:
            return {"success": False, "error": "Aucun champ valide à mettre à jour"}
        
        clean_data["updated_at"] = datetime.now().isoformat()
        
        # Mettre à jour avec user_id dynamique
        supabase.table("user_profile").update(clean_data).eq("user_id", user_id).execute()
        
        # Récupérer le profil mis à jour
        profile_result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        return {"success": True, "profile": profile_result.data[0] if profile_result.data else {}}
        
    except Exception as e:
        logger.error(f"Erreur update_identity: {e}")
        return {"success": False, "error": str(e)}


@app.put("/api/profile/children")
async def update_children(request: Request):
    """Met à jour uniquement les enfants"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Récupérer le user_id depuis les query params
        user_id = request.query_params.get("user_id")
        
        if not user_id:
            return {"success": False, "error": "user_id requis"}
        
        # Récupérer le body
        body = await request.json()
        children = body.get("children", [])
        
        # Mettre à jour avec user_id dynamique
        supabase.table("user_profile").update({
            "children": children,
            "updated_at": datetime.now().isoformat()
        }).eq("user_id", user_id).execute()
        
        # Récupérer le profil mis à jour
        profile_result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        return {"success": True, "profile": profile_result.data[0] if profile_result.data else {}}
        
    except Exception as e:
        logger.error(f"Erreur update_children: {e}")
        return {"success": False, "error": str(e)}


@app.put("/api/profile/projects")
async def update_projects(request: Request):
    """Met à jour uniquement les projets"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Récupérer le user_id depuis les query params
        user_id = request.query_params.get("user_id")
        
        if not user_id:
            return {"success": False, "error": "user_id requis"}
        
        # Récupérer le body
        body = await request.json()
        projects = body.get("projects", [])
        
        logger.info(f"📥 Mise à jour des projets pour {user_id}: {len(projects)} projets")
        
        # Mettre à jour avec user_id dynamique
        supabase.table("user_profile").update({
            "projects": projects,
            "updated_at": datetime.now().isoformat()
        }).eq("user_id", user_id).execute()
        
        # Récupérer le profil mis à jour
        profile_result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        return {"success": True, "profile": profile_result.data[0] if profile_result.data else {}}
        
    except Exception as e:
        logger.error(f"❌ Erreur update_projects: {e}")
        return {"success": False, "error": str(e)}


@app.put("/api/profile/goals")
async def update_goals(request: Request):
    """Met à jour uniquement les objectifs"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        # Récupérer le user_id depuis les query params
        user_id = request.query_params.get("user_id")
        
        if not user_id:
            return {"success": False, "error": "user_id requis"}
        
        # Récupérer le body
        body = await request.json()
        goals = body.get("current_goals", [])
        
        # Mettre à jour avec user_id dynamique
        supabase.table("user_profile").update({
            "current_goals": goals,
            "updated_at": datetime.now().isoformat()
        }).eq("user_id", user_id).execute()
        
        # Récupérer le profil mis à jour
        profile_result = supabase.table("user_profile").select("*").eq("user_id", user_id).execute()
        
        return {"success": True, "profile": profile_result.data[0] if profile_result.data else {}}
        
    except Exception as e:
        logger.error(f"Erreur update_goals: {e}")
        return {"success": False, "error": str(e)}
# =====================================================
# EXECUTION GUIDE - STEP BY STEP
# =====================================================

@app.post("/api/execute/step-by-step")
async def step_by_step_execution(request: Dict[str, Any]):
    """
    Transforme une demande en plan d'action étape par étape.
    Retourne une structure avec des étapes à cocher.
    """
    query = request.get("query", "")
    context = request.get("context", {})  # Contexte optionnel (mission_id, etc.)
    
    if not query:
        return {"success": False, "error": "Query requise"}
    user_id = get_request_user_id(request or {})
    
    try:
        # Récupérer le contexte utilisateur pour personnaliser
        profile_context = await get_profile_context(user_id=user_id)
        memory_context = await get_user_memory_context(user_id)
        
        system_prompt = f"""Tu es Becks, l'agent d'exécution de Rebecca. Tu transformes une demande en plan d'action CONCRET, ÉTAPE PAR ÉTAPE.

Tu connais Rebecca : {profile_context}
Informations sur elle : {memory_context}

RÈGLES IMPORTANTES :
1. Chaque étape doit être ACTIONNABLE (commencer par un verbe)
2. Chaque étape doit être RÉALISTE (max 10-15 min par étape)
3. Maximum 6 étapes par plan
4. Inclus des "victoires rapides" (étapes très courtes)
5. Si une info manque, transforme-la en étape "Trouver X"

Retourne UNIQUEMENT du JSON valide avec cette structure :

{{
  "title": "Titre du plan d'action",
  "estimated_duration": "durée estimée (ex: 30 minutes, 2 heures)",
  "steps": [
    {{ "description": "Étape 1", "action_type": "type d'action (task/email/document/call/decision/research)", "estimated_minutes": 5 }},
    {{ "description": "Étape 2", "action_type": "task", "estimated_minutes": 10 }}
  ],
  "success_criteria": "Ce qui définit la réussite",
  "next_steps_hint": "Ce qu'on pourra faire après"
}}

Types d'action possibles : task, email, document, call, decision, research, wait, celebrate

Si la demande est émotionnelle (fatigue, stress, etc.), retourne une version très courte (2-3 étapes) avec action_type="celebrate" ou "rest"."""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": query}
            ],
            temperature=0.3,
            max_tokens=800
        )
        
        result_text = response.choices[0].message.content
        result_text = result_text.replace("```json", "").replace("```", "").strip()
        execution_plan = json.loads(result_text)
        
        # Sauvegarder le plan en base pour suivi
        plan_id = str(uuid.uuid4())
        if supabase:
            supabase.table("execution_plans").insert({
                "id": plan_id,
                "title": execution_plan.get("title"),
                "steps": execution_plan.get("steps", []),
                "status": "active",
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }).execute()
        
        return {
            "success": True,
            "plan_id": plan_id,
            "plan": execution_plan
        }
        
    except Exception as e:
        logger.error(f"Erreur step_by_step: {e}")
        return {"success": False, "error": str(e), "fallback": True}


@app.post("/api/execute/complete-step")
async def complete_execution_step(request: Dict[str, Any]):
    """Marque une étape comme complétée"""
    plan_id = request.get("plan_id")
    step_index = request.get("step_index")
    user_id = get_request_user_id(request)
    
    if not plan_id or step_index is None:
        return {"success": False, "error": "plan_id et step_index requis"}
    
    try:
        if not supabase:
            return {"success": True, "message": "Pas de base, étape fictive"}
        
        # Récupérer le plan
        result = (
            supabase.table("execution_plans")
            .select("*")
            .eq("user_id", user_id)
            .eq("id", plan_id)
            .execute()
        )        
        if not result.data:
            return {"success": False, "error": "Plan non trouvé"}
        
        plan = result.data[0]
        steps = plan.get("steps", [])
        completed_steps = plan.get("completed_steps", [])
        
        if step_index not in completed_steps:
            completed_steps.append(step_index)
            
            # Envoyer une notification de progression
            send_notification_sync({
                "title": "✅ Étape accomplie",
                "body": f"'{steps[step_index].get('description', 'Étape')}' - Reste {len(steps) - len(completed_steps)} étape(s)",
                "url": "/chat",
                "user_id": user_id,
                "type": "task"
            })
        
        progress = int((len(completed_steps) / len(steps)) * 100) if steps else 0
        
        supabase.table("execution_plans").update({
            "completed_steps": completed_steps,
            "progress": progress,
            "updated_at": datetime.now().isoformat()
        }).eq("id", plan_id).execute()
        
        return {
            "success": True,
            "completed_steps": completed_steps,
            "progress": progress,
            "remaining": len(steps) - len(completed_steps),
            "is_complete": len(completed_steps) >= len(steps)
        }
        
    except Exception as e:
        logger.error(f"Erreur complete_step: {e}")
        return {"success": False, "error": str(e)}

@app.post("/api/memory/detect-patterns")
async def detect_memory_patterns(request: Dict[str, Any]):
    """
    Analyse les souvenirs pour détecter des patterns (récurrences, tendances)
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    user_id = get_request_user_id(request or {})
    
    try:
        # Récupérer tous les souvenirs
        memories = supabase.table("user_memory").select("*").eq("user_id", user_id).execute()
        
        if not memories.data:
            return {"success": True, "patterns": []}
        
        # Analyser par catégorie
        patterns = {}
        for mem in memories.data:
            cat = mem.get("category")
            if cat not in patterns:
                patterns[cat] = {"count": 0, "keys": []}
            patterns[cat]["count"] += 1
            patterns[cat]["keys"].append(mem.get("key"))
        
        # Détecter les catégories les plus utilisées
        sorted_cats = sorted(patterns.items(), key=lambda x: x[1]["count"], reverse=True)
        
        # Générer des insights
        insights = []
        for cat, data in sorted_cats[:3]:
            if data["count"] > 3:
                insights.append(f"Tu as {data['count']} souvenirs dans la catégorie {cat}")
        
        return {
            "success": True,
            "patterns": patterns,
            "insights": insights,
            "total_count": len(memories.data)
        }
        
    except Exception as e:
        logger.error(f"Erreur detect_patterns: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# RESCUE MODE - DÉTECTION AUTOMATIQUE DE SURCHARGE
# =====================================================


@app.post("/api/rescue/detect-overload")
async def detect_overload(request: Dict[str, Any] = None):
    """
    Analyse les données pour détecter si l'utilisateur est en surcharge.
    Retourne un score et des recommandations.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = require_user_id(user_id)
        today = datetime.now().date().isoformat()
        now = datetime.now()
        
        # 1. Récupérer l'humeur du jour
        mood_result = supabase.table("mood_entries").select("mood").eq("date", today).eq("user_id", user_id).execute()
        current_mood = mood_result.data[0]["mood"] if mood_result.data else None
        
        # 2. Tâches urgentes et en retard
        urgent_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).eq("status", "today").neq("status", "done").execute()
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        
        # 3. Documents en retard
        overdue_docs = supabase.table("documents").select("*").lt("due_date", today).neq("status", "approved").execute()
        
        # 4. Missions actives non avancées
        active_missions = supabase.table("missions").select("*").eq("status", "active").execute()
        stale_missions = [m for m in active_missions.data if not m.get("updated_at") or m["updated_at"] < (datetime.now() - timedelta(days=5)).isoformat()]
        
        # 5. Brain dumps non traités
        pending_brain_dumps = supabase.table("inbox").select("*").eq("needs_processing", True).execute()
        
        # Calcul du score de surcharge (0-100)
        overload_score = 0
        reasons = []
        
        # Humeur
        if current_mood == "stressée":
            overload_score += 25
            reasons.append("😰 Humeur stressée détectée")
        elif current_mood == "fatiguée":
            overload_score += 20
            reasons.append("😴 Humeur fatiguée détectée")
        
        # Tâches
        if len(overdue_tasks.data) > 0:
            overload_score += min(len(overdue_tasks.data) * 10, 30)
            reasons.append(f"⚠️ {len(overdue_tasks.data)} tâche(s) en retard")
        elif len(urgent_tasks.data) > 3:
            overload_score += min(len(urgent_tasks.data) * 5, 20)
            reasons.append(f"📋 {len(urgent_tasks.data)} tâches urgentes")
        
        # Documents
        if len(overdue_docs.data) > 0:
            overload_score += min(len(overdue_docs.data) * 8, 20)
            reasons.append(f"📄 {len(overdue_docs.data)} document(s) en retard")
        
        # Missions
        if len(stale_missions) > 0:
            overload_score += min(len(stale_missions) * 5, 15)
            reasons.append(f"🎯 {len(stale_missions)} mission(s) sans activité")
        
        # Brain dumps
        if len(pending_brain_dumps.data) > 3:
            overload_score += min(len(pending_brain_dumps.data) * 2, 10)
            reasons.append(f"🧠 {len(pending_brain_dumps.data)} idées non traitées")
        
        # Heure tardive (après 20h)
        if now.hour >= 20 and len(urgent_tasks.data) > 0:
            overload_score += 10
            reasons.append("🌙 Heure tardive + tâches restantes")
        
        overload_score = min(overload_score, 100)
        
        # Déterminer le niveau
        if overload_score >= 60:
            level = "critical"
        elif overload_score >= 35:
            level = "high"
        elif overload_score >= 15:
            level = "moderate"
        else:
            level = "low"
        
        # ========== GÉNÉRER LE MESSAGE D'ENCOURAGEMENT PAR IA ==========
        encouragement_message = None
        try:
            # Appel à l'endpoint de génération d'encouragement
            ai_response = await generate_rescue_encouragement({
                "overload_score": overload_score,
                "level": level,
                "reasons": reasons,
                "current_mood": current_mood,
                "urgent_tasks": len(urgent_tasks.data),
                "overdue_tasks": len(overdue_tasks.data)
            })
            if ai_response.get("success"):
                encouragement_message = ai_response.get("encouragement")
        except Exception as e:
            logger.error(f"Erreur génération encouragement IA: {e}")
        
        # Fallback si l'IA a échoué
        if not encouragement_message:
            if overload_score >= 60:
                encouragement_message = "⚠️ Tu es en surcharge sévère. Active le Rescue Mode immédiatement."
            elif overload_score >= 35:
                encouragement_message = "🟡 Charge élevée. Ralentis et priorise l'essentiel."
            elif overload_score >= 15:
                encouragement_message = "🟢 Charge modérée. Reste focus."
            else:
                encouragement_message = "🌿 Tout va bien. Profite de ce calme."
        
        # Générer des actions de secours
        rescue_actions = []
        
        if len(overdue_tasks.data) > 0:
            rescue_actions.append({
                "type": "focus_task",
                "title": "Faire la tâche la plus urgente",
                "task_id": overdue_tasks.data[0]["id"],
                "task_title": overdue_tasks.data[0]["title"]
            })
        
        if current_mood in ["stressée", "fatiguée"]:
            rescue_actions.append({
                "type": "breathing",
                "title": "Prendre 3 respirations profondes",
                "duration": 1
            })
            rescue_actions.append({
                "type": "reset",
                "title": "Faire une pause de 10 minutes",
                "duration": 10
            })
        
        if len(pending_brain_dumps.data) > 2:
            rescue_actions.append({
                "type": "brain_dump",
                "title": "Vider son esprit dans Brain Dump",
                "url": "/inbox"
            })
        
        rescue_actions.append({
            "type": "chat",
            "title": "Parler à Becks",
            "url": "/chat?mode=parle-moi"
        })
        
        return {
            "success": True,
            "overload_score": overload_score,
            "level": level,
            "message": encouragement_message,
            "reasons": reasons,
            "rescue_actions": rescue_actions[:4],  # Max 4 actions
            "stats": {
                "urgent_tasks": len(urgent_tasks.data),
                "overdue_tasks": len(overdue_tasks.data),
                "overdue_docs": len(overdue_docs.data),
                "stale_missions": len(stale_missions),
                "pending_brain_dumps": len(pending_brain_dumps.data)
            },
            "current_mood": current_mood
        }
        
    except Exception as e:
        logger.error(f"Erreur detect_overload: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/rescue/encouragement")
async def generate_rescue_encouragement(request: Dict[str, Any]):
    """
    Génère un message d'encouragement personnalisé pour le Rescue Mode.
    """
    overload_score = request.get("overload_score", 0)
    level = request.get("level", "medium")
    reasons = request.get("reasons", [])
    current_mood = request.get("current_mood")
    urgent_tasks = request.get("urgent_tasks", 0)
    overdue_tasks = request.get("overdue_tasks", 0)
    
    # Déterminer le niveau d'urgence pour le prompt
    if level == "critical":
        urgency_prompt = "La situation est CRITIQUE. Elle est en surcharge sévère."
    elif level == "high":
        urgency_prompt = "La situation est tendue. Elle a beaucoup de choses à gérer."
    else:
        urgency_prompt = "La situation est modérée, mais elle a besoin de recentrage."
    
    # Construire le prompt
    prompt = f"""Génère un message d'encouragement court et réconfortant pour Rebecca.

CONTEXTE :
- Niveau de surcharge : {level} (score: {overload_score}/100)
- Humeur actuelle : {current_mood if current_mood else "non renseignée"}
- Tâches urgentes : {urgent_tasks}
- Tâches en retard : {overdue_tasks}
- Raisons identifiées : {', '.join(reasons) if reasons else "Aucune raison spécifique"}

RÈGLES :
- Sois très douce et bienveillante
- Maximum 30 mots
- Ne propose PAS d'action (juste du réconfort)
- Si c'est critique → "Respire. Ralentis. Rien n'est aussi urgent."
- Si humeur stressée → "Je suis là. On va y aller doucement."
- Si humeur fatiguée → "Repose-toi d'abord. Le reste attendra."

Retourne UNIQUEMENT le message, rien d'autre."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=80
        )
        
        message = response.choices[0].message.content.strip()
        
        return {"success": True, "encouragement": message}
        
    except Exception as e:
        logger.error(f"Erreur génération encouragement: {e}")
        # Fallback humain
        fallbacks = {
            "critical": "🌿 Respire profondément. Rien n'est aussi urgent qu'il n'y paraît. Je suis là.",
            "high": "💖 Une chose à la fois. Tu n'es pas seule dans cette tempête.",
            "moderate": "✨ Ralentis. Priorise l'essentiel. Le reste attendra.",
            "low": "🌸 Tout va bien. Tu gères. Prends soin de toi."
        }
        return {"success": True, "encouragement": fallbacks.get(level, fallbacks["moderate"])}
# =====================================================
# NOTIFICATIONS PROACTIVES INTELLIGENTES
# =====================================================

@app.post("/api/notifications/intelligent-check")
async def intelligent_notification_check(request: Dict[str, Any] = None):
    """
    Analyse le contexte et envoie des notifications intelligentes si nécessaire.
    À appeler via cron toutes les 2-3 heures.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        now = datetime.now()
        today = now.date().isoformat()
        hour = now.hour
        
        notifications_sent = []
        
        # 1. Vérifier si déjà envoyé récemment (éviter spam)
        last_check = supabase.table("notifications_log").select("sent_at")\
            .eq("type", "intelligent_check")\
            .eq("user_id", user_id)\
            .order("sent_at", desc=True)\
            .limit(1)\
            .execute()
        
        if last_check.data:
            last_sent = datetime.fromisoformat(last_check.data[0]["sent_at"])
            if (now - last_sent).total_seconds() < 7200:  # 2 heures
                return {"success": True, "message": "Déjà vérifié récemment", "notifications_sent": []}
        
        # 2. Récupérer l'humeur du jour
        mood_result = supabase.table("mood_entries").select("mood").eq("date", today).eq("user_id", user_id).execute()
        current_mood = mood_result.data[0]["mood"] if mood_result.data else None
        
        # 3. Récupérer les tâches
        urgent_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).eq("status", "today").neq("status", "done").execute()
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        
        # 4. Récupérer les brain dumps non traités
        pending_brain_dumps = supabase.table("inbox").select("*").eq("needs_processing", True).execute()
        
        # 5. Récupérer les victoires récentes
        week_ago = (now.date() - timedelta(days=7)).isoformat()
        recent_wins = supabase.table("wins").select("*").gte("date", week_ago).execute()
        
        # 6. Générer des notifications contextuelles
        notifications = []
        
        # ========== NOTIFICATION SURCHARGE ==========
        if current_mood in ["stressée", "fatiguée"] and len(urgent_tasks.data) > 2:
            notifications.append({
                "title": "🌿 Prends soin de toi",
                "body": f"Je sens que tu es {current_mood} et tu as {len(urgent_tasks.data)} tâches. On respire. Une chose à la fois.",
                "url": "/rescue",
                "type": "rescue"
            })
        
        # ========== NOTIFICATION TÂCHES EN RETARD (version douce) ==========
        elif len(overdue_tasks.data) > 0:
            if current_mood in ["stressée", "fatiguée"]:
                notifications.append({
                    "title": "📋 Des tâches t'attendent",
                    "body": f"Tu as {len(overdue_tasks.data)} tâche(s) en retard. On les regarde ensemble ? Pas de pression.",
                    "url": "/tasks",
                    "type": "task"
                })
            else:
                notifications.append({
                    "title": "⚠️ Tâches en retard",
                    "body": f"Tu as {len(overdue_tasks.data)} tâche(s) en retard. Priorisons les plus importantes.",
                    "url": "/tasks",
                    "type": "task"
                })
        
        # ========== NOTIFICATION BRAIN DUMP ==========
        if len(pending_brain_dumps.data) > 2 and current_mood not in ["stressée", "fatiguée"]:
            notifications.append({
                "title": "🧠 Des idées t'attendent",
                "body": f"Tu as {len(pending_brain_dumps.data)} idées non traitées dans ton Brain Dump. Je peux les organiser pour toi.",
                "url": "/inbox",
                "type": "brain_dump"
            })
        
        # ========== NOTIFICATION VICTOIRES ==========
        if len(recent_wins.data) == 0 and current_mood in ["fatiguée", "neutre"]:
            notifications.append({
                "title": "🏆 Une petite victoire aujourd'hui ?",
                "body": "Même une petite chose mérite d'être célébrée. Ajoute ta victoire du jour ✨",
                "url": "/wins",
                "type": "win"
            })
        
        # ========== NOTIFICATION FIN DE JOURNÉE ==========
        if hour >= 19 and hour <= 21 and len(urgent_tasks.data) > 0:
            notifications.append({
                "title": "🌙 Fin de journée",
                "body": f"Il te reste {len(urgent_tasks.data)} tâche(s). Demain est un autre jour. Repose-toi bien.",
                "url": "/tasks",
                "type": "evening"
            })
        
        # ========== NOTIFICATION MORNING CHECK-IN ==========
        if 7 <= hour <= 9 and not notifications:
            notifications.append({
                "title": "☀️ Bonjour Rebecca",
                "body": "Comment te sens-tu aujourd'hui ? Je suis là pour t'aider.",
                "url": "/chat",
                "type": "morning"
            })
        
        # 7. Envoyer les notifications (max 2 par check)
        for notif in notifications[:2]:
            send_notification_sync({
                "title": notif["title"],
                "body": notif["body"],
                "url": notif["url"],
                "type": notif["type"],
                "sound": "/sounds/notification.mp3",
                "vibrate": [200, 100, 200],
                "user_id": user_id,
                "requireInteraction": notif["type"] in ["rescue", "task"]
            })
            notifications_sent.append(notif)
        
        # 8. Logger
        if notifications_sent:
            supabase.table("notifications_log").insert({
                "type": "intelligent_check",
                "date": today,
                "user_id": user_id,
                "sent_at": now.isoformat(),
                "metadata": {"notifications": len(notifications_sent), "mood": current_mood}
            }).execute()
        
        return {
            "success": True,
            "notifications_sent": len(notifications_sent),
            "details": notifications_sent,
            "context": {
                "mood": current_mood,
                "urgent_tasks": len(urgent_tasks.data),
                "overdue_tasks": len(overdue_tasks.data),
                "pending_brain_dumps": len(pending_brain_dumps.data),
                "recent_wins": len(recent_wins.data),
                "hour": hour
            }
        }
        
    except Exception as e:
        logger.error(f"Erreur intelligent_notification_check: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/notifications/preferences")
async def update_notification_preferences(request: Dict[str, Any]):
    """Met à jour les préférences de notifications de l'utilisateur"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        preferences = request.get("preferences", {})
        
        # Sauvegarder dans user_memory
        for key, value in preferences.items():
            await save_user_memory("notifications", key, str(value))
        
        return {"success": True, "preferences": preferences}
        
    except Exception as e:
        logger.error(f"Erreur update_preferences: {e}")
        return {"success": False, "error": str(e)}


@app.get("/api/notifications/preferences")
async def get_notification_preferences(user_id: Optional[str] = None):
    """Récupère les préférences de notifications"""
    if not supabase:
        return {"success": False, "preferences": {}}
    user_id = require_user_id(user_id)
    
    try:
    
        result = supabase.table("user_memory").select("*").eq("category", "notifications").eq("user_id", user_id).execute()
        
        preferences = {}
        for item in result.data:
            preferences[item["key"]] = item["value"] == "True" if item["value"] in ["True", "False"] else item["value"]
        
        # Valeurs par défaut
        default_preferences = {
            "morning_checkin": True,
            "evening_summary": True,
            "overdue_tasks": True,
            "brain_dump_reminder": True,
            "win_reminder": True,
            "intelligent_mode": True
        }
        
        default_preferences.update(preferences)
        
        return {"success": True, "preferences": default_preferences}
        
    except Exception as e:
        logger.error(f"Erreur get_preferences: {e}")
        return {"success": False, "preferences": {}}


# =====================================================
# CHECK-IN PROACTIF MATINAL
# =====================================================

@app.post("/api/morning-checkin")
async def send_morning_checkin(request: Dict[str, Any] = None):
    """
    Envoie une notification proactive le matin avec un message personnalisé.
    À appeler via cron tous les matins entre 7h et 9h.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        now = datetime.now()
        hour = now.hour
        today = now.date().isoformat()
        
        # Vérifier si déjà envoyé aujourd'hui
        existing = supabase.table("notifications_log").select("*")\
            .eq("type", "morning_checkin")\
            .eq("date", today)\
            .eq("user_id", user_id)\
            .execute()
        
        if existing.data:
            return {"success": True, "sent": False, "message": "Déjà envoyé aujourd'hui"}
        
        # 1. Récupérer l'humeur d'hier
        yesterday = (now.date() - timedelta(days=1)).isoformat()
        mood_result = supabase.table("mood_entries").select("mood").eq("date", yesterday).eq("user_id", user_id).execute()
        yesterday_mood = mood_result.data[0]["mood"] if mood_result.data else None
        
        # 2. Récupérer les tâches d'aujourd'hui
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        
        # 3. Récupérer les documents en retard
        overdue_docs = supabase.table("documents").select("*").lt("due_date", today).neq("status", "approved").execute()
        
        # 4. Récupérer les victoires d'hier
        wins_yesterday = supabase.table("wins").select("*").eq("date", yesterday).execute()
        
        # 5. Générer un message personnalisé
        greeting = "Bonjour"
        if hour < 9:
            greeting = "☀️ Bonjour"
        elif hour < 12:
            greeting = "🌤️ Bonjour"
        else:
            greeting = "👋 Bonjour"
        
        # Construire le message selon le contexte
        if yesterday_mood == "stressée":
            mood_message = "Je sens que hier était stressant. Aujourd'hui, on y va doucement."
        elif yesterday_mood == "fatiguée":
            mood_message = "Tu étais fatiguée hier. Priorise ton énergie aujourd'hui."
        elif yesterday_mood == "excellent":
            mood_message = "Tu étais en forme hier ! Continue sur cette lancée."
        else:
            mood_message = "J'espère que tu as bien dormi."
        
        # Message sur les tâches
        if len(tasks_today.data) > 0:
            task_message = f"Tu as {len(tasks_today.data)} tâche(s) aujourd'hui."
            first_task = tasks_today.data[0].get("title", "")
            if first_task:
                task_message += f" La plus importante : {first_task}"
        else:
            task_message = "Aucune tâche planifiée. Une journée pour respirer ?"
        
        # Message sur les documents
        if len(overdue_docs.data) > 0:
            doc_message = f"⚠️ {len(overdue_docs.data)} document(s) en retard. On les regarde ?"
        else:
            doc_message = "✅ Aucun document en retard."
        
        # Message sur les victoires
        if len(wins_yesterday.data) > 0:
            win_message = f"🏆 Hier, tu as célébré {len(wins_yesterday.data)} victoire(s) !"
        else:
            win_message = "✨ N'oublie pas de célébrer tes victoires, même petites."
        
        # Message final personnalisé
        final_message = f"""{greeting} Rebecca.

{mood_message}

📋 {task_message}
{doc_message}
{win_message}

Je suis là pour t'aider. Une chose à la fois. 👑"""

        # Envoyer la notification push
        send_notification_sync({
            "title": "🌅 Rebecca",
            "body": final_message[:200],  # Limite de caractères pour la notification
            "url": "/chat",
            "type": "morning_checkin",
            "sound": "/sounds/notification.mp3",
            "vibrate": [200, 100, 200],
            "requireInteraction": False,
            "user_id": user_id,
            "silent": False,
            "tag": f"morning_checkin_{today}"
        })
        
        # Logger l'envoi
        supabase.table("notifications_log").insert({
            "type": "morning_checkin",
            "date": today,
            "user_id": user_id,
            "sent_at": now.isoformat(),
            "metadata": {
                "tasks_count": len(tasks_today.data),
                "overdue_docs": len(overdue_docs.data),
                "wins_yesterday": len(wins_yesterday.data),
                "yesterday_mood": yesterday_mood
            }
        }).execute()
        
        return {
            "success": True,
            "sent": True,
            "message": final_message,
            "stats": {
                "tasks_today": len(tasks_today.data),
                "overdue_docs": len(overdue_docs.data),
                "wins_yesterday": len(wins_yesterday.data),
                "yesterday_mood": yesterday_mood
            }
        }
        
    except Exception as e:
        logger.error(f"Erreur morning_checkin: {e}")
        return {"success": False, "error": str(e)}


@app.get("/api/morning-checkin/test")
async def test_morning_checkin():
    """Endpoint de test pour le check-in matinal (ignore la vérification d'envoi)"""
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = DEFAULT_USER_ID
        now = datetime.now()
        today = now.date().isoformat()
        
        # Récupérer les données (sans vérifier si déjà envoyé)
        yesterday = (now.date() - timedelta(days=1)).isoformat()
        mood_result = supabase.table("mood_entries").select("mood").eq("date", yesterday).eq("user_id", user_id).execute()
        yesterday_mood = mood_result.data[0]["mood"] if mood_result.data else None
        
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        overdue_docs = supabase.table("documents").select("*").lt("due_date", today).neq("status", "approved").execute()
        wins_yesterday = supabase.table("wins").select("*").eq("date", yesterday).execute()
        
        greeting = "☀️ Bonjour" if now.hour < 12 else "👋 Bonjour"
        
        mood_message = ""
        if yesterday_mood == "stressée":
            mood_message = "Je sens que hier était stressant. Aujourd'hui, on y va doucement."
        elif yesterday_mood == "fatiguée":
            mood_message = "Tu étais fatiguée hier. Priorise ton énergie aujourd'hui."
        else:
            mood_message = "J'espère que tu as bien dormi."
        
        final_message = f"{greeting} Rebecca.\n\n{mood_message}\n\n📋 {len(tasks_today.data)} tâche(s) aujourd'hui.\n📄 {len(overdue_docs.data)} document(s) en retard.\n🏆 {len(wins_yesterday.data)} victoire(s) hier.\n\nJe suis là pour t'aider. 👑"
        
        return {
            "success": True,
            "test_message": final_message,
            "stats": {
                "tasks_today": len(tasks_today.data),
                "overdue_docs": len(overdue_docs.data),
                "wins_yesterday": len(wins_yesterday.data),
                "yesterday_mood": yesterday_mood
            }
        }
        
    except Exception as e:
        logger.error(f"Erreur test_morning_checkin: {e}")
        return {"success": False, "error": str(e)}



# =====================================================
# WEEKLY CEO VIEW - VERSION P
# =====================================================

def _generate_human_weekly_insight(completion_rate: int, wins_count: int, balance: int, overdue_docs: int, tasks_completed: int, tasks_created: int, overdue_tasks: int) -> str:
    """
    Génère un insight humain et personnalisé - version NATURELLE
    """
    import random
    
    # Phrases d'ouverture naturelles
    openers = [
        "Alors cette semaine,",
        "Voilà ce que je retiens,",
        "En regardant ta semaine,",
        "Tu sais quoi ?",
        "Franchement,",
        "Ce que je vois, c'est que"
    ]
    
    # Analyser la situation réelle
    has_progress = tasks_completed > 0 or wins_count > 0
    has_stress = overdue_docs > 0 or overdue_tasks > 0
    has_money_issue = balance < 0
    is_low_energy = completion_rate < 30 and tasks_created > 0
    is_great_week = completion_rate >= 70 and wins_count >= 2
    
    opener = random.choice(openers)
    
    # Semaine exceptionnelle
    if is_great_week:
        return f"{opener} {tasks_completed} tâches terminées, {wins_count} victoires célébrées. Franchement, t'as géré cette semaine. Garde cette énergie, elle te mène loin. 👑"
    
    # Semaine avec progrès mais stress
    if has_stress and has_progress:
        return f"{opener} tu as fait {tasks_completed} tâche(s), mais il reste {overdue_docs} document(s) et {overdue_tasks} tâche(s) qui traînent. C'est normal. On les prend un par un, pas la peine de tout gérer d'un coup. 🌿"
    
    # Semaine difficile sans progression
    if has_stress and not has_progress:
        if overdue_docs > 0:
            return f"{opener} cette semaine a été lourde. {overdue_docs} document(s) en retard, et tu n'as rien pu cocher. Parfois, tenir le coup est déjà une victoire. Demain est un autre jour. 💪"
        else:
            return f"{opener} tu as créé {tasks_created} tâche(s) mais rien n'est terminé. C'est peut-être un signe : tu as besoin de ralentir. Écoute-toi. 🌙"
    
    # Problème d'argent
    if has_money_issue:
        return f"{opener} le solde est négatif ({abs(balance):,.0f} CFA). Ce n'est qu'un chiffre, pas un jugement. Une petite action cette semaine peut inverser la tendance. On regarde ça ensemble ? 💰"
    
    # Basse énergie
    if is_low_energy:
        return f"{opener} tu as créé {tasks_created} tâche(s) mais peu sont terminées. Parfois, ralentir permet de mieux repartir. C'est peut-être ce dont tu avais besoin. 🌿"
    
    # Progression modérée
    if completion_rate >= 30:
        if wins_count > 0:
            return f"{opener} {completion_rate}% des tâches sont faites, et tu as {wins_count} victoire(s) à célébrer. Chaque pas compte, même petit. Continue comme ça. ✨"
        else:
            return f"{opener} {completion_rate}% des tâches sont terminées. Une progression modérée, mais une progression quand même. C'est bien. 🌱"
    
    # Message par défaut chaleureux
    encouragements = [
        f"{opener} {tasks_completed} tâche(s) terminée(s) cette semaine. Parfois, ralentir permet de mieux repartir. 🌿",
        f"{opener} {tasks_completed} tâche(s) accomplie(s). Chaque petite victoire construit ton empire. 👑",
        f"{opener} tu as fait ce que tu as pu. C'est suffisant. La semaine prochaine est une nouvelle page. 💖"
    ]
    
    return random.choice(encouragements)


@app.get("/api/weekly-ceo")
async def get_weekly_ceo(user_id: Optional[str] = None):
    """
    Retourne une vue stratégique hebdomadaire avec des messages humains.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = require_user_id(user_id)
        now = datetime.now()
        start_of_week = (now.date() - timedelta(days=now.weekday())).isoformat()
        end_of_week = (now.date() + timedelta(days=6 - now.weekday())).isoformat()
        
        # 1. Tâches complétées cette semaine
        completed_tasks = supabase.table("tasks").select("*")\
            .gte("updated_at", start_of_week)\
            .eq("user_id", user_id)\
            .eq("status", "done")\
            .execute()
        
        # 2. Tâches créées cette semaine
        new_tasks = supabase.table("tasks").select("*")\
            .eq("user_id", user_id)\
            .gte("created_at", start_of_week)\
            .execute()
        
        # 3. Missions actives
        active_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").execute()
        
        # 4. Missions avec le plus haut potentiel de revenu
        missions_by_revenue = supabase.table("missions").select("*")\
            .eq("user_id", user_id)\
            .order("revenue_potential", desc=True)\
            .limit(5)\
            .execute()
        
        # 5. Documents en attente ou en retard
        pending_docs = supabase.table("documents").select("*")\
            .eq("user_id", user_id)\
            .neq("status", "approved")\
            .execute()
        
        overdue_docs = supabase.table("documents").select("*")\
            .eq("user_id", user_id)\
            .lt("due_date", now.date().isoformat())\
            .neq("status", "approved")\
            .execute()
        
        # 6. Victoires cette semaine
        wins_this_week = supabase.table("wins").select("*")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        
        # 7. Dépenses cette semaine
        spending_this_week = supabase.table("spending").select("amount")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        total_spending = sum(s.get("amount", 0) for s in spending_this_week.data)
        
        # 8. Revenus cette semaine
        revenue_this_week = supabase.table("revenue").select("amount")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        total_revenue = sum(r.get("amount", 0) for r in revenue_this_week.data)
        
        # 9. Humeurs de la semaine
        moods_this_week = supabase.table("mood_entries").select("mood, date")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        
        # Calculer la mission la plus proche du cash
        closest_to_cash = None
        for mission in missions_by_revenue.data:
            if mission.get("status") == "active" and mission.get("revenue_potential", 0) > 3:
                closest_to_cash = mission.get("name")
                break
        
        # 10. Tâches en retard
        overdue_tasks = supabase.table("tasks").select("*")\
            .eq("user_id", user_id)\
            .lt("due_date", now.date().isoformat())\
            .eq("user_id", user_id)\
            .neq("status", "done")\
            .execute()
        
        # Calcul du taux de complétion
        completion_rate = 0
        if len(new_tasks.data) > 0:
            completion_rate = int((len(completed_tasks.data) / len(new_tasks.data)) * 100)
        
        # ========== GÉNÉRATION DES MESSAGES HUMAINS ==========
        
        # Message "what_moved" - version humaine
        if len(completed_tasks.data) == 0 and len(wins_this_week.data) == 0:
            what_moved_message = "Pas de mouvement enregistré cette semaine. Parfois, tenir le coup est déjà une victoire."
        elif len(completed_tasks.data) == 0:
            what_moved_message = f"Cette semaine, tu as célébré {len(wins_this_week.data)} victoire(s). Même sans tâches cochées, tu avances."
        elif len(wins_this_week.data) == 0:
            what_moved_message = f"Tu as terminé {len(completed_tasks.data)} tâche(s). C'est bien. N'oublie pas de célébrer, même les petites choses."
        else:
            what_moved_message = f"{len(completed_tasks.data)} tâche(s) accomplie(s) et {len(wins_this_week.data)} victoire(s). Une belle semaine."
        
        # Message "what_stalled" - version humaine
        stalled_parts = []
        if len(overdue_docs.data) > 0:
            if len(overdue_docs.data) == 1:
                stalled_parts.append(f"un document traîne : {overdue_docs.data[0]['name'][:30]}")
            else:
                doc_names = ", ".join([d["name"][:20] for d in overdue_docs.data[:2]])
                if len(overdue_docs.data) > 2:
                    doc_names += f" et {len(overdue_docs.data)-2} autre(s)"
                stalled_parts.append(f"{len(overdue_docs.data)} documents en retard : {doc_names}")
        
        if len(overdue_tasks.data) > 0:
            if len(overdue_tasks.data) == 1:
                stalled_parts.append(f"une tâche en retard : {overdue_tasks.data[0]['title'][:30]}")
            else:
                stalled_parts.append(f"{len(overdue_tasks.data)} tâches en retard")
        
        # Missions inactives
        stalled_missions = []
        for m in active_missions.data:
            if not m.get("updated_at") or m["updated_at"] < (datetime.now() - timedelta(days=14)).isoformat():
                stalled_missions.append(m["name"])
        
        if stalled_missions:
            if len(stalled_missions) == 1:
                stalled_parts.append(f"une mission n'a pas bougé depuis 2 semaines : {stalled_missions[0]}")
            else:
                stalled_parts.append(f"{len(stalled_missions)} missions n'ont pas avancé récemment")
        
        if stalled_parts:
            what_stalled_message = "⚠️ " + ". ".join(stalled_parts) + "."
        else:
            what_stalled_message = "✅ Rien de bloqué cette semaine. Tout roule."
        
        # Message "pending_documents_summary" - version humaine
        if len(overdue_docs.data) == 0:
            if len(pending_docs.data) == 0:
                pending_summary = "📄 Aucun document en attente. Tu es à jour."
            elif len(pending_docs.data) == 1:
                pending_summary = f"📄 Un document en attente : {pending_docs.data[0]['name'][:30]}. Rien d'urgent."
            else:
                pending_summary = f"📄 {len(pending_docs.data)} documents en attente. À garder dans un coin de ta tête."
        elif len(overdue_docs.data) == 1:
            pending_summary = f"⚠️ Un document est en retard : {overdue_docs.data[0]['name'][:40]}. On s'en occupe cette semaine ?"
        else:
            doc_names = ", ".join([d["name"][:25] for d in overdue_docs.data[:2]])
            if len(overdue_docs.data) > 2:
                doc_names += f" et {len(overdue_docs.data)-2} autre(s)"
            pending_summary = f"📋 {len(overdue_docs.data)} documents en retard : {doc_names}. Un par un, ça va le faire."
        
        # Message "closest_to_cash" - version humaine
        import random
        if closest_to_cash:
            cash_messages = [
                f"💰 {closest_to_cash} — c'est ta mission la plus prometteuse côté argent. Une petite action cette semaine peut débloquer des choses. On y va ?",
                f"💰 {closest_to_cash} a du potentiel. Même une heure cette semaine peut faire bouger les choses.",
                f"🎯 {closest_to_cash} — si tu veux du cash rapidement, c'est par là. Un petit pas chaque jour."
            ]
            closest_summary = random.choice(cash_messages)
        else:
            potential_mission = None
            for mission in missions_by_revenue.data:
                if mission.get("revenue_potential", 0) > 3:
                    potential_mission = mission.get("name")
                    break
            
            if potential_mission:
                closest_summary = f"💭 {potential_mission} a du potentiel mais n'est pas active. Et si tu lui consacrais une petite heure cette semaine ? Rien que pour voir."
            else:
                closest_summary = "💭 Aucune mission n'est proche du cash pour l'instant. C'est le moment d'en identifier une. Veux-tu qu'on en parle ?"
        
        # ========== GÉNÉRATION DES PRIORITÉS PAR IA ==========
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": """Tu es Becks, l'amie et stratège de Rebecca. Tu connais sa vie :
- Maman de 4 filles (Neriah, Nylah, Norah, Sheyi)
- Projets : Ifè Farm, Love & Fire Sport, Santé Plus, Bénin Relocation
- Elle est souvent surchargée et a besoin de simplicité

Génère 3 priorités pour la semaine prochaine.
Sois NATURELLE, HUMAINE, CONCRÈTE. Pas de jargon corporate.
Écris comme tu parlerais à une amie.

Exemples de bonnes priorités :
- "Prendre 10 minutes pour toi, sans culpabilité"
- "Vider ta tête dans le Brain Dump, ça fait du bien"
- "Avancer sur un seul projet, pas tous à la fois"

Retourne UNIQUEMENT du JSON : {"priorities": ["action 1", "action 2", "action 3"]}"""},
                    {"role": "user", "content": f"""Voici le contexte de sa semaine :
- Missions actives : {[m["name"] for m in active_missions.data[:5]]}
- {len(pending_docs.data)} documents en attente
- {len(overdue_docs.data)} documents en retard
- {len(wins_this_week.data)} victoire(s) cette semaine
- Solde net : {total_revenue - total_spending:,.0f} CFA

Génère 3 priorités naturelles pour la semaine prochaine."""}
                ],
                max_tokens=300,
                temperature=0.8
            )
            result_text = response.choices[0].message.content
            result_text = result_text.replace("```json", "").replace("```", "").strip()
            priorities_result = json.loads(result_text)
            next_week_priorities = priorities_result.get("priorities", [])
            
            # Nettoyer et humaniser les priorités
            next_week_priorities = [p.strip() for p in next_week_priorities if p and len(p) > 5]
            
            # Fallbacks humains
            if len(next_week_priorities) < 3:
                fallbacks = [
                    "🌿 Prendre un moment pour toi, ça compte",
                    "📋 Une petite tâche, une seule, pour avancer",
                    "💬 Parler à Becks si tu as besoin de vider ta tête"
                ]
                for fb in fallbacks:
                    if fb not in next_week_priorities and len(next_week_priorities) < 3:
                        next_week_priorities.append(fb)
                
        except Exception as e:
            logger.error(f"Erreur génération priorités: {e}")
            next_week_priorities = [
                "🌿 Respirer. Une seule chose importante aujourd'hui.",
                "📋 Regarder ce qui traîne et choisir une seule action",
                f"💬 Parler à Becks de {closest_to_cash or 'ce qui te pèse'}"
            ]
        
        # ========== INSIGHT GLOBAL ==========
        insight = _generate_human_weekly_insight(
            completion_rate, 
            len(wins_this_week.data), 
            total_revenue - total_spending, 
            len(overdue_docs.data),
            len(completed_tasks.data),
            len(new_tasks.data),
            len(overdue_tasks.data)
        )
        
        return {
            "success": True,
            "week_range": {
                "start": start_of_week,
                "end": end_of_week
            },
            "summary": {
                "tasks_completed": len(completed_tasks.data),
                "tasks_created": len(new_tasks.data),
                "completion_rate": completion_rate,
                "wins": len(wins_this_week.data),
                "total_spending": total_spending,
                "total_revenue": total_revenue,
                "net_balance": total_revenue - total_spending
            },
            "what_moved": {
                "message": what_moved_message,
                "completed_tasks": [{"title": t["title"], "project": t.get("project", "Général")} for t in completed_tasks.data[:5]],
                "wins": [{"title": w["title"], "celebration_emoji": w.get("celebration_emoji", "🎉")} for w in wins_this_week.data[:5]]
            },
            "what_stalled": {
                "message": what_stalled_message,
                "overdue_docs": [{"name": d["name"], "due_date": d["due_date"]} for d in overdue_docs.data[:5]],
                "overdue_tasks_count": len(overdue_tasks.data),
                "pending_docs_count": len(pending_docs.data),
                "stalled_missions": [{"name": m["name"]} for m in active_missions.data if not m.get("updated_at") or m["updated_at"] < (datetime.now() - timedelta(days=14)).isoformat()][:3]
            },
            "closest_to_cash": {
                "name": closest_to_cash or "Aucune pour l'instant",
                "message": closest_summary
            },
            "pending_documents_summary": pending_summary,
            "next_week_priorities": next_week_priorities,
            "mood_summary": [{"date": m["date"], "mood": m["mood"]} for m in moods_this_week.data],
            "insight": insight
        }
        
    except Exception as e:
        logger.error(f"Erreur weekly_ceo: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# CONTENT CALENDAR
# =====================================================

@app.get("/api/content/calendar")
async def get_content_calendar(month: int = None, year: int = None):
    """
    Retourne le calendrier éditorial avec les contenus programmés.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        now = datetime.now()
        target_month = month if month else now.month
        target_year = year if year else now.year
        
        start_date = datetime(target_year, target_month, 1).date()
        if target_month == 12:
            end_date = datetime(target_year + 1, 1, 1).date() - timedelta(days=1)
        else:
            end_date = datetime(target_year, target_month + 1, 1).date() - timedelta(days=1)
        
        contents = supabase.table("content").select("*")\
            .gte("publish_date", start_date.isoformat())\
            .lte("publish_date", end_date.isoformat())\
            .execute()
        
        calendar_data = {}
        for content in contents.data:
            date = content.get("publish_date")
            if date:
                if date not in calendar_data:
                    calendar_data[date] = []
                calendar_data[date].append(content)
        
        suggestions = []
        current_date = start_date
        while current_date <= end_date:
            date_str = current_date.isoformat()
            if date_str not in calendar_data and current_date.weekday() < 5:
                suggestions.append({
                    "date": date_str,
                    "suggested_platform": "instagram",
                    "suggested_type": "story",
                    "suggested_theme": _suggest_content_theme()
                })
            current_date += timedelta(days=1)
        
        return {
            "success": True,
            "year": target_year,
            "month": target_month,
            "calendar": calendar_data,
            "suggestions": suggestions[:10],
            "stats": {
                "total": len(contents.data),
                "scheduled": len([c for c in contents.data if c.get("status") == "scheduled"]),
                "published": len([c for c in contents.data if c.get("status") == "posted"]),
                "draft": len([c for c in contents.data if c.get("status") == "draft"])
            }
        }
        
    except Exception as e:
        logger.error(f"Erreur content_calendar: {e}")
        return {"success": False, "error": str(e)}


def _suggest_content_theme() -> str:
    """Génère une suggestion de thème de contenu aléatoire mais pertinente"""
    themes = [
        "🌱 Vie à la ferme", "💪 Sport adapté", "👩‍👧‍👦 Vie de maman",
        "💰 Opportunités business", "🌟 Victoire personnelle", "📚 Éducation",
        "🏠 Relocalisation", "❤️ Santé et bien-être", "🎯 Objectif de la semaine",
        "🙏 Gratitude", "🚀 Lancement à venir", "📖 Témoignage"
    ]
    import random
    return random.choice(themes)


@app.post("/api/content/generate-idea")
async def generate_content_idea(request: Dict[str, Any]):
    """
    Génère une idée de contenu avec l'IA.
    """
    platform = request.get("platform", "instagram")
    topic = request.get("topic", "")
    audience = request.get("audience", "")
    
    try:
        prompt = f"""Génère une idée de contenu pour {platform}.
Sujet: {topic if topic else 'général'}
Audience: {audience if audience else 'femmes entrepreneures, mamans'}

Retourne UNIQUEMENT du JSON:
{{
  "title": "titre accrocheur",
  "hook": "phrase d'accroche (max 100 caractères)",
  "content_type": "story/reel/carousel/post",
  "hashtags": ["#hashtag1", "#hashtag2"],
  "best_time": "meilleur horaire de publication"
}}"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.8
        )
        
        result_text = response.choices[0].message.content
        result_text = result_text.replace("```json", "").replace("```", "").strip()
        idea = json.loads(result_text)
        
        return {"success": True, "idea": idea}
        
    except Exception as e:
        logger.error(f"Erreur generate_content_idea: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# OPPORTUNITY SCANNER
# =====================================================

@app.post("/api/opportunities/scan")
async def scan_opportunities(request: Dict[str, Any] = None):
    """
    Scan toutes les sources pour détecter des opportunités.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        opportunities_found = []
        
        inbox_items = supabase.table("inbox").select("*").eq("needs_processing", False).limit(50).execute()
        missions = supabase.table("missions").select("name, notes").execute()
        conversations = supabase.table("conversation_messages").select("content").limit(30).execute()
        
        full_text = ""
        for item in inbox_items.data:
            full_text += item.get("content", "") + "\n"
        for mission in missions.data:
            full_text += mission.get("notes", "") + "\n"
        for conv in conversations.data:
            try:
                parsed = json.loads(conv.get("content", "{}"))
                full_text += parsed.get("content", "") + "\n"
            except:
                full_text += conv.get("content", "") + "\n"
        
        full_text = full_text[:8000]
        
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": """Tu es Becks, l'opportunity scanner. Analyse le texte et détecte les opportunités.

Une opportunité peut être :
- Grant / subvention
- Contrat public (DDA, eMMA, SAM.gov)
- Partenariat
- Client potentiel
- Collaboration
- Appel à projet
- Événement important

Retourne UNIQUEMENT du JSON avec cette structure :
{
  "opportunities": [
    {
      "title": "titre de l'opportunité",
      "type": "grant/contract/partnership/client/collaboration/event",
      "source": "d'où ça vient",
      "deadline": "date si connue",
      "value_estimate": "estimation en CFA",
      "confidence": "high/medium/low",
      "next_action": "action recommandée"
    }
  ]
}

Si aucune opportunité, retourne {"opportunities": []}"""},
                    {"role": "user", "content": f"Analyse ce texte :\n\n{full_text}"}
                ],
                max_tokens=1000,
                temperature=0.3
            )
            
            result_text = response.choices[0].message.content
            result_text = result_text.replace("```json", "").replace("```", "").strip()
            result = json.loads(result_text)
            
            opportunities_found = result.get("opportunities", [])
            
            for opp in opportunities_found:
                existing = supabase.table("opportunities").select("*")\
                    .ilike("title", f"%{opp.get('title', '')}%")\
                    .execute()
                
                if not existing.data:
                    supabase.table("opportunities").insert({
                        "title": opp.get("title"),
                        "type": opp.get("type"),
                        "stage": "idea",
                        "estimated_value": opp.get("value_estimate", 0),
                        "probability": opp.get("confidence", "medium"),
                        "notes": f"🔍 Détecté par scanner le {datetime.now().strftime('%d/%m/%Y')}\nSource: {opp.get('source', 'Inconnue')}\nAction recommandée: {opp.get('next_action', 'À définir')}",
                        "next_action": opp.get("next_action"),
                        "user_id": user_id
                    }).execute()
            
            return {
                "success": True,
                "opportunities_found": opportunities_found,
                "count": len(opportunities_found),
                "scanned_sources": {
                    "inbox": len(inbox_items.data),
                    "missions": len(missions.data),
                    "conversations": len(conversations.data)
                }
            }
            
        except Exception as e:
            logger.error(f"Erreur scan IA: {e}")
            return {"success": False, "error": str(e)}
        
    except Exception as e:
        logger.error(f"Erreur scan_opportunities: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# READY-TO-SEND GENERATION
# =====================================================

@app.post("/api/generate/ready-to-send")
async def generate_ready_to_send(request: Dict[str, Any]):
    """
    Génère un document prêt à être copié/envoyé.
    """
    doc_type = request.get("type", "email")
    context = request.get("context", "")
    tone = request.get("tone", "professional")
    recipient = request.get("recipient", "")
    sender = request.get("sender", "Rebecca")
    
    if not context:
        return {"success": False, "error": "Contexte requis"}
    user_id = get_request_user_id(request or {})
    
    memory_context = await get_user_memory_context(user_id)
    profile_context = await get_profile_context(user_id=user_id)
    
    prompt = f"""Tu es Becks, l'assistante de Rebecca. Génère un {doc_type} prêt à être copié et envoyé.

Contexte : {context}
Ton : {tone}
Destinataire : {recipient if recipient else 'Non spécifié'}
Expéditeur : {sender}

Informations sur Rebecca : {profile_context}
Ce que Becks sait d'elle : {memory_context}

RÈGLES IMPORTANTES :
1. Sois clair, concis, professionnel
2. Structure le document proprement
3. Si c'est un email, ajoute un objet pertinent
4. Si des informations manquent, laisse des placeholders [entre crochets]
5. Retourne UNIQUEMENT du JSON valide avec cette structure :

{{
  "subject": "objet (si applicable)",
  "body": "corps du message",
  "signature": "signature",
  "full_text": "texte complet prêt à copier"
}}

Ne retourne que le JSON, rien d'autre."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.5
        )
        
        result_text = response.choices[0].message.content
        result_text = result_text.replace("```json", "").replace("```", "").strip()
        generated = json.loads(result_text)
        
        if supabase:
            supabase.table("drafts").insert({
                "type": doc_type,
                "content": generated.get("full_text", generated.get("body", "")),
                "context": context,
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }).execute()
        
        return {"success": True, "generated": generated}
        
    except Exception as e:
        logger.error(f"Erreur generation: {e}")
        return {"success": False, "error": str(e)}


# =====================================================
# DECISION MODE
# =====================================================

@app.post("/api/decide/compare")
async def compare_options(request: Dict[str, Any]):
    """
    Compare deux options (A vs B) et aide à prendre une décision.
    """
    option_a = request.get("option_a", "")
    option_b = request.get("option_b", "")
    context = request.get("context", "")
    criteria_weights = request.get("criteria_weights", {
        "revenue_speed": 5,
        "strategic_value": 4,
        "effort": 3,
        "emotional_cost": 3,
        "urgency": 4
    })
    
    if not option_a or not option_b:
        return {"success": False, "error": "Deux options sont requises"}
    user_id = get_request_user_id(request or {})
    
    memory_context = await get_user_memory_context(user_id)
    profile_context = await get_profile_context(user_id=user_id)
    
    prompt = f"""Tu es Becks, conseillère stratégique. Aide Rebecca à choisir entre deux options.

Option A : {option_a}
Option B : {option_b}
Contexte : {context if context else "Aucun contexte spécifique"}

Critères de décision (poids de 1 à 5) :
- Revenu potentiel (rapidité) : {criteria_weights.get('revenue_speed', 5)}/5
- Valeur stratégique long terme : {criteria_weights.get('strategic_value', 4)}/5
- Effort / difficulté : {criteria_weights.get('effort', 3)}/5 (plus bas = mieux)
- Coût émotionnel : {criteria_weights.get('emotional_cost', 3)}/5 (plus bas = mieux)
- Urgence : {criteria_weights.get('urgency', 4)}/5

Retourne UNIQUEMENT du JSON avec cette structure :

{{
  "analysis": {{
    "option_a": {{
      "pros": ["avantage 1", "avantage 2"],
      "cons": ["inconvénient 1", "inconvénient 2"],
      "score": 0,
      "scores_detail": {{
        "revenue_speed": 0,
        "strategic_value": 0,
        "effort": 0,
        "emotional_cost": 0,
        "urgency": 0
      }}
    }},
    "option_b": {{
      "pros": ["avantage 1", "avantage 2"],
      "cons": ["inconvénient 1", "inconvénient 2"],
      "score": 0,
      "scores_detail": {{
        "revenue_speed": 0,
        "strategic_value": 0,
        "effort": 0,
        "emotional_cost": 0,
        "urgency": 0
      }}
    }}
  }},
  "recommendation": "option_a" ou "option_b",
  "recommendation_reason": "Pourquoi cette option est meilleure",
  "next_action": "Action concrète à prendre maintenant"
}}

Note : Pour chaque critère, note de 1 à 5. Le score total est la somme des notes pondérées par les poids."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.3
        )
        
        result_text = response.choices[0].message.content
        result_text = result_text.replace("```json", "").replace("```", "").strip()
        comparison = json.loads(result_text)
        
        if supabase:
            supabase.table("user_memory").insert({
                "category": "decisions",
                "key": f"decision_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                "value": json.dumps({
                    "option_a": option_a,
                    "option_b": option_b,
                    "recommendation": comparison.get("recommendation"),
                    "context": context
                }),
                "user_id": user_id,
                "created_at": datetime.now().isoformat()
            }).execute()
        
        return {"success": True, "comparison": comparison}
        
    except Exception as e:
        logger.error(f"Erreur compare: {e}")
        return {"success": False, "error": str(e)}



# =====================================================
# TEXT-TO-SPEECH AVEC ELEVENLABS (NOUVELLE API)
# =====================================================

ELEVENLABS_VOICE_ID = os.environ.get("ELEVENLABS_VOICE_ID", "uju3wxzG5OhpWcoi3SMy")  # Voix Bella


# Voix recommandées pour une voix naturelle et chaleureuse
VOICES = {
    "bella": "uju3wxzG5OhpWcoi3SMy",      # Douce, chaleureuse
    "rachel": "AeRdCCKzvd23BpJoofzx",    # Naturelle, professionnelle
    "grace": "nzFihrBIvB34imQBuxub",     # Élégante, claire
    "domi": "DODLEQrClDo8wCz460ld",      # Jeune, énergique
    "antoni": "dXtC3XhB9GtPusIpNtQx",    # Voix masculine douce
}
DEFAULT_VOICE_ID = VOICES["bella"]

@app.post("/api/tts/speak")
async def text_to_speech(request: Dict[str, Any]):
    """Convertit un texte en audio avec ElevenLabs (nouvelle API)"""
    text = request.get("text", "")
    voice_id = request.get("voice_id", DEFAULT_VOICE_ID)
    
    if not text:
        return {"success": False, "error": "Texte requis"}
    
    if not ELEVENLABS_API_KEY:
        return {"success": False, "error": "ElevenLabs non configuré", "fallback": True}
    
    try:
        # Nettoyer le texte des balises et caractères spéciaux
        clean_text = re.sub(r'\[ACTION:[^\]]*\]', '', text)
        clean_text = re.sub(r'\*\*.*?\*\*', '', clean_text)
        clean_text = re.sub(r'[✅🎯✨⚠️📋🎉]', '', clean_text)
        clean_text = ' '.join(clean_text.split())
        
        # Limiter la longueur (10k caractères max pour le plan gratuit)
        if len(clean_text) > 5000:
            clean_text = clean_text[:5000]
        
        # Utiliser la nouvelle API ElevenLabs
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}",
                headers={
                    "xi-api-key": ELEVENLABS_API_KEY,
                    "Content-Type": "application/json",
                    "Accept": "audio/mpeg"
                },
                json={
                    "text": clean_text,
                    "model_id": "eleven_multilingual_v2",  # Meilleure qualité
                    "voice_settings": {
                        "stability": 0.35,      # Plus d'émotion
                        "similarity_boost": 0.75,
                        "style": 0.2,
                        "use_speaker_boost": True
                    }
                }
            )
            
            if response.status_code == 200:
                import base64
                audio_base64 = base64.b64encode(response.content).decode('utf-8')
                return {
                    "success": True, 
                    "audio": audio_base64, 
                    "format": "mp3",
                    "voice_id": voice_id
                }
            else:
                error_msg = response.text
                logger.error(f"Erreur ElevenLabs: {error_msg}")
                
                # Si erreur de quota, fallback vers Edge TTS
                if "quota" in error_msg.lower() or "unusual_activity" in error_msg.lower():
                    return await edge_tts_fallback(clean_text)
                
                return {"success": False, "error": error_msg, "fallback": True}
                
    except Exception as e:
        logger.error(f"Erreur TTS ElevenLabs: {e}")
        # Fallback vers Edge TTS
        return await edge_tts_fallback(clean_text if 'clean_text' in locals() else text)

async def edge_tts_fallback(text: str):
    """Fallback vers Edge TTS (gratuit, illimité)"""
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            # Encoder le texte pour l'URL
            import urllib.parse
            encoded_text = urllib.parse.quote(text)
            
            response = await client.get(
                f"https://speech.platform.bing.com/consumer/speech/synthesize/readaloud/edge/v1?TrustedClientToken=6A5AA1D4EAFF4E9FB37E23D68491D6F4&Voice=fr-FR-DeniseNeural&Text={encoded_text}",
                follow_redirects=True
            )
            
            if response.status_code == 200:
                import base64
                audio_base64 = base64.b64encode(response.content).decode('utf-8')
                return {
                    "success": True,
                    "audio": audio_base64,
                    "format": "mp3",
                    "fallback": True
                }
    except Exception as e:
        logger.error(f"Erreur Edge TTS fallback: {e}")
    
    return {"success": False, "error": "Aucune voix disponible", "fallback": True}

# =====================================================
# TEXT-TO-SPEECH AVEC DEEPGRAM
# =====================================================

DEEPGRAM_API_KEY = os.environ.get("DEEPGRAM_API_KEY")

@app.post("/api/tts/deepgram")
async def deepgram_speak(request: Dict[str, Any]):
    """Convertit un texte en audio avec Deepgram"""
    text = request.get("text", "")
    voice = request.get("voice", "aura-2-athena-en")
    
    if not text:
        return {"success": False, "error": "Texte requis"}
    
    if not DEEPGRAM_API_KEY:
        return {"success": False, "error": "Deepgram non configuré"}
    
    try:
        clean_text = re.sub(r'\[ACTION:[^\]]*\]', '', text)
        clean_text = re.sub(r'\*\*.*?\*\*', '', clean_text)
        clean_text = re.sub(r'[✅🎯✨⚠️📋🎉]', '', clean_text)
        clean_text = ' '.join(clean_text.split())
        
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"https://api.deepgram.com/v1/speak?model={voice}",
                headers={
                    "Authorization": f"Token {DEEPGRAM_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={"text": clean_text}
            )
            
            if response.status_code == 200:
                # Retourner l'audio en base64 (JSON valide)
                import base64
                audio_base64 = base64.b64encode(response.content).decode('utf-8')
                return {"success": True, "audio": audio_base64, "format": "mp3"}
            else:
                return {"success": False, "error": response.text}
                
    except Exception as e:
        logger.error(f"Erreur Deepgram: {e}")
        return {"success": False, "error": str(e)}
# =====================================================
# API ROUTES - GENERIC CRUD (EXISTANT)
# =====================================================
@app.get("/{table}")
def get_table(table: str, limit: int = 100):
    if table not in AVAILABLE_TABLES:
        raise HTTPException(status_code=404, detail=f"Table '{table}' non trouvée")
    return db_query(table, limit=limit)


@app.post("/{table}")
async def create_item(table: str, request: WriteRequest):
    if table not in AVAILABLE_TABLES:
        raise HTTPException(status_code=404, detail=f"Table '{table}' non trouvée")
    result = await db_insert(table, request.data)
    return result

@app.put("/{table}/{item_id}")
def update_item(table: str, item_id: str, request: UpdateRequest):
    if table not in AVAILABLE_TABLES:
        raise HTTPException(status_code=404, detail=f"Table '{table}' non trouvée")
    return db_update(table, item_id, request.data)


@app.delete("/{table}/{item_id}")
def delete_item(table: str, item_id: str):
    if table not in AVAILABLE_TABLES:
        raise HTTPException(status_code=404, detail=f"Table '{table}' non trouvée")
    return db_delete(table, item_id)

@app.post("/api/generate-greeting")
async def generate_greeting(request: Dict[str, Any]):
    tasks_count = request.get("tasks_count", 0)
    overdue_count = request.get("overdue_count", 0)
    wins_count = request.get("wins_count", 0)
    missions_count = request.get("missions_count", 0)
    mood = request.get("mood")
    hour = request.get("hour", 12)
    
    # Déterminer le moment de la journée
    if hour < 12:
        time_context = "matin"
        emoji = "☀️"
    elif hour < 18:
        time_context = "après-midi"
        emoji = "🌤️"
    else:
        time_context = "soir"
        emoji = "🌙"
    
    # Adapter selon l'humeur
    mood_context = ""
    if mood == "fatiguée":
        mood_context = "Elle est fatiguée. Sois douce et propose une micro-pause."
    elif mood == "stressée":
        mood_context = "Elle est stressée. Propose de respirer et de prioriser UNE seule chose."
    elif mood == "excellent":
        mood_context = "Elle est en pleine forme ! Propose d'attaquer une tâche importante."
    elif mood == "bien":
        mood_context = "Elle va bien. Propose une action équilibrée."
    
    prompt = f"""Génère un message d'accueil pour Rebecca, comme si tu étais son amie et assistante personnelle.

CONTEXTE :
- Moment : {time_context} {emoji}
- Tâches aujourd'hui : {tasks_count}
- Tâches en retard : {overdue_count}
- Victoires récentes : {wins_count}
- Missions actives : {missions_count}
- {mood_context}

STYLE : 
- Humain, chaleureux, naturel (pas robotique)
- Parle à la première personne ("Je vois que...", "Je sens que...")
- Maximum 40 mots
- Structure ton message ainsi :
  1. Une phrase sur son état ou le moment
  2. Un constat sur sa charge (tâches/missions)
  3. Une proposition d'action simple ET une question ouverte

EXEMPLES de bons messages :
- "☀️ Bonjour Rebecca. Je vois que tu as 3 tâches aujourd'hui, dont une en retard. On commence par celle-là ? Je suis là. 💖"
- "🌤️ Salut. 2 missions actives et 1 victoire récente, bien joué ! Tu veux qu'on avance sur laquelle ?"
- "🌙 Bonsoir. Tu es fatiguée ? Ralentissons. Une seule petite chose pour ce soir ?"

Retourne UNIQUEMENT le message, rien d'autre."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=150
        )
        
        greeting = response.choices[0].message.content.strip()
        return {"success": True, "greeting": greeting}
        
    except Exception as e:
        logger.error(f"Erreur génération greeting: {e}")
        # Fallback humain
        fallback = f"{emoji} Salut Rebecca. {tasks_count} tâche(s) aujourd'hui. On y va doucement. 👑"
        return {"success": True, "greeting": fallback}


@app.post("/api/weekly-summary")
async def send_weekly_summary(request: Dict[str, Any] = None):
    """
    Envoie un résumé hebdomadaire personnalisé par email et notification push.
    À appeler par cron-job.org tous les dimanches à 19h.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    # Vérifier si on est dimanche (optionnel, la cron s'en charge)
    if datetime.now().weekday() != 6:
        return {"success": True, "sent": False, "message": "Pas dimanche, résumé non envoyé"}

    
    try:
        user_id = get_request_user_id(request or {})
        now = datetime.now()
        start_of_week = (now.date() - timedelta(days=now.weekday())).isoformat()
        end_of_week = now.date().isoformat()
        
        # ========== RÉCUPÉRER LES DONNÉES DE LA SEMAINE ==========
        
        # 1. Tâches complétées
        completed_tasks = supabase.table("tasks").select("*")\
            .gte("updated_at", start_of_week)\
            .eq("status", "done")\
            .eq("user_id", user_id)\
            .execute()
        completed_count = len(completed_tasks.data)
        completed_list = [t["title"] for t in completed_tasks.data[:5]]
        
        # 2. Tâches créées
        new_tasks = supabase.table("tasks").select("*")\
            .eq("user_id", user_id)\
            .gte("created_at", start_of_week)\
            .execute()
        new_count = len(new_tasks.data)
        
        # 3. Taux de complétion
        completion_rate = int((completed_count / new_count) * 100) if new_count > 0 else 0
        
        # 4. Tâches en retard
        overdue_tasks = supabase.table("tasks").select("*")\
            .eq("user_id", user_id)\
            .lt("due_date", now.date().isoformat())\
            .neq("status", "done")\
            .execute()
        overdue_count = len(overdue_tasks.data)
        
        # 5. Victoires de la semaine
        wins = supabase.table("wins").select("*")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        wins_count = len(wins.data)
        wins_list = [w["title"] for w in wins.data[:3]]
        
        # 6. Dépenses et revenus
        spending = supabase.table("spending").select("amount")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        total_spending = sum(s.get("amount", 0) for s in spending.data)
        
        revenue = supabase.table("revenue").select("amount")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        total_revenue = sum(r.get("amount", 0) for r in revenue.data)
        balance = total_revenue - total_spending
        
        # 7. Humeurs de la semaine
        moods = supabase.table("mood_entries").select("mood")\
            .eq("user_id", user_id)\
            .gte("date", start_of_week)\
            .execute()
        mood_counts = {}
        for m in moods.data:
            mood_counts[m["mood"]] = mood_counts.get(m["mood"], 0) + 1
        
        # 8. Missions actives
        active_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").execute()
        
        # 9. Récupérer le nom
        profile = supabase.table("user_profile").select("preferred_name").eq("user_id", user_id).execute()
        user_name = profile.data[0].get("preferred_name", "Rebecca") if profile.data else "Rebecca"
        
        # ========== GÉNÉRATION IA ==========
        
        prompt = f"""Génère un résumé hebdomadaire chaleureux et encourageant pour Rebecca.

DONNÉES DE LA SEMAINE :
- Tâches terminées : {completed_count}
- Tâches créées : {new_count}
- Taux de complétion : {completion_rate}%
- Tâches en retard : {overdue_count}
- Victoires célébrées : {wins_count}
- Dépenses : {total_spending:,.0f} CFA
- Revenus : {total_revenue:,.0f} CFA
- Solde : {balance:,.0f} CFA
- Humeurs : {mood_counts}
- Missions actives : {len(active_missions.data)}

RÈGLES :
- Maximum 100 mots
- Commence par "Bonsoir {user_name},"
- Mentionne les points forts (ce qu'elle a accompli)
- Mentionne doucement ce qui peut être amélioré
- Termine par un encouragement pour la semaine à venir
- Soit naturelle, pas corporate

Retourne UNIQUEMENT le message, rien d'autre."""

        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7,
                max_tokens=300
            )
            message = response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Erreur IA weekly summary: {e}")
            # Fallback humain
            message = f"""Bonsoir {user_name},

Cette semaine, tu as terminé {completed_count} tâche(s) sur {new_count} créées ({completion_rate}%).
Tu as célébré {wins_count} victoire(s) ! 👑

💰 Finances : {total_revenue:,.0f} CFA de revenus, {total_spending:,.0f} CFA de dépenses.

{len(active_missions.data)} mission(s) active(s) en cours.

La semaine prochaine, concentre-toi sur l'essentiel. Une chose à la fois. Tu as tout ce qu'il faut. 👑

Passe une bonne soirée et repose-toi bien. 💖"""
        
        # ========== ENVOI ==========
        
        # Notification push
        push_sent = False
        try:
            send_notification_sync({
                "title": "📊 Bilan de la semaine",
                "body": f"{completed_count} tâches faites • {wins_count} victoires",
                "url": "/vision-strategy?tab=weekly",
                "user_id": user_id,
                "type": "report"
            })
            push_sent = True
        except Exception as e:
            logger.error(f"Erreur envoi push weekly: {e}")
        
        # Email
        email_sent = False
        if BREVO_API_KEY:
            try:
                email_body = message.replace("\n", "<br>")
                await send_email(EmailRequest(
                    to="jbillcataria@gmail.com",
                    subject=f"📊 Bilan hebdomadaire - {now.strftime('%d/%m')} au {end_of_week}",
                    body=email_body
                ))
                email_sent = True
                logger.info("📧 Email bilan hebdomadaire envoyé")
            except Exception as e:
                logger.error(f"Erreur envoi email weekly: {e}")
        
        # Sauvegarder dans la base pour historique
        try:
            supabase.table("weekly_summaries").insert({
                "week_start": start_of_week,
                "week_end": end_of_week,
                "tasks_completed": completed_count,
                "tasks_created": new_count,
                "completion_rate": completion_rate,
                "wins_count": wins_count,
                "total_spending": total_spending,
                "total_revenue": total_revenue,
                "message": message,
                "sent_at": now.isoformat()
            }).execute()
        except Exception as e:
            logger.error(f"Erreur sauvegarde weekly summary: {e}")
        
        return {
            "success": True,
            "message": message,
            "stats": {
                "tasks_completed": completed_count,
                "tasks_created": new_count,
                "completion_rate": completion_rate,
                "overdue_tasks": overdue_count,
                "wins_count": wins_count,
                "total_spending": total_spending,
                "total_revenue": total_revenue,
                "balance": balance
            },
            "push_sent": push_sent,
            "email_sent": email_sent
        }
        
    except Exception as e:
        logger.error(f"Erreur weekly summary: {e}")
        return {"success": False, "error": str(e)}



@app.post("/api/chat/generate-title")
async def generate_chat_title(request: Dict[str, Any]):
    """
    Génère un titre dynamique pour une conversation à partir du premier message.
    """
    first_message = request.get("first_message", "")
    
    if not first_message or len(first_message) < 5:
        return {"success": True, "title": f"Conversation du {datetime.now().strftime('%d/%m/%Y')}"}
    
    prompt = f"""Génère un titre COURT (max 6 mots) pour cette conversation.

Message : "{first_message}"

RÈGLES :
- Soit descriptif et pertinent
- Pas de date dans le titre
- Pas de "Nouvelle conversation"
- Exemples : "Planning ferme", "Dossier relocation", "Budget mois", "Idées content"

Retourne UNIQUEMENT le titre, rien d'autre."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=30
        )
        
        title = response.choices[0].message.content.strip()
        # Nettoyer les guillemets éventuels
        title = title.replace('"', '').replace("'", "")
        
        return {"success": True, "title": title}
        
    except Exception as e:
        logger.error(f"Erreur génération titre: {e}")
        # Fallback
        return {"success": True, "title": f"Conversation du {datetime.now().strftime('%d/%m/%Y')}"}



@app.post("/api/notifications/smart-group")
async def send_smart_notifications(request: Dict[str, Any] = None):
    """
    Analyse les notifications en attente et les regroupe intelligemment.
    À appeler par cron-job.org toutes les heures.
    """
    if not supabase:
        return {"success": False, "error": "Supabase non configuré"}
    
    try:
        user_id = get_request_user_id(request or {})
        today = datetime.now().date().isoformat()
        now = datetime.now()
        
        # Vérifier si déjà envoyé dans la dernière heure
        last_check = supabase.table("notifications_log").select("sent_at")\
            .eq("type", "smart_group")\
            .eq("user_id", user_id)\
            .order("sent_at", desc=True)\
            .limit(1)\
            .execute()
        
        if last_check.data:
            last_sent = datetime.fromisoformat(last_check.data[0]["sent_at"])
            if (now - last_sent).total_seconds() < 3600:  # 1 heure
                return {"success": True, "sent": False, "message": "Déjà envoyé récemment"}
        
        # ========== RÉCUPÉRER LES NOTIFICATIONS POTENTIELLES ==========
        
        # 1. Tâches du jour
        tasks_today = supabase.table("tasks").select("*").eq("user_id", user_id).eq("due_date", today).neq("status", "done").execute()
        tasks_today_count = len(tasks_today.data)
        tasks_today_list = [t["title"] for t in tasks_today.data[:3]]
        
        # 2. Tâches en retard
        overdue_tasks = supabase.table("tasks").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "done").execute()
        overdue_count = len(overdue_tasks.data)
        
        # 3. Documents en retard ou proches
        overdue_docs = supabase.table("documents").select("*").eq("user_id", user_id).lt("due_date", today).neq("status", "approved").execute()
        expiring_docs = supabase.table("documents").select("*").eq("user_id", user_id).gte("due_date", today).lte("due_date", (datetime.now().date() + timedelta(days=3)).isoformat()).neq("status", "approved").execute()
        
        # 4. Missions inactives (5+ jours)
        five_days_ago = (datetime.now() - timedelta(days=5)).isoformat()
        stale_missions = supabase.table("missions").select("*").eq("user_id", user_id).eq("status", "active").lt("updated_at", five_days_ago).execute()
        
        # 5. Victoires récentes (7 jours)
        week_ago = (datetime.now().date() - timedelta(days=7)).isoformat()
        recent_wins = supabase.table("wins").select("*").eq("user_id", user_id).gte("date", week_ago).execute()
        
        # 6. Événements familiaux aujourd'hui
        family_today = supabase.table("family_events").select("*").eq("user_id", user_id).eq("date", today).neq("status", "done").execute()
        
        # ========== CONSTRUIRE LE MESSAGE REGROUPÉ ==========
        
        notifications = []
        
        # Tâches
        if overdue_count > 0:
            notifications.append(f"⚠️ {overdue_count} tâche(s) en retard")
        elif tasks_today_count > 0:
            task_list = ", ".join(tasks_today_list[:2])
            if len(tasks_today_list) > 2:
                task_list += f" et {tasks_today_count - 2} autre(s)"
            notifications.append(f"📋 {tasks_today_count} tâche(s) aujourd'hui : {task_list}")
        
        # Documents
        if len(overdue_docs.data) > 0:
            notifications.append(f"📄 {len(overdue_docs.data)} document(s) en retard")
        elif len(expiring_docs.data) > 0:
            notifications.append(f"📄 {len(expiring_docs.data)} document(s) bientôt dû(s)")
        
        # Missions
        if len(stale_missions.data) > 0:
            missions_names = ", ".join([m["name"] for m in stale_missions.data[:2]])
            if len(stale_missions.data) > 2:
                missions_names += f" et {len(stale_missions.data) - 2} autre(s)"
            notifications.append(f"🎯 {len(stale_missions.data)} mission(s) sans activité récente : {missions_names}")
        
        # Victoires
        if len(recent_wins.data) == 0:
            notifications.append(f"🏆 Pas de victoire cette semaine. C'est le moment d'en célébrer une !")
        elif len(recent_wins.data) > 0:
            notifications.append(f"🏆 {len(recent_wins.data)} victoire(s) récente(s) ! Continue comme ça.")
        
        # Famille
        if len(family_today.data) > 0:
            events_names = ", ".join([e["title"] for e in family_today.data[:2]])
            notifications.append(f"👨‍👩‍👧‍👦 Aujourd'hui : {events_names}")
        
        # Si rien à signaler
        if not notifications:
            return {"success": True, "sent": False, "message": "Rien à signaler"}
        
        # ========== CRÉER LE MESSAGE FINAL ==========
        
        # Récupérer le nom
        profile = supabase.table("user_profile").select("preferred_name").eq("user_id", user_id).execute()
        user_name = profile.data[0].get("preferred_name", "Rebecca") if profile.data else "Rebecca"
        
        # Heure pour la salutation
        hour = now.hour
        if hour < 12:
            greeting = "☀️ Bonjour"
        elif hour < 18:
            greeting = "🌤️ Bon après-midi"
        else:
            greeting = "🌙 Bonsoir"
        
        body = "\n".join(notifications)
        
        message = f"""{greeting} {user_name},

Voici ce qui t'attend :

{body}

Becks est là si tu as besoin. 👑"""
        
        # ========== ENVOI D'UNE SEULE NOTIFICATION ==========
        
        # Envoyer une seule notification push
        send_notification_sync({
            "title": f"📋 {len(notifications)} point(s) important(s)",
            "body": notifications[0] + (" • ..." if len(notifications) > 1 else ""),
            "url": "/",
            "user_id": user_id,
            "type": "smart_group"
        })
        
        # Logger l'envoi
        supabase.table("notifications_log").insert({
            "type": "smart_group",
            "date": today,
            "user_id": user_id,
            "sent_at": now.isoformat(),
            "metadata": {
                "notifications_count": len(notifications),
                "categories": [n.split(" ")[0] for n in notifications]
            }
        }).execute()
        
        return {
            "success": True,
            "sent": True,
            "message": message,
            "notifications_count": len(notifications),
            "details": notifications
        }
        
    except Exception as e:
        logger.error(f"Erreur smart group notifications: {e}")
        return {"success": False, "error": str(e)}



@app.post("/api/welcome-message")
async def generate_welcome_message(request: Dict[str, Any]):
    """
    Génère un message de bienvenue personnalisé après connexion.
    """
    user_name = request.get("user_name", "Rebecca")
    hour = request.get("hour", datetime.now().hour)
    last_visit_days = request.get("last_visit_days", 0)
    
    # Déterminer le moment de la journée
    if hour < 12:
        time_greeting = "Bonjour"
        emoji = "☀️"
    elif hour < 18:
        time_greeting = "Bon après-midi"
        emoji = "🌤️"
    else:
        time_greeting = "Bonsoir"
        emoji = "🌙"
    
    # Message selon la dernière visite
    if last_visit_days == 0:
        visit_context = "contenu de revenir"
    elif last_visit_days == 1:
        visit_context = "ravi de te revoir aujourd'hui"
    elif last_visit_days <= 3:
        visit_context = f"content de te revoir après {last_visit_days} jours"
    else:
        visit_context = f"ravi de te revoir ! Cela faisait {last_visit_days} jours"
    
    prompt = f"""Génère un message de bienvenue court et chaleureux pour Rebecca.

Contexte :
- Prénom : {user_name}
- Moment : {time_greeting} ({emoji})
- Dernière visite : {visit_context}

RÈGLES :
- Maximum 20 mots
- Pas de données chiffrées
- Soit naturelle et encourageante
- Pas de "je suis désolée" ou "pardon"

Exemples :
- "Bonjour Rebecca ! Ravie de te revoir. Prête pour une belle journée ? ✨"
- "Bonsoir Rebecca. Content de te retrouver. Une petite victoire à célébrer ? 🌙"
- "Bon après-midi Rebecca ! J'espère que tu vas bien. Je suis là. 💖"

Retourne UNIQUEMENT le message, rien d'autre."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=60
        )
        
        message = response.choices[0].message.content.strip()
        # Nettoyer les guillemets
        message = message.replace('"', '').replace("'", "")
        
        return {"success": True, "message": message}
        
    except Exception as e:
        logger.error(f"Erreur génération message bienvenue: {e}")
        # Fallback
        fallbacks = [
            f"{emoji} {time_greeting} {user_name}. Ravie de te revoir. 💖",
            f"{emoji} {time_greeting} {user_name}. Je suis là si tu as besoin. 👑",
            f"{emoji} Content de te voir, {user_name}. Prête pour avancer ? ✨"
        ]
        import random
        return {"success": True, "message": random.choice(fallbacks)}


@app.post("/api/suggest-next-action")
async def suggest_next_action(request: Dict[str, Any]):
    """
    Suggère une action proactive basée sur l'historique et le contexte.
    """
    current_page = request.get("current_page", "dashboard")
    last_completed_task = request.get("last_completed_task")
    recent_tasks = request.get("recent_tasks", [])
    active_missions = request.get("active_missions", [])
    hour = request.get("hour", datetime.now().hour)
    last_area = request.get("last_area")  # farm, money, family, etc.
    
    # Contexte supplémentaire depuis la base
    if not active_missions and supabase:
        missions = supabase.table("missions").select("name").eq("status", "active").limit(3).execute()
        active_missions = [m["name"] for m in missions.data]
    
    # Déterminer la zone par défaut selon la page
    page_to_area = {
        "dashboard": "général",
        "farm": "ferme",
        "money-opportunities": "finances",
        "family": "famille",
        "agenda": "organisation",
        "missions-business": "business",
        "relocation": "relocalisation",
        "inbox": "brain dump",
        "chat": "conversation"
    }
    area = page_to_area.get(current_page, "général")
    
    # Si une tâche vient d'être terminée, l'utiliser
    recent_context = ""
    if last_completed_task:
        recent_context = f"Elle vient de terminer : '{last_completed_task}'"
    elif recent_tasks:
        recent_context = f"Ses dernières tâches : {', '.join(recent_tasks[:2])}"
    
    # Heure pour adapter la suggestion
    if hour < 12:
        time_context = "c'est le matin, elle a de l'énergie"
    elif hour < 18:
        time_context = "c'est l'après-midi, elle peut avancer sur une tâche moyenne"
    else:
        time_context = "c'est le soir, une petite tâche rapide ou une pause serait bien"
    
    prompt = f"""Suggère une action proactive pour Rebecca.

CONTEXTE :
- Page actuelle : {current_page}
- Zone : {area}
- Dernière activité : {recent_context if recent_context else "Aucune récente"}
- Missions actives : {', '.join(active_missions) if active_missions else 'Aucune'}
- Moment : {time_context}

RÈGLES :
- Maximum 15 mots
- Une seule suggestion
- Pas de question fermée (oui/non)
- Propose une action CONCRÈTE et RAPIDE (max 10 min)
- Sois naturelle, pas corporate

EXEMPLES :
- "On continue sur le dossier de la ferme ?"
- "Tu veux qu'on avance sur le budget du mois ?"
- "Ajouter une victoire ? Ça fait toujours du bien."
- "Faire le point sur les tâches du jour ?"
- "Prendre 5 minutes pour respirer ?"

Retourne UNIQUEMENT la suggestion, rien d'autre."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=50
        )
        
        suggestion = response.choices[0].message.content.strip()
        suggestion = suggestion.replace('"', '').replace("'", "")
        
        return {"success": True, "suggestion": suggestion}
        
    except Exception as e:
        logger.error(f"Erreur génération suggestion: {e}")
        # Fallback
        fallbacks = [
            "On avance sur une tâche ?",
            "Une petite victoire à célébrer ?",
            "Faire le point sur la journée ?",
            "Prendre une pause ?",
            "Vider ta tête dans le Brain Dump ?"
        ]
        import random
        return {"success": True, "suggestion": random.choice(fallbacks)}
